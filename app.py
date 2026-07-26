"""
AI Fingerprint Detector v3.3 — Unified PDF Report Build
تقرير PDF يعمل على ملفات PDF الأصلية فقط مع تظليل النقاط المشكوك بها فوق الأصل نفسه
تقرير فوق الصفحات الأصلية مع تظليل شفاف، واستبعاد محافظ للمراجع والاستشهادات من الحساب
"""
import re
import math
import collections
import io
import base64
from datetime import datetime

import streamlit as st

# ── مكتبات اختيارية ──────────────────────────────────────────────────────────
try:
    import fitz  # PyMuPDF
    FITZ_OK = True
except Exception:
    FITZ_OK = False

try:
    import docx
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.text.paragraph import Paragraph
    from docx.table import Table, _Cell
    from docx.document import Document as _Document
    DOCX_OK = True
except Exception:
    DOCX_OK = False

try:
    from reportlab.pdfgen import canvas as rl_canvas
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.colors import HexColor
    RLAB_OK = True
except Exception:
    RLAB_OK = False

try:
    from PIL import Image
    PIL_OK = True
except Exception:
    PIL_OK = False


# ══════════════════════════════════════════════════════════════════════════════
# FINGERPRINT ENGINE
# ══════════════════════════════════════════════════════════════════════════════

T1_PHRASES = [
    'it is worth noting','it is important to note','it is crucial to',
    'it is essential to','in the realm of','delve into','delve deeper',
    'let us delve','pave the way','at the heart of',
    'a multifaceted','a holistic approach','a comprehensive understanding',
    'nuanced understanding','multifaceted nature','transformative potential',
    'unprecedented','rapidly evolving landscape','ever-evolving',
    'plays a pivotal role','plays a crucial role','plays a fundamental role',
    'serves as a cornerstone','serves as a catalyst','serves as a testament',
    'serves as a foundation','in today\'s rapidly','in today\'s ever-changing',
    'foster a deeper understanding','harness the power of',
    'navigate the complexities','navigate the challenges',
    'unlock the potential','leverage the power of','reimagine the way',
    'underscore the importance','underscores the need','it is imperative',
    'first and foremost','state-of-the-art','best practices',
    'robust framework','comprehensive framework','evidence-based approach',
    'data-driven insights','key stakeholders','synergistic',
    'cutting-edge solutions','groundbreaking research',
    'innovative solutions','sustainable development',
    'this study examines','this paper examines','this paper investigates',
    'this paper addresses','this paper aims','this work examines',
    'these findings indicate','these findings suggest','these findings demonstrate',
    'these findings confirm','these findings reveal',
    'the results show a','the results demonstrate','the results indicate',
    'the results suggest','the results reveal','the results confirm',
    'the observed improvement','the observed increase','the observed decrease',
    'the observed pattern','the observed behavior','the observed trend',
    'when compared with recent','when compared to previous',
    'when compared with existing','when compared to traditional',
    'despite rapid progress','despite significant progress',
    'despite recent advances','despite considerable advances',
    'the feasibility of using','the feasibility of integrating',
    'the feasibility of applying','the feasibility of implementing',
    'instead of separating','rather than being treated',
    'rather than relying on','rather than focusing on',
    'emerges through repeated','emerges from physical',
    'emerges through sensorimotor','emerged through interaction',
    'at the same time, several limitations',
    'at the same time, some limitations',
    'future work should','future research should','future studies should',
    'future investigation should','future work could','future research could',
    'the main contribution','the key contribution','the primary contribution',
    'the central contribution','the core contribution',
    'in practical terms','in concrete terms',
    'further research is needed','additional research is needed',
    'has become increasingly important','has become increasingly prevalent',
    'has become increasingly complex','has become increasingly clear',
    'has emerged as a','has emerged as an',
    'highlights the importance of','highlights the need for',
    'highlights the potential of','highlights the significance of',
    'the paper focuses on','it investigates whether',
    'the study therefore supports','the study therefore confirms',
    'the study therefore demonstrates','the study therefore shows',
    'meaningful adaptive behavior','consistent learning pattern',
    'meaningful improvement','measurable improvement','measurable learning',
    'this makes the case','this makes it relevant','this makes the study',
    'not only.*but also','embodied intelligence','embodied learning',
    'resource-constrained','real-world constraints',
    'the present study','the present paper','the present work',
    'the proposed approach','the proposed method','the proposed framework',
    'the proposed system','the proposed model',
    'علاوة على ذلك','بالإضافة إلى ذلك','من ناحية أخرى',
    'في هذا السياق','في ضوء ذلك','تجدر الإشارة',
    'لا يمكن إغفال','لا يمكن إنكار','من المسلم به',
    'في الختام','خلاصة القول','مما سبق يتضح',
    'منهجية متكاملة','نهج شامل','رؤية متكاملة',
    'في عالمنا المعاصر','في ظل التطورات المتسارعة',
    'يُعدّ من أبرز','يُمثّل ركيزة أساسية',
    'يكتسب أهمية بالغة','يحتل مكانة محورية',
    'تحديات جمّة','فرص واعدة','آفاق رحبة',
    'متعدد الأوجه','متعدد الأبعاد','شامل ومتكامل',
]

T2_PATTERNS = [
    r'\bthis (?:study|paper|article|research|work) (?:examines?|investigates?|explores?|addresses?|presents?|demonstrates?|aims? to|argues?|proposes?|introduces?)\b',
    r'\bthese findings (?:indicate|suggest|demonstrate|show|reveal|confirm|support|highlight|underscore)\b',
    r'\bthe (?:results?|findings?|data|evidence|analysis) (?:show|suggest|indicate|demonstrate|reveal|confirm|highlight)\b',
    r'\bthe (?:observed|noted|measured|recorded|reported) (?:improvement|increase|decrease|reduction|pattern|trend|behavior|change|difference)\b',
    r'\bthis (?:makes?|renders?) (?:the|it|this) (?:study|approach|work|finding|result|system|prototype|case|paper|contribution)\b',
    r'\bat the same time,? (?:several?|some|various|important|key|notable) (?:limitations?|challenges?|issues?|concerns?|caveats?)\b',
    r'\b(?:meaningful|measurable|clear|consistent|significant|substantial|notable) (?:improvement|learning|progress|adaptation|behavior|pattern|results?|gains?|increase|decrease)\b',
    r'\bwhen compared (?:with|to) (?:recent|previous|existing|traditional|other|prior|current|state-of-the-art)\b',
    r'\bfuture (?:work|research|studies?|investigation|directions?) (?:should|could|may|might|will|can) (?:explore|examine|investigate|test|consider|address|focus|extend|incorporate)\b',
    r'\bdespite (?:rapid|significant|recent|considerable|substantial|great|notable) (?:progress|advances?|development|growth|improvement|interest)\b',
    r'\bthe (?:feasibility|practicality|effectiveness|utility|applicability|viability) of (?:using|applying|integrating|implementing|employing|adopting)\b',
    r'\binstead of (?:separating|relying on|using|treating|employing|assuming)\b',
    r'\brather than (?:being treated|relying on|focusing on|separating|using|assuming|requiring)\b',
    r'\b(?:emerges?|emerged?) (?:from|through|via) (?:repeated?|real|physical|sensorimotor|direct|online|trial)\b',
    r'\bthis (?:approach|method|framework|system|design|setup|configuration|architecture) (?:allows?|enables?|facilitates?|supports?|provides?|offers?|ensures?)\b',
    r'\bthe (?:study|paper|experiment|research|work) (?:therefore|thus|hence|consequently) (?:supports?|confirms?|demonstrates?|shows?|suggests?|indicates?)\b',
    r'\bdespit(?:e|ing) (?:these|the|its|their|such|various) (?:challenges?|limitations?|drawbacks?|obstacles?|constraints?)\b',
    r'\bthis (?:article|paper|essay|study|work|chapter|section) (?:aims?|seeks?|explores?|examines?|discusses?|presents?|provides?|argues?|demonstrates?|shows?)\b',
    r'\ba (?:comprehensive|thorough|detailed|rigorous|systematic|careful|nuanced|in-depth) (?:analysis|examination|overview|review|investigation|evaluation|assessment|study)\b',
    r'\b(?:significant|substantial|considerable|notable|marked|profound|dramatic) (?:impact|effect|influence|improvement|shift|difference|change|reduction|increase)\b',
    r'\bthe (?:main|primary|key|central|core|fundamental|principal) (?:contribution|value|finding|advantage|limitation|challenge|goal|aim|objective|novelty) of\b',
    r'\b(?:plays?|played?) a (?:vital|crucial|key|central|pivotal|important|major|fundamental|critical|essential|significant) role in\b',
    r'\b(?:moreover|furthermore|additionally|besides),?\s+(?:the|this|its|these|such|various|several)\b',
    r'\b(?:has|have) (?:significantly|greatly|substantially|considerably|dramatically|rapidly|increasingly) (?:grown|expanded|increased|developed|improved|enhanced|advanced|transformed|evolved)\b',
    r'\b(?:is|are) characterized by (?:a|an|the|its|their)\b',
    r'\b(?:serves?|acts?) as (?:a|an|the) (?:major|key|primary|central|critical|vital|important|significant|crucial|fundamental|useful|practical)\b',
    r'\b(?:driven by|fueled by|powered by|propelled by|guided by|motivated by) (?:the|its|their|a|an)\b',
    r'\bin (?:conclusion|summary|closing|summation),?\s+(?:it is|we can|this|the|these)\b',
    r'\bit (?:is|was) (?:therefore|thus|hence|consequently) (?:demonstrated?|confirmed?|shown|established|evident|clear|argued|suggested)\b',
    r'\bthe (?:present|current|proposed|described|outlined|discussed) (?:study|work|paper|approach|method|system|framework|research)\b',
    r'\bby (?:examining|exploring|analyzing|investigating|considering|addressing|comparing|evaluating) (?:the|these|this|how|why|what|whether)\b',
    r'\bit (?:should|must|cannot|can) be (?:noted|emphasized|stressed|acknowledged|recognized|mentioned|highlighted) that\b',
    r'\bthe (?:importance|significance|relevance|role|impact|value|potential|need|necessity|utility) of (?:this|these|the|such|incorporating|using|applying)\b',
    r'\b(?:not only).{3,60}\b(?:but also|but it also|but they also|but also helps|but also provides)\b',
    r'\bthis (?:narrow(?:er)?|broad(?:er)?|simpl(?:e|er)|more (?:complex|focused|practical|realistic)) (?:focus|scope|approach|perspective|question|contribution)\b',
    r'\b(?:overall|in general|taken together|collectively),?\s+(?:the|these|this|our|the results)\b',
    r'\baddress(?:ing|es|ed)? (?:the|these|this|key|critical|important|pressing|growing|real) (?:issue|challenge|problem|concern|need|question|gap|limitation|challenge)\b',
    r'\bيُسهم في\b|\bيُعزز من\b|\bيُرسخ\b|\bيُفضي إلى\b',
    r'\bوتجدر الإشارة\b|\bتجدر الإشارة إلى\b',
    r'\bمن خلال(?:ه|ها|هم)?\s+(?:يمكن|نستطيع|تتضح|يتضح)\b',
    r'\bوفي هذا الإطار\b|\bضمن هذا السياق\b',
]


# تجهيز النص قبل التحليل: استبعاد المراجع والاستشهادات داخل المتن
# ══════════════════════════════════════════════════════════════════════════════

_ARABIC_DIACRITICS_RE = re.compile(r'[\u0610-\u061A\u064B-\u065F\u0670\u06D6-\u06ED]')


def _count_text_words(text: str) -> int:
    return len(re.findall(r'\b[a-zA-Z\u0600-\u06FF]+\b', text or ''))


def _normalize_reference_heading(text: str) -> str:
    """توحيد عنوان قسم المراجع قبل مطابقته."""
    s = _ARABIC_DIACRITICS_RE.sub('', text or '')
    s = s.replace('ـ', '')
    s = re.sub(r'\s+', ' ', s).strip()
    s = re.sub(r'^[\s\-–—:؛،,.()\[\]]+', '', s)
    s = re.sub(r'[\s\-–—:؛،,.()\[\]]+$', '', s)
    return s


REF_SECTION_HEADERS = re.compile(
    r'''^(?:
        (?:(?:chapter|section|appendix|الفصل|القسم)\s+)?(?:\d+(?:\.\d+)*)?\s*[:.\-–—]?\s*
    )?(?:
        references?|reference\s+list|bibliography|works\s+cited|literature\s+cited|sources?
        |المراجع(?:\s+(?:العربية|الاجنبية|الأجنبية))?
        |قائمة\s+(?:المراجع|المصادر)
        |المصادر(?:\s+والمراجع)?
        |المراجع\s+والمصادر
    )$''',
    re.I | re.X
)


def _is_reference_header(text: str) -> bool:
    """
    يكتشف عنوان قسم المراجع حتى لو سبقه رقم فصل أو تبعته نقطتان.
    لا يعتمد على وجود كلمة References داخل فقرة طويلة.
    """
    for raw_line in re.split(r'[\r\n]+', text or ''):
        line = _normalize_reference_heading(raw_line)
        if not line:
            continue
        if len(line.split()) <= 8 and REF_SECTION_HEADERS.fullmatch(line):
            return True
    return False


REF_LINE_PATTERNS = [
    r'^\s*\[\s*\d+(?:\s*[-–,]\s*\d+)*\s*\]',
    r'^\s*\d{1,4}\s*[.)]\s+',
    r'\bdoi\s*:\s*10\.\d{4,9}/\S+',
    r'\bhttps?://\S+',
    r'\b(?:ISBN|ISSN)\b',
    r'\b(?:Vol\.?|Volume)\s*\d+',
    r'\b(?:No\.?|Issue)\s*\d+',
    r'\bpp?\.\s*\d+(?:\s*[-–]\s*\d+)?',
    r'\b(?:Journal|Proceedings|Transactions|Conference|Press|Publisher)\b',
    r'\bet\s+al\.?',
]


def _reference_line_score(line: str) -> float:
    """درجة احتمالية أن السطر مدخل ببليوجرافي وليس فقرة من المتن."""
    s = re.sub(r'\s+', ' ', (line or '').strip())
    if not s:
        return 0.0

    score = 0.0
    if re.search(r'^\s*\[\s*\d+(?:\s*[-–,]\s*\d+)*\s*\]', s):
        score += 2.5
    if re.search(r'^\s*\d{1,4}\s*[.)]\s+', s):
        score += 1.8
    if re.search(r'\b(?:18|19|20)\d{2}[a-z]?\b', s, re.I):
        score += 0.8
    if re.search(r'\bdoi\s*:|doi\.org/10\.|\b10\.\d{4,9}/\S+', s, re.I):
        score += 2.2
    if re.search(r'https?://', s, re.I):
        score += 1.7
    if re.search(r'\bet\s+al\.?\b', s, re.I):
        score += 1.0
    if re.search(r'\b(?:Journal|Proceedings|Transactions|Conference|Press|Publisher|Springer|Elsevier|Wiley|IEEE|ACM)\b', s, re.I):
        score += 1.1
    if re.search(r'\b(?:Vol\.?|Volume|No\.?|Issue|pp?\.)\s*\d+', s, re.I):
        score += 1.0
    if re.search(r'\b(?:ISBN|ISSN)\b', s, re.I):
        score += 1.5
    if re.search(r'^[A-ZÀ-ÖØ-Ý][A-Za-zÀ-ÖØ-öø-ÿ\'’-]+,\s*(?:[A-Z]\.?\s*){1,4}', s):
        score += 1.2
    if re.search(r'\([12]\d{3}[a-z]?\)', s, re.I):
        score += 0.8
    if s.count(',') >= 2 and s.count('.') >= 2:
        score += 0.5
    if len(s) >= 45:
        score += 0.2
    return score


def _is_reference_line(line: str) -> bool:
    """
    تصنيف محافظ للسطر المرجعي. الأنماط القوية تكفي منفردة،
    وإلا نحتاج اجتماع أكثر من علامة ببليوجرافية لتجنب حذف متن الباحث.
    """
    s = re.sub(r'\s+', ' ', (line or '').strip())
    if not s:
        return False
    if _is_reference_header(s):
        return True
    if re.search(r'^\s*\[\s*\d+(?:\s*[-–,]\s*\d+)*\s*\]', s):
        return True
    if re.search(r'\bdoi\s*:|doi\.org/10\.|\b10\.\d{4,9}/\S+', s, re.I):
        return True
    return _reference_line_score(s) >= 2.6


_INLINE_NUMERIC_CITATION_RE = re.compile(
    r'\[(?:\s*\d+\s*(?:[-–]\s*\d+)?\s*)(?:[,;]\s*\d+\s*(?:[-–]\s*\d+)?\s*)*\]'
)
_INLINE_AUTHOR_YEAR_RE = re.compile(
    r'''\(
        (?=[^()]{0,180}\b(?:18|19|20)\d{2}[a-z]?\b)
        (?:[^()]*?[A-ZÀ-ÖØ-Ý\u0600-\u06FF][^()]*?)
        (?:18|19|20)\d{2}[a-z]?
        (?:\s*[,;]\s*[^()]*?(?:18|19|20)\d{2}[a-z]?)*
    \)''',
    re.I | re.X
)
_INLINE_YEAR_ONLY_RE = re.compile(r'\((?:18|19|20)\d{2}[a-z]?\)')
_INLINE_URL_RE = re.compile(r'https?://\S+|www\.\S+', re.I)
_INLINE_DOI_RE = re.compile(r'\b(?:doi\s*:\s*)?10\.\d{4,9}/\S+', re.I)


def _strip_inline_citations(text: str) -> str:
    """يحذف علامات الاستشهاد فقط ويُبقي الجملة العلمية نفسها للتحليل."""
    s = text or ''
    s = _INLINE_NUMERIC_CITATION_RE.sub(' ', s)
    s = _INLINE_AUTHOR_YEAR_RE.sub(' ', s)
    s = _INLINE_YEAR_ONLY_RE.sub(' ', s)
    s = _INLINE_DOI_RE.sub(' ', s)
    s = _INLINE_URL_RE.sub(' ', s)
    s = re.sub(r'\s+', ' ', s)
    return s.strip()


def _find_unheaded_reference_tail(lines) -> int | None:
    """
    إذا غاب عنوان المراجع، يبحث فقط في الجزء الأخير من المستند عن كتلة
    كثيفة من المداخل الببليوجرافية. هذا يمنع إسقاط فقرات وسط البحث.
    """
    nonempty = [(i, line.strip()) for i, line in enumerate(lines) if line.strip()]
    if len(nonempty) < 6:
        return None

    earliest_pos = max(0, int(len(nonempty) * 0.50))
    for pos in range(earliest_pos, len(nonempty)):
        window = nonempty[pos:pos + 6]
        if len(window) < 3:
            continue
        ref_hits = sum(1 for _, line in window if _is_reference_line(line))
        if ref_hits < 3:
            continue
        first_ref_offset = next(
            (j for j, (_, line) in enumerate(window) if _is_reference_line(line)),
            None
        )
        if first_ref_offset is None:
            continue
        tail = nonempty[pos + first_ref_offset:]
        density = sum(1 for _, line in tail if _is_reference_line(line)) / max(len(tail), 1)
        if density >= 0.55:
            return tail[0][0]
    return None


def _prepare_analysis_text(text: str) -> dict:
    """
    يفصل متن البحث عن قسم المراجع ويزيل الاستشهادات داخل المتن.
    يعيد إحصاءات الاستبعاد لتسهيل المراجعة والاختبار.
    """
    raw_text = text or ''
    lines = raw_text.replace('\r\n', '\n').replace('\r', '\n').split('\n')

    ref_start = None
    for idx, line in enumerate(lines):
        if _is_reference_header(line):
            ref_start = idx
            break

    if ref_start is None:
        ref_start = _find_unheaded_reference_tail(lines)

    if ref_start is None:
        main_lines = lines
        ref_lines = []
        header_found = False
    else:
        main_lines = lines[:ref_start]
        ref_lines = lines[ref_start:]
        header_found = any(_is_reference_header(line) for line in ref_lines[:3])

    main_text_raw = '\n'.join(main_lines)
    analysis_text = _strip_inline_citations(main_text_raw)
    references_text = '\n'.join(ref_lines)

    return {
        'analysis_text': analysis_text,
        'main_text_raw': main_text_raw,
        'references_text': references_text,
        'reference_section_found': ref_start is not None,
        'reference_header_found': header_found,
        'original_words': _count_text_words(raw_text),
        'analyzed_words': _count_text_words(analysis_text),
        'reference_words_excluded': _count_text_words(references_text),
        'inline_citation_words_removed': max(
            0,
            _count_text_words(main_text_raw) - _count_text_words(analysis_text)
        ),
    }


LOW_SCORE_HIGHLIGHT_THRESHOLD = 20.0
MASKED_PERCENT_LABEL = "*%"


def _coerce_percentage(value) -> float | None:
    """تحويل آمن للنسبة مع رفض القيم غير الرقمية أو غير المنتهية."""
    try:
        pct = float(value)
    except (TypeError, ValueError):
        return None
    if not math.isfinite(pct):
        return None
    return pct


def _is_masked_percentage(value) -> bool:
    """السياسة الموحدة: كل نتيجة من 0% إلى 20% شاملة تُخفى تماماً."""
    pct = _coerce_percentage(value)
    return pct is None or 0.0 <= pct <= LOW_SCORE_HIGHLIGHT_THRESHOLD


def _format_percentage(value: float, mask_low: bool = True) -> str:
    """يعرض 0–20% بالشكل *%، ولا يكشف القيمة الرقمية الحقيقية."""
    pct = _coerce_percentage(value)
    if pct is None:
        return MASKED_PERCENT_LABEL
    if mask_low and _is_masked_percentage(pct):
        return MASKED_PERCENT_LABEL
    rounded_pct = int(math.floor(pct + 0.5))
    # إذا كانت النتيجة أعلى من 20% لكن تقريبها سيعرض 20%، نعرض منزلة عشرية
    # حتى لا يظهر التقرير 20% مع وجود تظليل.
    if mask_low and pct > LOW_SCORE_HIGHLIGHT_THRESHOLD and rounded_pct <= 20:
        return f'{pct:.1f}%'
    return f'{rounded_pct}%'


def _format_human_percentage(ai_value: float, human_value: float) -> str:
    """لا يعرض النسبة البشرية المعكوسة عندما تكون نسبة AI مخفية."""
    if _is_masked_percentage(ai_value):
        return 'مرتفع'
    human_pct = _coerce_percentage(human_value)
    if human_pct is None:
        return 'غير محدد'
    return f'{human_pct:.0f}%'


def _highlighting_allowed(ai_percentage: float) -> bool:
    """
    يسمح بالتظليل فقط عندما تكون نسبة AI أكبر من 20% فعلياً.
    من 0% إلى 20% شاملة: لا PDF annotations، ولا DOCX shading، ولا تظليل نصي.
    """
    pct = _coerce_percentage(ai_percentage)
    return pct is not None and pct > LOW_SCORE_HIGHLIGHT_THRESHOLD


def _visual_percentage(value: float) -> float | None:
    """
    يعيد موضعاً بصرياً للنسبة فقط إذا كانت غير مخفية.
    يمنع شريط الواجهة ومؤشر الغلاف من كشف قيمة 0–20% بصورة غير مباشرة.
    """
    pct = _coerce_percentage(value)
    if pct is None or _is_masked_percentage(pct):
        return None
    return max(0.0, min(100.0, pct))


def _collect_unique_pattern_matches(text: str, patterns) -> tuple[int, list]:
    """
    يجمع مطابقات الأنماط مع إزالة التطابقات المتداخلة التي تصف العبارة نفسها.
    هذا يمنع احتساب جملة واحدة مرتين لأن نمطين عامين التقطا الجزء نفسه.
    """
    raw_matches = []
    for pat_idx, pat in enumerate(patterns):
        try:
            for match in re.finditer(pat, text, re.I):
                start, end = match.span()
                if end > start:
                    raw_matches.append((start, end, pat_idx, pat))
        except Exception:
            continue

    raw_matches.sort(key=lambda x: (x[0], -(x[1] - x[0]), x[2]))
    unique = []
    for item in raw_matches:
        start, end, pat_idx, pat = item
        duplicate = False
        for kept in unique:
            ks, ke, _, _ = kept
            overlap = max(0, min(end, ke) - max(start, ks))
            smaller = max(1, min(end - start, ke - ks))
            if overlap / smaller >= 0.60:
                duplicate = True
                break
        if not duplicate:
            unique.append(item)

    pattern_counts = collections.Counter(item[2] for item in unique)
    matched_patterns = [
        (patterns[idx], pattern_counts[idx])
        for idx in sorted(pattern_counts)
    ]
    return len(unique), matched_patterns


# ══════════════════════════════════════════════════════════════════════════════
# FINGERPRINT ENGINE
# ══════════════════════════════════════════════════════════════════════════════

def _ai_fingerprint_score(text: str) -> dict:
    prep = _prepare_analysis_text(text)
    analyzed_text = prep['analysis_text']
    tl = analyzed_text.lower()
    words   = re.findall(r'\b[a-zA-Z\u0600-\u06FF]+\b', analyzed_text)
    sents   = [s.strip() for s in re.split(r'[.!?؟]+', analyzed_text) if len(s.strip()) > 8]
    n_words = max(len(words), 1)
    n_sents = max(len(sents), 1)

    t1_hits = []
    for p in T1_PHRASES:
        try:
            if '.*' in p:
                if re.search(p, tl, re.I):
                    t1_hits.append(p)
            elif p in tl:
                t1_hits.append(p)
        except Exception:
            continue
    t1_count = len(t1_hits)

    if   t1_count >= 20: t1_score = min(0.97, 0.90 + (t1_count - 20) * 0.007)
    elif t1_count >= 14: t1_score = 0.80 + (t1_count - 14) * 0.0167
    elif t1_count >= 8:  t1_score = 0.55 + (t1_count - 8)  * 0.0417
    elif t1_count >= 4:  t1_score = 0.30 + (t1_count - 4)  * 0.0625
    elif t1_count >= 2:  t1_score = 0.12 + (t1_count - 2)  * 0.09
    elif t1_count == 1:  t1_score = 0.07
    else:                t1_score = 0.02

    t2_total, t2_matched = _collect_unique_pattern_matches(tl, T2_PATTERNS)

    hit_ratio = t2_total / n_sents
    if   hit_ratio >= 0.35: t2_score = min(0.97, 0.90 + (hit_ratio - 0.35) * 1.00)
    elif hit_ratio >= 0.22: t2_score = 0.75 + (hit_ratio - 0.22) * 1.15
    elif hit_ratio >= 0.12: t2_score = 0.48 + (hit_ratio - 0.12) * 2.70
    elif hit_ratio >= 0.06: t2_score = 0.18 + (hit_ratio - 0.06) * 5.00
    elif hit_ratio >= 0.02: t2_score = 0.05 + (hit_ratio - 0.02) * 3.25
    else:                   t2_score = 0.02

    avg_len = sum(len(s.split()) for s in sents) / n_sents if sents else 0.0
    if   15 <= avg_len <= 26: style_score = 0.72
    elif 12 <= avg_len <= 28: style_score = 0.50
    elif 10 <= avg_len <= 30: style_score = 0.28
    else:                     style_score = 0.10 if words else 0.0

    self_ref = len(re.findall(
        r'\b(?:the|this)\s+(?:study|paper|article|experiment|research|work|prototype|system|approach|method|model|framework|analysis)\b',
        tl
    ))
    _STOP = {
        'which','about','their','there','these','those','where','while',
        'after','before','since','would','could','should','might','other',
        'study','paper','using','based','also','been','have','with','from',
        'that','this','than','into','upon','each','both','such','when',
        'then','they','were','some','more','most','over','only','well',
        'same','even','much','many','within','through','between','during',
        'because','however','therefore','system','method','approach',
        'model','results','section','table','figure',
    }
    _fw = collections.Counter(
        w for w in re.findall(r'\b[a-z]{5,}\b', tl) if w not in _STOP
    )
    top_kw_count = _fw.most_common(1)[0][1] if _fw else 0
    top_kw_word  = _fw.most_common(1)[0][0] if _fw else ''

    sr_ratio = self_ref / n_sents
    kw_ratio = top_kw_count / n_words
    struct_boost = 0.0
    if sr_ratio >= 0.12:    struct_boost += 0.03
    elif sr_ratio >= 0.07:  struct_boost += 0.015
    if kw_ratio >= 0.012:   struct_boost += 0.02
    elif kw_ratio >= 0.007: struct_boost += 0.01

    base = (t1_score * 0.45) + (t2_score * 0.40) + (style_score * 0.15)
    raw_final = max(0.0, min(1.0, base + struct_boost))

    # بوابات محافظة لتقليل الإدانة الخاطئة: الأسلوب وحده لا يكفي.
    evidence_units = max(t1_count, min(t2_total, 8))
    evidence_families = int(t1_count > 0) + int(t2_total > 0) + int(self_ref > 0)
    final = raw_final
    if evidence_units == 0:
        final = min(final, 0.12)
    elif evidence_units == 1:
        final = min(final, 0.24)
    elif evidence_units == 2 and evidence_families <= 1:
        final = min(final, 0.36)

    # النصوص القصيرة لا تسمح بحكم قوي؛ نخفض الثقة والدرجة بدلاً من ظلم الباحث.
    if n_words < 80:
        final = min(final * 0.70, 0.45)
    elif n_words < 150:
        final = min(final * 0.85, 0.62)
    final = max(0.0, min(1.0, final))

    length_factor = min(1.0, n_words / 800.0)
    signal_factor = min(1.0, evidence_units / 12.0)
    diversity_factor = min(1.0, evidence_families / 3.0)
    confidence = max(0.05, min(0.99,
        0.15 + (0.40 * length_factor) + (0.30 * signal_factor) + (0.15 * diversity_factor)
    ))

    if   final >= 0.80: decision, verdict, color = 'STRONG_SIGNALS',  'مؤشرات أسلوبية قوية — يلزم تحقق بشري', '#c0392b'
    elif final >= 0.62: decision, verdict, color = 'HIGH_SIGNALS',    'مؤشرات مرتفعة — يحتاج مراجعة',          '#e67e22'
    elif final >= 0.42: decision, verdict, color = 'MIXED',           'مختلط — لا يكفي للحكم منفرداً',          '#f39c12'
    elif final >= 0.22: decision, verdict, color = 'LIMITED_SIGNALS', 'مؤشرات محدودة',                          '#27ae60'
    else:               decision, verdict, color = 'NO_CLEAR_SIGNAL', 'لا توجد مؤشرات كافية للحكم',             '#2ecc71'

    percentage = round(final * 100, 1)
    return {
        'score':        round(final, 4),
        'raw_score_before_safety_gates': round(raw_final, 4),
        'percentage':   percentage,
        'display_percentage': _format_percentage(percentage),
        'human_score':  round((1.0 - final) * 100, 1),
        'confidence':   round(confidence, 4),
        'confidence_percentage': round(confidence * 100, 1),
        'decision':     decision,
        'verdict':      verdict,
        'color':        color,
        't1_count':     t1_count,
        't1_score':     round(t1_score, 4),
        't1_hits':      t1_hits[:12],
        't2_total':     t2_total,
        't2_hit_ratio': round(hit_ratio, 4),
        't2_score':     round(t2_score, 4),
        't2_matched':   [(m[1], m[0][:60]) for m in t2_matched[:8]],
        'style_score':  round(style_score, 4),
        'avg_sent_len': round(avg_len, 1),
        'struct_boost': round(struct_boost, 4),
        'self_ref':     self_ref,
        'top_kw':       top_kw_word,
        'top_kw_count': top_kw_count,
        'n_words':      n_words,
        'n_sents':      n_sents,
        'original_words': prep['original_words'],
        'reference_words_excluded': prep['reference_words_excluded'],
        'inline_citation_words_removed': prep['inline_citation_words_removed'],
        'reference_section_found': prep['reference_section_found'],
        'reference_header_found': prep['reference_header_found'],
        'analysis_text': analyzed_text,
        'breakdown': {
            'T1 (إشارات لغوية)': round(t1_score * 0.45, 4),
            'T2 (أنماط جملة)':   round(t2_score * 0.40, 4),
            'Style (أسلوب)':     round(style_score * 0.15, 4),
            'Struct (بنية)':     round(struct_boost, 4),
        },
    }



def _score_paragraph_for_highlight(para: str) -> dict:
    """
    يحسب درجة محلية محافظة للفقرة بعد حذف الاستشهادات.
    لا يسمح للأسلوب أو طول الجملة وحدهما بإنشاء تظليل.
    """
    raw_text = (para or '').strip()
    if not raw_text or _is_reference_header(raw_text) or _is_reference_line(raw_text):
        return {
            'text': '', 'word_count': 0, 't1_count': 0, 't2_total': 0,
            'self_ref': 0, 'style_score': 0.0, 'evidence_score': 0.0,
            'density': 0.0, 'explicit_signal': False,
        }

    text = _strip_inline_citations(raw_text)
    tl = text.lower()
    words = re.findall(r'\b[a-zA-Z\u0600-\u06FF]+\b', text)
    sents = [s.strip() for s in re.split(r'[.!?؟]+', text) if len(s.strip()) > 8]
    n_words = len(words)
    n_sents = max(len(sents), 1)

    if n_words < 4:
        return {
            'text': text, 'word_count': n_words, 't1_count': 0, 't2_total': 0,
            'self_ref': 0, 'style_score': 0.0, 'evidence_score': 0.0,
            'density': 0.0, 'explicit_signal': False,
        }

    t1_count = 0
    for p in T1_PHRASES:
        try:
            if '.*' in p:
                t1_count += len(re.findall(p, tl, re.I))
            else:
                t1_count += tl.count(p)
        except Exception:
            continue

    t2_total, _ = _collect_unique_pattern_matches(tl, T2_PATTERNS)

    self_ref = len(re.findall(
        r'\b(?:the|this)\s+(?:study|paper|article|experiment|research|work|prototype|system|approach|method|model|framework|analysis)\b',
        tl
    ))

    avg_len = sum(len(s.split()) for s in sents) / n_sents if sents else n_words
    if 15 <= avg_len <= 26:
        style_score = 0.72
    elif 12 <= avg_len <= 28:
        style_score = 0.50
    elif 10 <= avg_len <= 30:
        style_score = 0.28
    else:
        style_score = 0.10

    hit_ratio = t2_total / n_sents
    t1_norm = min(1.0, t1_count / 2.0)
    t2_norm = min(1.0, hit_ratio / 0.20) if hit_ratio > 0 else 0.0
    struct_norm = min(1.0, self_ref / n_sents)

    # الإشارات الصريحة فقط هي التي تفتح باب التظليل.
    explicit_signal = (t1_count > 0) or (t2_total > 0)
    evidence_score = (
        (t1_norm * 0.50) +
        (t2_norm * 0.38) +
        (style_score * 0.07) +
        (struct_norm * 0.05)
    )
    if not explicit_signal:
        evidence_score = 0.0

    # فقرة بإشارة واحدة شائعة لا تُرفع تلقائياً لدرجة قوية.
    if (t1_count + t2_total) == 1:
        evidence_score = min(evidence_score, 0.34)

    density = evidence_score / max(math.sqrt(n_words), 1.0)

    return {
        'text': text,
        'word_count': n_words,
        't1_count': t1_count,
        't2_total': t2_total,
        'self_ref': self_ref,
        'style_score': round(style_score, 4),
        'evidence_score': round(evidence_score, 6),
        'density': round(density, 6),
        'explicit_signal': explicit_signal,
    }


def _merge_adjacent_blocks(blocks, gap_threshold: float = 18.0):
    """
    يدمج الـ blocks النصية المتجاورة رأسياً في فقرة واحدة حقيقية.
    يعيد قائمة من الفقرات، كل فقرة = قائمة line_data مع bbox كل سطر على حدة.
    """
    text_blocks = [b for b in blocks if b.get("type", 0) == 0]
    text_blocks.sort(key=lambda b: b["bbox"][1])

    paragraphs = []

    for block in text_blocks:
        lines = block.get("lines", [])
        line_data = []
        for line in lines:
            spans = line.get("spans", [])
            txt = "".join((sp.get("text", "") or "") for sp in spans).strip()
            if txt:
                line_data.append({
                    'text':  txt,
                    'bbox':  fitz.Rect(line["bbox"]),
                    'y0':    line["bbox"][1],
                    'y1':    line["bbox"][3],
                    # نحفظ span-level data لتحديد نقطة بدء/نهاية دقيقة
                    'spans': line.get("spans", []),
                })
        if not line_data:
            continue

        if paragraphs:
            last_para = paragraphs[-1]
            last_y1   = last_para['lines'][-1]['y1']
            this_y0   = line_data[0]['y0']
            gap       = this_y0 - last_y1
            last_x0   = last_para['lines'][0]['bbox'].x0
            this_x0   = line_data[0]['bbox'].x0
            if gap <= gap_threshold and abs(last_x0 - this_x0) < 40.0:
                last_para['lines'].extend(line_data)
                last_para['full_text'] = " ".join(d['text'] for d in last_para['lines'])
                continue

        paragraphs.append({
            'lines':     line_data,
            'full_text': " ".join(d['text'] for d in line_data),
        })

    return paragraphs


def _sentence_start_x(line_data_item) -> float:
    """
    يجد الـ x الذي تبدأ منه أول جملة في السطر.
    إذا كان السطر يبدأ بحرف كبير بعد نقطة سابقة → نقطة البداية الفعلية للجملة.
    في حالات أخرى نأخذ x0 السطر كما هو.
    """
    return line_data_item['bbox'].x0


def _punct_end_x(line_data_item) -> float:
    """
    يجد الـ x الذي تنتهي عنده آخر علامة ترقيم (. أو ،) في السطر.
    نبحث في الـ spans عن آخر حرف قبل المسافات النهائية.
    إذا لم يوجد → نأخذ x1 السطر كاملاً.
    """
    txt = line_data_item['text'].rstrip()
    if not txt:
        return line_data_item['bbox'].x1
    # آخر حرف هو علامة ترقيم؟
    if re.search(r'[.،,;؟?!]\s*$', txt):
        # نأخذ x1 كاملاً (الترقيم هو الحرف الأخير)
        return line_data_item['bbox'].x1
    return line_data_item['bbox'].x1


def _build_line_rects_turnitin(chunk_lines):
    """
    يبني قائمة من المستطيلات بأسلوب Turnitin:
    - كل سطر له مستطيله المستقل (سطر تلو الآخر)
    - السطر الأول: من بداية الجملة (x0 السطر)
    - السطر الأخير: حتى آخر علامة ترقيم (x1 أو موضع الترقيم)
    - الأسطر الوسطى: من x0 إلى x1 كاملاً
    """
    rects = []
    n = len(chunk_lines)
    for i, ld in enumerate(chunk_lines):
        b = ld['bbox']
        x0 = b.x0
        x1 = b.x1
        y0 = b.y0
        y1 = b.y1

        if i == 0:
            # السطر الأول: من بداية الجملة
            x0 = _sentence_start_x(ld)
        if i == n - 1:
            # السطر الأخير: حتى الترقيم
            x1 = _punct_end_x(ld)

        rects.append(fitz.Rect(x0, y0, x1, y1))
    return rects


def _collect_page_block_candidates(page, in_refs_flag: bool):
    """
    يجمع مرشحات التظليل من متن الصفحة فقط.
    يعيد أيضاً عدد جميع كلمات المتن القابلة للتحليل، وليس كلمات المرشحات فقط.
    """
    candidates = []
    in_refs = in_refs_flag
    analyzable_words = 0

    MIN_LINES = 4
    MAX_LINES = 5
    MIN_LOCAL_EVIDENCE = 0.20

    try:
        page_dict = page.get_text("dict")
        raw_blocks = page_dict.get("blocks", [])
    except Exception:
        raw_blocks = []

    paragraphs = _merge_adjacent_blocks(raw_blocks)

    for para in paragraphs:
        original_lines = para['lines']
        if not original_lines:
            continue

        # قد يندمج عنوان المراجع مع block سابق أو لاحق؛ نعالج ما قبله فقط.
        header_idx = next(
            (i for i, ld in enumerate(original_lines) if _is_reference_header(ld.get('text', ''))),
            None
        )
        set_refs_after = header_idx is not None

        if in_refs:
            continue

        if header_idx == 0:
            in_refs = True
            continue
        elif header_idx is not None:
            line_data = original_lines[:header_idx]
        else:
            line_data = original_lines

        btext = " ".join(d['text'] for d in line_data).strip()
        if not btext:
            if set_refs_after:
                in_refs = True
            continue

        if _is_reference_line(btext):
            if set_refs_after:
                in_refs = True
            continue

        clean_btext = _strip_inline_citations(btext)
        analyzable_words += _count_text_words(clean_btext)
        total_lines = len(line_data)

        if total_lines < MIN_LINES:
            local = _score_paragraph_for_highlight(btext)
            if (
                local['word_count'] < 4
                or not local['explicit_signal']
                or local['evidence_score'] < MIN_LOCAL_EVIDENCE
            ):
                if set_refs_after:
                    in_refs = True
                continue

            line_rects = _build_line_rects_turnitin(line_data)
            combined = line_rects[0]
            for r in line_rects[1:]:
                combined = combined | r
            candidates.append({
                'rects': line_rects,
                'rect': combined,
                'text': btext,
                'word_count': local['word_count'],
                'score': local['evidence_score'],
                'density': local['density'],
                't1_count': local['t1_count'],
                't2_total': local['t2_total'],
                'self_ref': local['self_ref'],
                'explicit_signal': local['explicit_signal'],
                'line_count': total_lines,
            })
            if set_refs_after:
                in_refs = True
            continue

        best_candidate = None
        for end_i in range(MIN_LINES - 1, total_lines):
            end_text = line_data[end_i]['text'].rstrip()
            ends_with_punct = bool(re.search(r'[.،,;؟?!]\s*$', end_text))

            for chunk_len in range(MIN_LINES, MAX_LINES + 1):
                start_i = end_i - chunk_len + 1
                if start_i < 0:
                    continue
                chunk_lines = line_data[start_i:end_i + 1]
                chunk_text = " ".join(d['text'] for d in chunk_lines)
                local = _score_paragraph_for_highlight(chunk_text)

                if (
                    local['word_count'] < 4
                    or not local['explicit_signal']
                    or local['evidence_score'] < MIN_LOCAL_EVIDENCE
                ):
                    continue

                line_rects = _build_line_rects_turnitin(chunk_lines)
                combined = line_rects[0]
                for r in line_rects[1:]:
                    combined = combined | r

                score_val = local['evidence_score'] * (1.08 if ends_with_punct else 1.0)
                if best_candidate is None or score_val > best_candidate['score']:
                    best_candidate = {
                        'rects': line_rects,
                        'rect': combined,
                        'text': chunk_text,
                        'word_count': local['word_count'],
                        'score': score_val,
                        'density': local['density'],
                        't1_count': local['t1_count'],
                        't2_total': local['t2_total'],
                        'self_ref': local['self_ref'],
                        'explicit_signal': local['explicit_signal'],
                        'line_count': len(chunk_lines),
                    }

        if best_candidate is not None:
            candidates.append(best_candidate)

        if set_refs_after:
            in_refs = True

    return candidates, in_refs, analyzable_words


def _select_highlight_plan(orig_doc, target_pct: float):
    """
    يختار فقط الفقرات ذات الدليل الصريح، ويحسب النسبة من جميع كلمات المتن.
    إذا لم توجد أدلة كافية فلن يملأ النسبة قسراً بفقرات عادية.
    """
    pages_candidates = []
    in_refs_global = False
    total_words = 0

    for page_idx in range(len(orig_doc)):
        page = orig_doc[page_idx]
        candidates, in_refs_global, page_words = _collect_page_block_candidates(
            page, in_refs_global
        )
        total_words += page_words
        for item in candidates:
            item['page_idx'] = page_idx
        pages_candidates.extend(candidates)

    total_words = max(total_words, 1)

    # النسبة المخفية *% (0–20%) لا يُسمح معها بأي تظليل مطلقاً.
    if not _highlighting_allowed(target_pct):
        return {
            'plan_by_page': {},
            'total_words': total_words,
            'target_words': 0,
            'covered_words': 0,
            'achieved_pct': 0.0,
            'selected_blocks': 0,
            'available_candidate_words': sum(i['word_count'] for i in pages_candidates),
            'suppressed_low_score': True,
        }

    target_words = int(round(
        total_words * max(0.0, min(100.0, target_pct)) / 100.0
    ))

    ranked = sorted(
        pages_candidates,
        key=lambda x: (x['density'], x['score'], x['t1_count'] + x['t2_total'], -x['word_count']),
        reverse=True
    )

    selected = []
    covered_words = 0
    for item in ranked:
        if target_words > 0 and covered_words >= target_words:
            break
        selected.append(item)
        covered_words += item['word_count']

    plan_by_page = {}
    for item in selected:
        plan_by_page.setdefault(item['page_idx'], []).extend(
            item.get('rects') or [item['rect']]
        )

    achieved_pct = round((covered_words / total_words) * 100.0, 1)

    return {
        'plan_by_page': plan_by_page,
        'total_words': total_words,
        'target_words': target_words,
        'covered_words': covered_words,
        'achieved_pct': achieved_pct,
        'selected_blocks': len(selected),
        'available_candidate_words': sum(i['word_count'] for i in ranked),
        'suppressed_low_score': False,
    }


def generate_report_pdf_from_original(original_pdf_bytes: bytes,
                                       text: str,
                                       doc_name: str = "Document") -> bytes:
    """
    يستخدم ملف PDF الأصلي مباشرة:
    - يحتفظ بكل صفحة كما هي (نص + صور + جداول + تنسيق)
    - يُضيف تظليل أصفر/برتقالي شفاف فوق الفقرات المشبوهة كاملة فقط
    - يجعل مجموع كلمات الفقرات المظللة قريباً من نسبة AI النهائية من كلمات الملف (بدون المراجع)
    - وحدة التظليل هي الفقرة كاملة بكل أسطرها، وليس جزءاً من الجملة أو بعض الأسطر
    - لا يُظلّل قسم المراجع
    - يُضيف صفحة غلاف احترافية في البداية
    - يُضيف صفحة ملخص في النهاية
    - إذا كانت النتيجة *% (0–20%) يُنشئ التقرير بلا أي تظليل نهائياً
    """
    if not FITZ_OK or not RLAB_OK:
        return b""

    result = _ai_fingerprint_score(text)
    pct    = result['percentage']
    today  = datetime.now().strftime("%Y-%m-%d  %H:%M")

    try:
        orig_doc = fitz.open(stream=original_pdf_bytes, filetype="pdf")
    except Exception:
        return b""

    allow_highlighting = _highlighting_allowed(pct)
    highlight_plan = _select_highlight_plan(orig_doc, pct)
    # دفاع إضافي: حتى لو تغيرت خطة الاختيار لاحقاً، النطاق 0–20% يبقى بلا تظليل.
    if not allow_highlighting:
        highlight_plan['plan_by_page'] = {}
        highlight_plan['target_words'] = 0
        highlight_plan['covered_words'] = 0
        highlight_plan['achieved_pct'] = 0.0
        highlight_plan['selected_blocks'] = 0
        highlight_plan['suppressed_low_score'] = True

    result['highlight_target_words'] = highlight_plan['target_words']
    result['highlight_covered_words'] = highlight_plan['covered_words']
    result['highlight_total_words'] = highlight_plan['total_words']
    result['highlight_achieved_pct'] = highlight_plan['achieved_pct']
    result['highlight_selected_blocks'] = highlight_plan['selected_blocks']
    result['highlight_suppressed_low_score'] = highlight_plan.get('suppressed_low_score', False)

    cover_buf = io.BytesIO()
    W_pt, H_pt = A4
    cv = rl_canvas.Canvas(cover_buf, pagesize=A4)
    _draw_cover_rl(cv, W_pt, H_pt, result, doc_name, today)
    cv.save()
    cover_buf.seek(0)

    cover_doc = fitz.open(stream=cover_buf.read(), filetype="pdf")

    out_doc = fitz.open()
    out_doc.insert_pdf(cover_doc)

    for page_idx in range(len(orig_doc)):
        out_doc.insert_pdf(orig_doc, from_page=page_idx, to_page=page_idx)
        new_page = out_doc[-1]

        highlights = (
            highlight_plan['plan_by_page'].get(page_idx, [])
            if allow_highlighting else []
        )
        for line_rect in highlights:
            # كل line_rect هو مستطيل سطر مستقل → annotation منفصل (أسلوب Turnitin)
            annot = new_page.add_highlight_annot(line_rect)
            annot.set_colors(stroke=[0.0, 0.75, 1.0])   # أزرق سماوي
            annot.set_opacity(0.40)
            annot.update()

        _stamp_page_header(new_page, page_idx + 1, len(orig_doc), pct, result['color'])

    summary_buf = io.BytesIO()
    cv2 = rl_canvas.Canvas(summary_buf, pagesize=A4)
    _draw_summary_page(cv2, W_pt, H_pt, result, doc_name, today)
    cv2.save()
    summary_buf.seek(0)
    summary_doc = fitz.open(stream=summary_buf.read(), filetype="pdf")
    out_doc.insert_pdf(summary_doc)

    out_buf = io.BytesIO()
    out_doc.save(out_buf, garbage=4, deflate=True)
    orig_doc.close()
    cover_doc.close()
    summary_doc.close()
    out_doc.close()
    return out_buf.getvalue()


def _stamp_page_header(page, page_num: int, total_pages: int, pct: float, col_hex: str):
    """يُضيف شريط معلومات خفي في أعلى كل صفحة"""
    pw = page.rect.width
    ph = page.rect.height

    # شريط علوي شفاف
    strip_rect = fitz.Rect(0, 0, pw, 18)
    page.draw_rect(strip_rect, color=None, fill=(0.04, 0.08, 0.16), fill_opacity=0.82)

    # نص على اليسار
    page.insert_text(
        (6, 13), f"AI Fingerprint Detector  |  AI Score: {_format_percentage(pct)}",
        fontsize=6.5, color=(0.7, 0.85, 1.0)
    )
    # رقم الصفحة على اليمين
    page.insert_text(
        (pw - 55, 13), f"Page {page_num}/{total_pages}",
        fontsize=6.5, color=(0.7, 0.85, 1.0)
    )


def _draw_cover_rl(c, W, H, result, doc_name, doc_date):
    """صفحة الغلاف الاحترافية"""
    pct = result['percentage']

    score_colors = {
        '#c0392b': (0.75, 0.22, 0.17),
        '#e67e22': (0.90, 0.49, 0.13),
        '#f39c12': (0.95, 0.61, 0.07),
        '#27ae60': (0.15, 0.68, 0.38),
        '#2ecc71': (0.18, 0.80, 0.44),
    }
    score_color_rl = score_colors.get(result['color'], (0.75, 0.22, 0.17))

    # خلفية داكنة
    c.setFillColor(HexColor('#060d1a'))
    c.rect(0, 0, W, H, fill=1, stroke=0)

    # زخارف دائرية
    for (cx, cy, r, cr, ca) in [
        (W*0.1, H*0.88, 200, '#00bcd4', 0.05),
        (W*0.9, H*0.12, 170, '#f4a300', 0.04),
        (W*0.78, H*0.70, 120, '#7c3aed', 0.035),
    ]:
        rgb = tuple(int(cr[i:i+2], 16)/255 for i in (1, 3, 5))
        c.setFillColorRGB(*rgb, alpha=ca)
        c.circle(cx, cy, r, fill=1, stroke=0)

    # شبكة نقاط خفيفة
    c.setFillColorRGB(0.2, 0.5, 0.7, alpha=0.08)
    for gx in range(0, int(W), 30):
        for gy in range(0, int(H), 30):
            c.circle(gx, gy, 1.0, fill=1, stroke=0)

    # شرائط علوية
    for i, (col, th) in enumerate([('#00bcd4', 7), ('#f4a300', 4), ('#7c3aed', 2)]):
        offset = sum(t for _, t in [('#00bcd4', 7), ('#f4a300', 4), ('#7c3aed', 2)][:i])
        c.setFillColor(HexColor(col))
        c.rect(0, H - offset - th, W, th, fill=1, stroke=0)

    # شعار دائري
    lcx, lcy = W/2, H - 105
    c.setFillColorRGB(0.0, 0.74, 0.83, alpha=0.12)
    c.circle(lcx, lcy, 52, fill=1, stroke=0)
    c.setFillColor(HexColor('#00bcd4'))
    c.circle(lcx, lcy, 38, fill=1, stroke=0)
    c.setFillColor(HexColor('#060d1a'))
    c.setFont("Helvetica-Bold", 20)
    c.drawCentredString(lcx, lcy - 7, "AI")

    # عنوان
    c.setFillColor(HexColor('#ffffff'))
    c.setFont("Helvetica-Bold", 26)
    c.drawCentredString(W/2, H - 170, "AI Fingerprint Detector")
    c.setFillColor(HexColor('#7fb3c8'))
    c.setFont("Helvetica", 11)
    c.drawCentredString(W/2, H - 190, "Academic AI-Style Signal Report")

    # فاصل
    c.setStrokeColor(HexColor('#1e3a5f'))
    c.setLineWidth(0.8)
    c.line(60, H - 210, W - 60, H - 210)

    # دائرة النتيجة
    cx, cy = W/2, H - 340
    c.setFillColorRGB(0.06, 0.10, 0.18, alpha=1.0)
    c.circle(cx, cy, 78, fill=1, stroke=0)
    c.setStrokeColorRGB(*score_color_rl)
    c.setLineWidth(4)
    c.circle(cx, cy, 78, fill=0, stroke=1)
    c.setFillColorRGB(*score_color_rl)
    c.setFont("Helvetica-Bold", 46)
    c.drawCentredString(cx, cy - 16, _format_percentage(pct))
    c.setFillColor(HexColor('#aaaaaa'))
    c.setFont("Helvetica", 9)
    c.drawCentredString(cx, cy + 30, "AI-Style Signal Score")

    # الحكم
    c.setFillColorRGB(*score_color_rl)
    c.setFont("Helvetica-Bold", 17)
    c.drawCentredString(W/2, H - 466, result['verdict'])
    c.setFillColor(HexColor('#7fb3c8'))
    c.setFont("Helvetica", 10)
    c.drawCentredString(W/2, H - 484,
        f"Human: {_format_human_percentage(pct, result['human_score'])}    |    AI: {_format_percentage(pct)}")

    # شريط اللون
    bx, by = 55, H - 538
    bw, bh = W - 110, 18
    for sc, s0, s1 in [
        ('#2ecc71', 0.0, 0.22), ('#27ae60', 0.22, 0.42),
        ('#f39c12', 0.42, 0.62), ('#e67e22', 0.62, 0.80),
        ('#c0392b', 0.80, 1.00),
    ]:
        sx = bx + bw * s0
        sw = bw * (s1 - s0) - 1
        c.setFillColor(HexColor(sc))
        c.roundRect(sx, by, sw, bh, 3, fill=1, stroke=0)

    visual_pct = _visual_percentage(pct)
    if visual_pct is not None:
        ind_x = bx + bw * (visual_pct / 100.0)
        c.setFillColor(HexColor('#ffffff'))
        c.setStrokeColor(HexColor('#000000'))
        c.setLineWidth(1.0)
        p = c.beginPath()
        p.moveTo(ind_x, by - 2)
        p.lineTo(ind_x - 6, by - 13)
        p.lineTo(ind_x + 6, by - 13)
        p.close()
        c.drawPath(p, fill=1, stroke=1)
    else:
        c.setFillColor(HexColor('#7fb3c8'))
        c.setFont("Helvetica", 7.5)
        c.drawCentredString(W/2, by - 12, "Exact 0-20% position is masked")

    # صناديق الإحصائيات
    sy = H - 628
    sw_each = (W - 110) / 4
    for i, (lbl, val, vc) in enumerate([
        ("الكلمات",   f"{result['n_words']:,}", '#00bcd4'),
        ("الجمل",     f"{result['n_sents']}",   '#7c3aed'),
        ("T1 إشارات", f"{result['t1_count']}",  '#f4a300'),
        ("T2 أنماط",  f"{result['t2_total']}",  '#e74c3c'),
    ]):
        bx2 = 55 + i * sw_each
        c.setFillColor(HexColor('#0b1829'))
        c.roundRect(bx2 + 3, sy - 32, sw_each - 6, 52, 7, fill=1, stroke=0)
        c.setStrokeColor(HexColor(vc))
        c.setLineWidth(0.8)
        c.roundRect(bx2 + 3, sy - 32, sw_each - 6, 52, 7, fill=0, stroke=1)
        c.setFillColor(HexColor(vc))
        c.setFont("Helvetica-Bold", 17)
        c.drawCentredString(bx2 + sw_each/2, sy + 4, val)
        c.setFillColor(HexColor('#888888'))
        c.setFont("Helvetica", 8)
        c.drawCentredString(bx2 + sw_each/2, sy - 20, lbl)

    # معلومات المستند
    iy = H - 708
    c.setFillColor(HexColor('#0b1829'))
    c.roundRect(55, iy - 18, W - 110, 52, 8, fill=1, stroke=0)
    c.setStrokeColor(HexColor('#1e3a5f'))
    c.setLineWidth(0.5)
    c.roundRect(55, iy - 18, W - 110, 52, 8, fill=0, stroke=1)
    c.setFillColor(HexColor('#aaaaaa'))
    c.setFont("Helvetica", 8.5)
    c.drawString(70, iy + 20, f"Document: {doc_name[:55]}")
    c.drawString(70, iy + 7,  f"Analysis Date: {doc_date}")
    h_pct = result.get('highlight_achieved_pct')
    if result.get('highlight_suppressed_low_score'):
        c.drawString(70, iy - 6, "Engine: AI Fingerprint v3.3  |  No highlighting for masked score *%")
    elif h_pct is not None:
        c.drawString(70, iy - 6,  f"Engine: AI Fingerprint v3.3  |  Highlighted coverage: {_format_percentage(h_pct, mask_low=False)}")
    else:
        c.drawString(70, iy - 6,  "Engine: AI Fingerprint v3.3  |  Original format preserved")

    # شريط أسفل
    c.setFillColor(HexColor('#060d1a'))
    c.rect(0, 0, W, 26, fill=1, stroke=0)
    c.setFillColor(HexColor('#00bcd4'))
    c.rect(0, 25, W, 1.5, fill=1, stroke=0)
    c.setFillColor(HexColor('#555'))
    c.setFont("Helvetica", 7)
    c.drawCentredString(W/2, 8,
        "AI Fingerprint Detector  •  Confidential Analysis Report  •  Academic & Research Use Only")


def _draw_summary_page(c, W, H, result, doc_name, doc_date):
    """صفحة ملخص الإشارات"""
    c.setFillColor(HexColor('#f8fbff'))
    c.rect(0, 0, W, H, fill=1, stroke=0)

    # رأس
    c.setFillColor(HexColor('#060d1a'))
    c.rect(0, H - 40, W, 40, fill=1, stroke=0)
    c.setFillColor(HexColor('#00bcd4'))
    c.rect(0, H - 41, W, 1.5, fill=1, stroke=0)
    c.setFillColor(HexColor('#aaaaaa'))
    c.setFont("Helvetica", 8)
    c.drawString(18, H - 24, "AI Fingerprint Detector")
    c.drawRightString(W - 18, H - 24, f"Summary  |  AI Score: {_format_percentage(result['percentage'])}")

    y = H - 70
    c.setFillColor(HexColor('#00bcd4'))
    c.setFont("Helvetica-Bold", 14)
    c.drawString(52, y, "Detected Style Signals Summary")
    y -= 8
    c.setStrokeColor(HexColor('#1e3a5f'))
    c.setLineWidth(0.8)
    c.line(52, y, W - 52, y)
    y -= 20

    # سياسة التظليل
    low_score_no_highlight = bool(result.get("highlight_suppressed_low_score")) or not _highlighting_allowed(result.get('percentage', 0))
    if low_score_no_highlight:
        c.setFillColor(HexColor('#e8f5e9'))
        c.roundRect(52, y - 5, W - 104, 22, 4, fill=1, stroke=0)
        c.setStrokeColor(HexColor('#27ae60'))
        c.setLineWidth(0.6)
        c.roundRect(52, y - 5, W - 104, 22, 4, fill=0, stroke=1)
        c.setFillColor(HexColor('#1b5e20'))
        c.setFont("Helvetica-Bold", 9)
        c.drawString(58, y + 6, "No highlighting applied: the exact 0-20% score is masked as *%.")
        y -= 30
    else:
        c.setFillColor(HexColor('#fff9c4'))
        c.roundRect(52, y - 4, 14, 12, 2, fill=1, stroke=0)
        c.setStrokeColor(HexColor('#f4a300'))
        c.setLineWidth(0.5)
        c.roundRect(52, y - 4, 14, 12, 2, fill=0, stroke=1)
        c.setFillColor(HexColor('#555'))
        c.setFont("Helvetica", 9)
        c.drawString(70, y + 4, "Highlighted text = selected suspicious lines; references are excluded.")
        y -= 22

        # نسبة التغطية الفعلية للتظليل
        h_pct = result.get("highlight_achieved_pct")
        h_cov = result.get("highlight_covered_words")
        h_tot = result.get("highlight_total_words")
        if h_pct is not None and h_cov is not None and h_tot is not None:
            c.setFillColor(HexColor('#0b1829'))
            c.roundRect(52, y - 5, W - 104, 18, 4, fill=1, stroke=0)
            c.setFillColor(HexColor('#00bcd4'))
            c.setFont("Helvetica-Bold", 9)
            c.drawString(58, y + 5, f"Highlighted coverage (non-references): {_format_percentage(h_pct, mask_low=False)}  |  {h_cov:,} / {h_tot:,} words")
            y -= 26

    # عبارات T1
    if result['t1_hits']:
        c.setFillColor(HexColor('#e67e22'))
        c.setFont("Helvetica-Bold", 11)
        c.drawString(52, y, f"T1 Phrases Detected ({result['t1_count']} hits):")
        y -= 16

        col_idx = 0
        col_w = (W - 110) / 2
        col_x = 52

        for hit in result['t1_hits']:
            if y < 60:
                break
            tag_w = min(len(hit) * 5.5 + 14, col_w - 8)
            c.setFillColor(HexColor('#fff3e0'))
            c.roundRect(col_x, y - 3, tag_w, 14, 4, fill=1, stroke=0)
            c.setStrokeColor(HexColor('#e65100'))
            c.setLineWidth(0.4)
            c.roundRect(col_x, y - 3, tag_w, 14, 4, fill=0, stroke=1)
            c.setFillColor(HexColor('#bf360c'))
            c.setFont("Helvetica", 8)
            c.drawString(col_x + 5, y + 4, hit[:45])
            col_idx += 1
            if col_idx % 2 == 0:
                col_x = 52
                y -= 22
            else:
                col_x = 52 + col_w

        y -= 28

    # أنماط T2
    if result['t2_matched'] and y > 100:
        c.setFillColor(HexColor('#3498db'))
        c.setFont("Helvetica-Bold", 11)
        c.drawString(52, y, f"T2 Patterns Detected ({result['t2_total']} matches):")
        y -= 16

        for cnt, pat in result['t2_matched'][:6]:
            if y < 60:
                break
            c.setFillColor(HexColor('#0b1829'))
            c.roundRect(52, y - 3, W - 110, 14, 3, fill=1, stroke=0)
            c.setFillColor(HexColor('#00bcd4'))
            c.setFont("Helvetica-Bold", 8)
            c.drawString(58, y + 4, f"×{cnt}")
            c.setFillColor(HexColor('#aaaaaa'))
            c.setFont("Helvetica", 7.5)
            c.drawString(78, y + 4, pat[:75])
            y -= 20

    # ذيل
    c.setFillColor(HexColor('#060d1a'))
    c.rect(0, 0, W, 26, fill=1, stroke=0)
    c.setFillColor(HexColor('#f4a300'))
    c.rect(0, 25, W, 1, fill=1, stroke=0)
    c.setFillColor(HexColor('#555'))
    c.setFont("Helvetica", 7)
    c.drawCentredString(W/2, 8, "AI Fingerprint Detector — Confidential Analysis Report")


def _para_is_suspicious(para: str) -> bool:
    """
    يحدد إذا كانت الفقرة مشبوهة (AI) بناءً على نفس منهجية التظليل الأصلية.
    يُستخدم في تقرير DOCX/النص المباشر فقط.
    """
    local = _score_paragraph_for_highlight(para)
    return local['explicit_signal'] and local['evidence_score'] >= 0.20


def generate_report_pdf_text_only(text: str, doc_name: str = "Document") -> bytes:
    """
    نسخة احتياطية: عندما لا يكون هناك ملف PDF أصلي.
    تُنشئ تقريراً نصياً مع تظليل الفقرات المشبوهة.
    المراجع لا تُظلَّل.
    """
    if not RLAB_OK:
        return b""

    buf    = io.BytesIO()
    result = _ai_fingerprint_score(text)
    pct    = result['percentage']
    col_h  = result['color']
    allow_highlighting = _highlighting_allowed(pct)
    result['highlight_suppressed_low_score'] = not allow_highlighting
    result['highlight_target_words'] = 0 if not allow_highlighting else None
    result['highlight_covered_words'] = 0 if not allow_highlighting else None
    result['highlight_total_words'] = result.get('n_words', 0)
    result['highlight_achieved_pct'] = 0.0 if not allow_highlighting else None
    result['highlight_selected_blocks'] = 0 if not allow_highlighting else None
    today  = datetime.now().strftime("%Y-%m-%d  %H:%M")

    W, H   = A4
    ML, MR = 52, 52
    MT, MB = 45, 40
    TW     = W - ML - MR
    LH     = 13.5
    PS     = 9.0

    c = rl_canvas.Canvas(buf, pagesize=A4)

    # غلاف
    _draw_cover_rl(c, W, H, result, doc_name, today)
    c.showPage()

    # صفحات المحتوى
    page_num = 2
    y = H - MT - 36

    # رأس الصفحة
    _rl_page_header(c, W, H, page_num, doc_name, pct)

    # legend
    c.setFillColor(HexColor('#00bcd4'))
    c.setFont("Helvetica-Bold", 12)
    c.drawString(ML, y, "Document Content Analysis")
    y -= 7
    c.setStrokeColor(HexColor('#1e3a5f'))
    c.setLineWidth(0.7)
    c.line(ML, y, W - MR, y)
    y -= 18

    lx = ML
    for lcol, border, ltxt in [
        ('#fff9c4', '#f4a300', '   AI-Suspicious (orange highlight)'),
        ('#f8fafc', '#dce8f5', '   Normal / Human text'),
        ('#f3e5f5', '#c897d8', '   References — not analyzed'),
    ]:
        c.setFillColor(HexColor(lcol))
        c.roundRect(lx, y - 3, 14, 10, 2, fill=1, stroke=0)
        c.setStrokeColor(HexColor(border))
        c.setLineWidth(0.4)
        c.roundRect(lx, y - 3, 14, 10, 2, fill=0, stroke=1)
        c.setFillColor(HexColor('#333333'))
        c.setFont("Helvetica", 8)
        c.drawString(lx + 17, y + 3, ltxt)
        lx += 158
    y -= 22

    # فقرات النص
    paras_raw = [p.strip() for p in text.split('\n') if p.strip()]
    in_refs = False
    para_groups = []
    for para in paras_raw:
        if _is_reference_header(para):
            in_refs = True
        is_ref  = in_refs or _is_reference_line(para)
        is_susp = allow_highlighting and (not is_ref) and _para_is_suspicious(para)
        para_groups.append((para, is_susp, is_ref))

    def new_page_rl():
        nonlocal page_num, y
        _rl_page_footer(c, W, page_num)
        c.showPage()
        page_num += 1
        y = H - MT - 40
        _rl_page_header(c, W, H, page_num, doc_name, pct)

    def draw_para_rl(txt, is_susp, is_ref):
        nonlocal y

        if is_ref:
            bg, border, tc, fn, fs = '#f3e5f5', '#c897d8', '#6a1b9a', "Helvetica-Oblique", 9.0
        elif is_susp:
            bg, border, tc, fn, fs = '#fff9c4', '#f4a300', '#5d3a00', "Helvetica", 10.5
        else:
            bg, border, tc, fn, fs = '#f8fafc', '#dce8f5', '#1a1a2a', "Helvetica", 10.5

        avg_char_w = fs * 0.52
        max_c = max(10, int(TW / avg_char_w))
        words_ = txt.split()
        lines  = []
        cur    = ""
        for w in words_:
            test = (cur + " " + w).strip()
            if len(test) <= max_c:
                cur = test
            else:
                if cur: lines.append(cur)
                cur = w
        if cur: lines.append(cur)

        bh = len(lines) * LH + 8

        if y - bh < MB + 26:
            new_page_rl()

        # خلفية
        c.setFillColor(HexColor(bg))
        c.roundRect(ML - 5, y - bh + 4, TW + 10, bh + 2, 4, fill=1, stroke=0)
        if not is_ref:
            c.setStrokeColor(HexColor(border))
            c.setLineWidth(0.3)
            c.roundRect(ML - 5, y - bh + 4, TW + 10, bh + 2, 4, fill=0, stroke=1)

        # شريط جانبي للمشبوه — مثل تظليل Turnitin
        if is_susp:
            c.setFillColor(HexColor('#f4a300'))
            c.rect(ML - 5, y - bh + 4, 4, bh + 2, fill=1, stroke=0)

        # النص
        c.setFillColor(HexColor(tc))
        c.setFont(fn, fs)
        ty = y
        for line in lines:
            if ty < MB + 26:
                new_page_rl()
                ty = y
            c.drawString(ML + 6, ty, line)
            ty -= LH

        # علامة AI
        if is_susp:
            tag_y = y + 1
            c.setFillColor(HexColor('#f4a300'))
            c.roundRect(W - MR - 28, tag_y - 1, 28, 12, 3, fill=1, stroke=0)
            c.setFillColor(HexColor('#ffffff'))
            c.setFont("Helvetica-Bold", 6.5)
            c.drawCentredString(W - MR - 14, tag_y + 4, "AI ✓")

        y = ty - PS

    for pt, ps_, pr in para_groups:
        draw_para_rl(pt, ps_, pr)

    # صفحة ملخص
    _rl_page_footer(c, W, page_num)
    c.showPage()
    page_num += 1
    _draw_summary_page(c, W, H, result, doc_name, today)
    c.save()
    return buf.getvalue()


def _rl_page_header(c, W, H, page_num, doc_name, pct):
    c.setFillColor(HexColor('#060d1a'))
    c.rect(0, H - 30, W, 30, fill=1, stroke=0)
    c.setFillColor(HexColor('#00bcd4'))
    c.rect(0, H - 31, W, 1.5, fill=1, stroke=0)
    c.setFillColor(HexColor('#aaaaaa'))
    c.setFont("Helvetica", 7.5)
    c.drawString(18, H - 19, "AI Fingerprint Detector")
    c.drawCentredString(W/2, H - 19, doc_name[:42])
    c.drawRightString(W - 18, H - 19, f"AI: {_format_percentage(pct)}  |  p.{page_num}")


def _rl_page_footer(c, W, page_num):
    c.setFillColor(HexColor('#060d1a'))
    c.rect(0, 0, W, 24, fill=1, stroke=0)
    c.setFillColor(HexColor('#f4a300'))
    c.rect(0, 23, W, 1, fill=1, stroke=0)
    c.setFillColor(HexColor('#555'))
    c.setFont("Helvetica", 7)
    c.drawCentredString(W/2, 7, f"Page {page_num}  —  AI Fingerprint Detector  —  Confidential")


# ══════════════════════════════════════════════════════════════════════════════
# HELPERS: قراءة الملفات
# ══════════════════════════════════════════════════════════════════════════════

def extract_text_from_pdf(file_bytes: bytes) -> str:
    if not FITZ_OK: return ""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        return "\n".join(page.get_text() for page in doc)
    except Exception:
        return ""


def extract_text_from_docx(file_bytes: bytes) -> str:
    if not DOCX_OK:
        return ""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        return "\n".join(
            p.text for p in _iter_docx_paragraphs(doc) if (p.text or "").strip()
        )
    except Exception:
        return ""



def _iter_docx_paragraphs(container):
    """
    يمر على فقرات Word والجداول بترتيب ظهورها الحقيقي في المستند.
    هذا يمنع انتقال حالة "داخل المراجع" إلى أجزاء سابقة بسبب ترتيب غير صحيح.
    """
    try:
        if isinstance(container, _Document):
            parent_element = container.element.body
            parent = container
        elif isinstance(container, _Cell):
            parent_element = container._tc
            parent = container
        else:
            for p in getattr(container, "paragraphs", []):
                yield p
            return

        for child in parent_element.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                table = Table(child, parent)
                seen_cells = set()
                for row in table.rows:
                    for cell in row.cells:
                        # الخلايا المدمجة قد تظهر أكثر من مرة؛ لا نكرر نصها.
                        # احتفظ بعنصر XML نفسه، لا بـ id رقمي قد تعيد بايثون استخدامه
                        # أثناء إنشاء wrappers للخلايا؛ إعادة استخدام id كانت تُسقط خلايا صحيحة.
                        cell_key = cell._tc
                        if cell_key in seen_cells:
                            continue
                        seen_cells.add(cell_key)
                        yield from _iter_docx_paragraphs(cell)
    except Exception:
        return


def _set_run_docx_shading(run, fill: str = "9FE2FF"):
    """
    يضيف تظليل خلفية حقيقي على run داخل DOCX
    للحفاظ على النص الأصلي وتنسيقه بالكامل.
    """
    try:
        r_pr = run._r.get_or_add_rPr()
        shd = r_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            r_pr.append(shd)
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
    except Exception:
        pass


def generate_highlighted_docx_bytes(docx_bytes: bytes,
                                    text: str = "",
                                    doc_name: str = "Document") -> bytes:
    """
    يطبّق نفس مبدأ PDF لكن مباشرة على ملف Word الأصلي:
    - لا يعيد كتابة التقرير كنص جديد
    - يحتفظ بالتنسيق والجداول والأشكال كما هي
    - يظلل الفقرات المشبوهة داخل DOCX نفسه
    - يستثني قسم المراجع من التظليل
    """
    if not DOCX_OK:
        return b""

    # احتياطياً: إذا كانت النتيجة *% نعيد ملف Word الأصلي بلا أي تعديل.
    if text and not _highlighting_allowed(_ai_fingerprint_score(text).get('percentage', 0)):
        return docx_bytes

    try:
        doc = docx.Document(io.BytesIO(docx_bytes))
    except Exception:
        return b""

    in_refs = False
    highlighted = 0

    try:
        for para in _iter_docx_paragraphs(doc):
            para_text = (para.text or "").strip()
            if not para_text:
                continue

            if _is_reference_header(para_text):
                in_refs = True

            is_ref = in_refs or _is_reference_line(para_text)
            is_susp = (not is_ref) and _para_is_suspicious(para_text)

            if not is_susp:
                continue

            # نطبّق التظليل على كل run موجود للحفاظ على تنسيق كل جزء كما هو
            non_empty_run_found = False
            for run in para.runs:
                if (run.text or "").strip():
                    _set_run_docx_shading(run, fill="9FE2FF")
                    non_empty_run_found = True

            # إذا كانت الفقرة بلا runs واضحة، ننشئ run أخيراً بدون المساس بالمحتوى
            if not non_empty_run_found and para_text:
                run = para.add_run("")
                _set_run_docx_shading(run, fill="9FE2FF")

            highlighted += 1
    except Exception:
        return b""

    out = io.BytesIO()
    try:
        doc.save(out)
        return out.getvalue()
    except Exception:
        return b""



_DOCX_CONVERSION_DIAGNOSTIC = {
    "ok": False,
    "engine": "",
    "message": "لم تبدأ محاولة التحويل بعد.",
    "details": "",
}


def _set_docx_conversion_diagnostic(ok: bool, engine: str, message: str, details: str = "") -> None:
    """يحفظ آخر حالة تحويل Word حتى تعرض الواجهة سبب النجاح أو الفشل الحقيقي."""
    global _DOCX_CONVERSION_DIAGNOSTIC
    _DOCX_CONVERSION_DIAGNOSTIC = {
        "ok": bool(ok),
        "engine": str(engine or ""),
        "message": str(message or ""),
        "details": str(details or "")[:5000],
    }


def get_last_docx_conversion_diagnostic() -> dict:
    """يعيد نسخة آمنة من آخر تشخيص لتحويل Word إلى PDF."""
    return dict(_DOCX_CONVERSION_DIAGNOSTIC)


def convert_docx_to_pdf_bytes(docx_bytes: bytes) -> bytes:
    """
    يحوّل DOCX إلى PDF مع الحفاظ على تخطيط المستند قدر الإمكان.

    ترتيب المحركات:
    1) Microsoft Word عبر docx2pdf على Windows/macOS.
    2) Microsoft Word COM على Windows.
    3) LibreOffice Writer على Linux وStreamlit Community Cloud.

    ملاحظة نشر مهمة:
    - requirements.txt يثبت مكتبات Python فقط.
    - على Streamlit Community Cloud يجب تثبيت LibreOffice كاعتماد نظام
      من خلال packages.txt في جذر المستودع.
    """
    import gc
    import os
    import platform
    import shutil
    import socket
    import stat
    import subprocess
    import tempfile
    import time
    from contextlib import contextmanager
    from pathlib import Path

    _set_docx_conversion_diagnostic(False, "", "بدأت محاولة تحويل Word إلى PDF.")

    if not docx_bytes or len(docx_bytes) < 100:
        _set_docx_conversion_diagnostic(
            False,
            "input",
            "ملف Word فارغ أو غير صالح.",
        )
        return b""

    def _read_valid_pdf(pdf_path: str) -> bytes:
        try:
            if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 100:
                with open(pdf_path, "rb") as f:
                    data = f.read()
                if data.startswith(b"%PDF"):
                    return data
        except Exception:
            pass
        return b""

    def _force_remove(path: str) -> None:
        """يزيل ملفًا للقراءة فقط عند استدعائه من shutil.rmtree."""
        try:
            os.chmod(path, stat.S_IWRITE)
            if os.path.isdir(path):
                os.rmdir(path)
            else:
                os.remove(path)
        except Exception:
            pass

    def _cleanup_temp_tree(path: str, retries: int = 8) -> None:
        """
        تنظيف متسامح مع Windows وWord/LibreOffice.

        بعض محركات التحويل تُبقي ملف lock أو ملفًا مخفيًا لجزء من الثانية بعد
        انتهاء التصدير. إعادة المحاولة تمنع WinError 145 من إسقاط التطبيق،
        بينما يبقى فشل التنظيف الأخير غير قاتل لأن ملف PDF قُرئ بالفعل للذاكرة.
        """
        if not path:
            return

        for attempt in range(retries):
            if not os.path.exists(path):
                return
            try:
                shutil.rmtree(path, onerror=lambda func, item, exc: _force_remove(item))
                if not os.path.exists(path):
                    return
            except FileNotFoundError:
                return
            except OSError:
                pass

            gc.collect()
            time.sleep(0.15 * (attempt + 1))

        # محاولة أخيرة غير قاتلة. قد يبقى مجلد صغير إلى أن يحرره Word/مضاد الفيروسات.
        try:
            shutil.rmtree(path, ignore_errors=True)
        except Exception:
            pass

    @contextmanager
    def _safe_temporary_directory(prefix: str):
        tmp_path = tempfile.mkdtemp(prefix=prefix)
        try:
            yield tmp_path
        finally:
            _cleanup_temp_tree(tmp_path)

    with _safe_temporary_directory(prefix="ai_detector_docx_") as tmpdir:
        docx_path = os.path.join(tmpdir, "input.docx")
        expected_pdf_path = os.path.join(tmpdir, "input.pdf")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)

        system_name = platform.system()
        errors = []

        # المسار الأدق على الأجهزة التي تحتوي Microsoft Word.
        if system_name in {"Windows", "Darwin"}:
            try:
                from docx2pdf import convert as docx2pdf_convert

                docx2pdf_convert(docx_path, expected_pdf_path)
                data = _read_valid_pdf(expected_pdf_path)
                if data:
                    _set_docx_conversion_diagnostic(
                        True,
                        "Microsoft Word / docx2pdf",
                        "تم تصدير Word إلى PDF بواسطة Microsoft Word.",
                    )
                    return data
                errors.append("docx2pdf انتهى دون إنشاء PDF صالح.")
            except Exception as exc:
                errors.append(f"docx2pdf: {type(exc).__name__}: {exc}")

        # Word COM على Windows.
        if os.name == "nt":
            word_app = None
            word_doc = None
            pythoncom_initialized = False
            try:
                import pythoncom
                import win32com.client

                pythoncom.CoInitialize()
                pythoncom_initialized = True
                word_app = win32com.client.DispatchEx("Word.Application")
                word_app.Visible = False
                word_app.DisplayAlerts = 0
                word_doc = word_app.Documents.Open(
                    os.path.abspath(docx_path),
                    ConfirmConversions=False,
                    ReadOnly=True,
                    AddToRecentFiles=False,
                )
                word_doc.ExportAsFixedFormat(
                    OutputFileName=os.path.abspath(expected_pdf_path),
                    ExportFormat=17,
                    OpenAfterExport=False,
                    OptimizeFor=0,
                    Range=0,
                    Item=0,
                    IncludeDocProps=True,
                    KeepIRM=True,
                    CreateBookmarks=1,
                    DocStructureTags=True,
                    BitmapMissingFonts=True,
                    UseISO19005_1=False,
                )
                data = _read_valid_pdf(expected_pdf_path)
                if data:
                    _set_docx_conversion_diagnostic(
                        True,
                        "Microsoft Word COM",
                        "تم تصدير Word إلى PDF بواسطة Microsoft Word.",
                    )
                    return data
                errors.append("Word COM انتهى دون إنشاء PDF صالح.")
            except Exception as exc:
                errors.append(f"Word COM: {type(exc).__name__}: {exc}")
            finally:
                try:
                    if word_doc is not None:
                        word_doc.Close(False)
                except Exception:
                    pass
                try:
                    if word_app is not None:
                        word_app.Quit()
                except Exception:
                    pass
                if pythoncom_initialized:
                    try:
                        pythoncom.CoUninitialize()
                    except Exception:
                        pass

        # LibreOffice: المحرك المطلوب على Streamlit Community Cloud/Linux.
        env = os.environ.copy()
        env["TMPDIR"] = tmpdir
        env["HOME"] = tmpdir
        env["XDG_RUNTIME_DIR"] = tmpdir
        env["SAL_USE_VCLPLUGIN"] = "svp"
        env.setdefault("LANG", "C.UTF-8")
        env.setdefault("LC_ALL", "C.UTF-8")

        # حل احتياطي لبيئات Linux المقيدة التي تمنع AF_UNIX.
        af_unix = getattr(socket, "AF_UNIX", None)
        if af_unix is not None and os.name != "nt":
            try:
                sock = socket.socket(af_unix, socket.SOCK_STREAM)
                sock.close()
            except OSError:
                shim_path = "/tmp/lo_socket_shim.so"
                if os.path.exists(shim_path):
                    env["LD_PRELOAD"] = shim_path

        candidates = [shutil.which("soffice"), shutil.which("libreoffice")]
        if os.name == "nt":
            candidates.extend([
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            ])
        elif system_name == "Darwin":
            candidates.append("/Applications/LibreOffice.app/Contents/MacOS/soffice")
        else:
            candidates.extend([
                "/usr/bin/soffice",
                "/usr/bin/libreoffice",
                "/usr/local/bin/soffice",
                "/usr/local/bin/libreoffice",
            ])

        soffice_bin = next((c for c in candidates if c and os.path.exists(c)), None)
        if not soffice_bin:
            details = "\n".join(errors)
            _set_docx_conversion_diagnostic(
                False,
                "LibreOffice",
                "LibreOffice غير مثبت في بيئة الاستضافة. أضف ملف packages.txt إلى جذر مستودع GitHub ثم أعد تشغيل التطبيق.",
                details,
            )
            return b""

        # تأكد أن الملف التنفيذي يعمل، وهذا يعطي رسالة أوضح من فشل صامت.
        try:
            version_proc = subprocess.run(
                [soffice_bin, "--version"],
                env=env,
                capture_output=True,
                timeout=30,
                text=True,
            )
            version_text = (version_proc.stdout or version_proc.stderr or "").strip()
            if version_proc.returncode != 0:
                errors.append(
                    f"LibreOffice --version returncode={version_proc.returncode}: {version_text}"
                )
        except Exception as exc:
            version_text = ""
            errors.append(f"LibreOffice version check: {type(exc).__name__}: {exc}")

        profile_dir = Path(tmpdir, "lo_profile")
        profile_dir.mkdir(parents=True, exist_ok=True)
        cmd = [
            soffice_bin,
            f"-env:UserInstallation={profile_dir.as_uri()}",
            "--headless",
            "--invisible",
            "--nologo",
            "--nodefault",
            "--nolockcheck",
            "--norestore",
            "--nofirststartwizard",
            "--convert-to", "pdf:writer_pdf_Export",
            "--outdir", tmpdir,
            docx_path,
        ]

        try:
            proc = subprocess.run(
                cmd,
                env=env,
                capture_output=True,
                timeout=240,
                text=True,
            )
        except subprocess.TimeoutExpired as exc:
            _set_docx_conversion_diagnostic(
                False,
                "LibreOffice",
                "انتهت مهلة تحويل Word إلى PDF داخل بيئة الاستضافة.",
                f"timeout={exc.timeout}; errors={' | '.join(errors)}",
            )
            return b""
        except Exception as exc:
            _set_docx_conversion_diagnostic(
                False,
                "LibreOffice",
                "تعذر تشغيل LibreOffice داخل بيئة الاستضافة.",
                f"{type(exc).__name__}: {exc}\n" + "\n".join(errors),
            )
            return b""

        # LibreOffice قد يعيد تسمية الملف؛ لذلك نبحث عن أي PDF ناتج.
        pdf_candidates = [Path(expected_pdf_path)] + sorted(Path(tmpdir).glob("*.pdf"))
        for candidate in pdf_candidates:
            data = _read_valid_pdf(str(candidate))
            if data:
                details = "\n".join(
                    part for part in [
                        version_text,
                        (proc.stdout or "").strip(),
                        (proc.stderr or "").strip(),
                    ] if part
                )
                _set_docx_conversion_diagnostic(
                    True,
                    "LibreOffice Writer",
                    "تم تصدير Word إلى PDF بواسطة LibreOffice Writer.",
                    details,
                )
                return data

        details = "\n".join(
            part for part in [
                f"returncode={proc.returncode}",
                (proc.stdout or "").strip(),
                (proc.stderr or "").strip(),
                "\n".join(errors),
            ] if part
        )
        _set_docx_conversion_diagnostic(
            False,
            "LibreOffice",
            "عمل LibreOffice لكنه لم يُنشئ ملف PDF صالحًا.",
            details,
        )
        return b""


def _build_logo_data_uri() -> str:
    svg = """<svg xmlns='http://www.w3.org/2000/svg' width='220' height='220' viewBox='0 0 220 220'>
      <defs>
        <linearGradient id='lg1' x1='0%' y1='0%' x2='100%' y2='100%'>
          <stop offset='0%' stop-color='#0a0e27'/>
          <stop offset='100%' stop-color='#0d1f4a'/>
        </linearGradient>
        <linearGradient id='lg2' x1='0%' y1='0%' x2='100%' y2='100%'>
          <stop offset='0%' stop-color='#00d4ff'/>
          <stop offset='50%' stop-color='#7b61ff'/>
          <stop offset='100%' stop-color='#ff6b9d'/>
        </linearGradient>
        <linearGradient id='lg3' x1='0%' y1='0%' x2='100%' y2='100%'>
          <stop offset='0%' stop-color='#00d4ff' stop-opacity='0.3'/>
          <stop offset='100%' stop-color='#7b61ff' stop-opacity='0.1'/>
        </linearGradient>
        <filter id='glow'>
          <feGaussianBlur stdDeviation='3' result='coloredBlur'/>
          <feMerge><feMergeNode in='coloredBlur'/><feMergeNode in='SourceGraphic'/></feMerge>
        </filter>
      </defs>
      <!-- خلفية -->
      <rect x='0' y='0' width='220' height='220' rx='44' fill='url(#lg1)'/>
      <!-- حلقات توهج خلفية -->
      <circle cx='110' cy='110' r='90' fill='none' stroke='url(#lg2)' stroke-width='0.5' opacity='0.2'/>
      <circle cx='110' cy='110' r='72' fill='none' stroke='url(#lg2)' stroke-width='0.5' opacity='0.3'/>
      <!-- شبكة نقاط خفية -->
      <g opacity='0.12' fill='#00d4ff'>
        <circle cx='40' cy='40' r='1.5'/><circle cx='70' cy='40' r='1.5'/><circle cx='100' cy='40' r='1.5'/>
        <circle cx='130' cy='40' r='1.5'/><circle cx='160' cy='40' r='1.5'/><circle cx='180' cy='40' r='1.5'/>
        <circle cx='40' cy='70' r='1.5'/><circle cx='180' cy='70' r='1.5'/>
        <circle cx='40' cy='180' r='1.5'/><circle cx='70' cy='180' r='1.5'/><circle cx='130' cy='180' r='1.5'/>
        <circle cx='160' cy='180' r='1.5'/><circle cx='180' cy='180' r='1.5'/>
      </g>
      <!-- دائرة رئيسية متوهجة -->
      <circle cx='110' cy='105' r='52' fill='url(#lg3)' filter='url(#glow)'/>
      <circle cx='110' cy='105' r='44' fill='none' stroke='url(#lg2)' stroke-width='2.5' opacity='0.9'/>
      <!-- خطوط بصمة إصبع -->
      <g fill='none' stroke='url(#lg2)' stroke-linecap='round' opacity='0.7'>
        <path d='M90 105 Q110 88 130 105 Q110 122 90 105' stroke-width='2'/>
        <path d='M83 105 Q110 80 137 105 Q110 130 83 105' stroke-width='1.5'/>
        <path d='M76 105 Q110 72 144 105 Q110 138 76 105' stroke-width='1.2'/>
        <path d='M70 105 Q110 65 150 105 Q110 145 70 105' stroke-width='1'/>
      </g>
      <!-- نقطة مركزية -->
      <circle cx='110' cy='105' r='6' fill='#00d4ff' filter='url(#glow)'/>
      <circle cx='110' cy='105' r='3' fill='#ffffff'/>
      <!-- خطوط ماسحة رادار -->
      <line x1='110' y1='105' x2='110' y2='65' stroke='#00d4ff' stroke-width='1.5' opacity='0.6'/>
      <line x1='110' y1='105' x2='140' y2='80' stroke='#7b61ff' stroke-width='1' opacity='0.5'/>
      <!-- نص AI -->
      <text x='110' y='168' text-anchor='middle' font-size='13' font-family='Arial Black,Arial' font-weight='900'
            letter-spacing='4' fill='url(#lg2)' opacity='0.9'>AI·FP</text>
      <!-- إطار خارجي فاخر -->
      <rect x='3' y='3' width='214' height='214' rx='42' fill='none' stroke='url(#lg2)' stroke-width='1.5' opacity='0.4'/>
    </svg>"""
    return "data:image/svg+xml;base64," + base64.b64encode(svg.encode("utf-8")).decode("ascii")


logo_data_uri = _build_logo_data_uri()

st.markdown(f"""
<style>
@import url('https://fonts.googleapis.com/css2?family=Cairo:wght@400;600;700;900&family=Space+Grotesk:wght@400;600;700&display=swap');

* {{ box-sizing: border-box; }}

body,
[data-testid="stAppViewContainer"],
[data-testid="stVerticalBlock"],
.main {{
    direction: rtl;
    font-family: 'Cairo', 'Segoe UI', Tahoma, sans-serif;
}}

pre, code, .stat-num, .layer-val, .ltr {{
    direction: ltr !important;
    text-align: left !important;
    font-family: 'Space Grotesk', 'Courier New', monospace;
}}

[data-testid="stAppViewContainer"] {{
    background:
        radial-gradient(ellipse 80% 60% at 10% 0%, rgba(0,212,255,0.07) 0%, transparent 60%),
        radial-gradient(ellipse 60% 70% at 90% 100%, rgba(123,97,255,0.09) 0%, transparent 60%),
        radial-gradient(ellipse 50% 50% at 50% 50%, rgba(255,107,157,0.04) 0%, transparent 70%),
        linear-gradient(160deg, #07091a 0%, #0b0f2a 40%, #0e0a22 100%);
    min-height: 100vh;
}}

[data-testid="stHeader"] {{
    background: rgba(7,9,26,0.85) !important;
    backdrop-filter: blur(20px);
    border-bottom: 1px solid rgba(0,212,255,0.12);
}}

.main .block-container {{
    max-width: 980px;
    padding-top: 2rem;
    padding-bottom: 4rem;
    padding-right: 1.4rem;
    padding-left: 1.4rem;
}}

.hero-card {{
    background: linear-gradient(145deg, rgba(13,16,42,0.95), rgba(10,14,35,0.98));
    border: 1px solid rgba(0,212,255,0.18);
    border-radius: 28px;
    padding: 1.8rem 2rem 1.5rem;
    box-shadow: 0 0 0 1px rgba(123,97,255,0.08) inset, 0 30px 80px rgba(0,0,0,0.5);
    margin-bottom: 1.4rem;
    position: relative;
    overflow: hidden;
}}

.hero-card::before {{
    content: '';
    position: absolute;
    top: -80px; right: -80px;
    width: 260px; height: 260px;
    background: radial-gradient(circle, rgba(0,212,255,0.07), transparent 70%);
    pointer-events: none;
}}

.hero-card::after {{
    content: '';
    position: absolute;
    bottom: -60px; left: -60px;
    width: 200px; height: 200px;
    background: radial-gradient(circle, rgba(123,97,255,0.06), transparent 70%);
    pointer-events: none;
}}

.hero-card-topbar {{
    position: absolute;
    top: 0; left: 0; right: 0; height: 3px;
    background: linear-gradient(90deg, #00d4ff, #7b61ff, #ff6b9d, #00d4ff);
    background-size: 200% 100%;
    animation: shimmer 4s linear infinite;
    border-radius: 28px 28px 0 0;
}}

@keyframes shimmer {{
    0%   {{ background-position: 0% 0%; }}
    100% {{ background-position: 200% 0%; }}
}}

.hero-grid {{
    display: grid;
    grid-template-columns: 110px 1fr;
    gap: 1.4rem;
    align-items: center;
    direction: rtl;
}}

.hero-logo {{
    width: 100px; height: 100px;
    border-radius: 24px;
    display: flex;
    align-items: center;
    justify-content: center;
    background: linear-gradient(145deg, #0d1030, #0a0e27);
    border: 1px solid rgba(0,212,255,0.25);
    box-shadow: 0 0 30px rgba(0,212,255,0.15), 0 10px 30px rgba(0,0,0,0.4);
    overflow: hidden;
    flex-shrink: 0;
}}

.hero-logo img {{
    width: 92px; height: 92px;
    object-fit: contain;
}}

.hero-text {{ text-align: right; }}

.hero-title {{
    font-size: 2.1rem;
    font-weight: 900;
    background: linear-gradient(135deg, #00d4ff 0%, #7b61ff 50%, #ff6b9d 100%);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
    line-height: 1.1;
    margin-bottom: 0.45rem;
    letter-spacing: -0.5px;
    direction: ltr;
    text-align: left;
}}

.hero-subtitle {{
    color: rgba(180,200,240,0.75);
    font-size: 0.93rem;
    line-height: 1.9;
    direction: rtl;
    text-align: right;
}}

.hero-badges {{
    display: flex;
    flex-wrap: wrap;
    gap: 0.5rem;
    margin-top: 1rem;
    justify-content: flex-end;
    direction: rtl;
}}

.hero-badge {{
    background: rgba(0,212,255,0.07);
    border: 1px solid rgba(0,212,255,0.18);
    color: #7ecfea;
    border-radius: 999px;
    padding: 0.3rem 0.85rem;
    font-size: 0.78rem;
    font-weight: 700;
    white-space: nowrap;
    transition: all 0.2s;
}}

.features-wrap {{
    background: rgba(10,14,35,0.85);
    border: 1px solid rgba(123,97,255,0.15);
    border-radius: 24px;
    padding: 1.4rem;
    box-shadow: 0 20px 50px rgba(0,0,0,0.3);
    margin-bottom: 0.6rem;
    position: relative;
    overflow: hidden;
}}

.features-wrap::before {{
    content: '';
    position: absolute;
    top: 0; left: 0; right: 0; height: 2px;
    background: linear-gradient(90deg, transparent, rgba(123,97,255,0.6), transparent);
}}

.features-title {{
    font-weight: 900;
    color: #c5d5f0;
    margin-bottom: 1rem;
    font-size: 1rem;
    text-align: right;
}}

.features-grid {{
    display: grid;
    grid-template-columns: repeat(2, 1fr);
    gap: 0.85rem;
}}

.feature-card {{
    background: rgba(15,20,50,0.7);
    border: 1px solid rgba(0,212,255,0.10);
    border-radius: 18px;
    padding: 1.1rem;
    transition: transform 0.25s, box-shadow 0.25s, border-color 0.25s;
    text-align: right;
    position: relative;
    overflow: hidden;
}}

.feature-card:hover {{
    transform: translateY(-3px);
    box-shadow: 0 16px 40px rgba(0,0,0,0.4), 0 0 20px rgba(0,212,255,0.07);
    border-color: rgba(0,212,255,0.22);
}}

.feature-icon {{
    width: 40px; height: 40px;
    border-radius: 12px;
    display: inline-flex;
    align-items: center;
    justify-content: center;
    background: linear-gradient(135deg, rgba(0,212,255,0.15), rgba(123,97,255,0.15));
    border: 1px solid rgba(0,212,255,0.15);
    font-size: 1.15rem;
    margin-bottom: 0.55rem;
}}

.feature-card strong {{
    display: block;
    color: #c8deff;
    margin-bottom: 0.3rem;
    font-size: 0.92rem;
    font-weight: 800;
}}

.feature-card span {{
    color: rgba(140,165,210,0.7);
    font-size: 0.82rem;
    line-height: 1.75;
}}

.section-card {{
    background: rgba(10,14,35,0.88);
    border: 1px solid rgba(0,212,255,0.14);
    border-radius: 24px;
    padding: 1.4rem;
    box-shadow: 0 20px 50px rgba(0,0,0,0.3);
    position: relative;
}}

.section-card::before {{
    content: '';
    position: absolute;
    top: 0; left: 0; right: 0; height: 2px;
    background: linear-gradient(90deg, transparent, rgba(0,212,255,0.5), transparent);
    border-radius: 24px 24px 0 0;
}}

.meter-wrap {{
    background: rgba(10,14,35,0.9);
    border: 1px solid rgba(0,212,255,0.15);
    border-radius: 24px;
    padding: 2.4rem 2rem 2rem;
    margin: 1.5rem 0;
    text-align: center;
    box-shadow: 0 30px 70px rgba(0,0,0,0.4);
    position: relative;
    overflow: hidden;
}}

.meter-wrap::before {{
    content: '';
    position: absolute;
    top: 0; left: 0; right: 0; height: 2px;
    background: linear-gradient(90deg, transparent, rgba(0,212,255,0.6), rgba(123,97,255,0.6), transparent);
}}

.meter-pct {{
    font-size: 5.5rem;
    font-weight: 900;
    line-height: 1;
    letter-spacing: -4px;
    font-family: 'Space Grotesk', sans-serif;
    direction: ltr;
    display: block;
}}

.meter-label {{
    font-size: 1.2rem;
    margin-top: 0.6rem;
    font-weight: 800;
    direction: rtl;
}}

.meter-human {{
    font-size: 0.9rem;
    color: rgba(140,165,210,0.65);
    margin-top: 0.5rem;
    direction: rtl;
    font-family: 'Space Grotesk', sans-serif;
}}

.bar-track {{
    background: rgba(255,255,255,0.06);
    border-radius: 999px;
    height: 14px;
    margin: 1.5rem 0;
    overflow: hidden;
    border: 1px solid rgba(255,255,255,0.04);
}}

.bar-fill {{
    height: 100%;
    border-radius: 999px;
    transition: width 0.8s cubic-bezier(.4, 0, .2, 1);
}}

.meter-mask-note {{
    margin-top: -0.65rem;
    font-size: 0.82rem;
    color: rgba(140,165,210,0.78);
    direction: rtl;
}}

.layer-card {{
    background: rgba(12,17,45,0.85);
    border: 1px solid rgba(0,212,255,0.10);
    border-radius: 16px;
    padding: 1rem 1.2rem;
    margin-bottom: 0.75rem;
    direction: rtl;
    text-align: right;
    transition: border-color 0.2s;
    overflow: hidden;
}}

.layer-card:hover {{
    border-color: rgba(0,212,255,0.22);
}}

.layer-title {{
    font-weight: 800;
    font-size: 0.9rem;
    color: #b0c8e8;
    direction: rtl;
    display: inline-block;
}}

.layer-val {{
    float: left;
    font-weight: 900;
    font-size: 0.92rem;
    direction: ltr !important;
    text-align: left !important;
    font-family: 'Space Grotesk', monospace;
}}

.layer-bar {{
    clear: both;
    height: 8px;
    background: rgba(255,255,255,0.06);
    border-radius: 999px;
    margin-top: 0.75rem;
    overflow: hidden;
}}

.layer-bar-fill {{
    height: 100%;
    border-radius: 999px;
}}

.hit-tag {{
    display: inline-block;
    background: rgba(255,107,107,0.08);
    border: 1px solid rgba(255,107,107,0.22);
    border-radius: 8px;
    padding: 4px 12px;
    margin: 3px;
    font-size: 0.77rem;
    font-family: 'Space Grotesk', 'Courier New', monospace;
    color: #ff9999;
    direction: ltr !important;
    text-align: left !important;
    transition: background 0.2s;
}}

.stat-row {{
    display: flex;
    gap: 0.75rem;
    flex-wrap: wrap;
    margin: 1rem 0;
    direction: rtl;
}}

.stat-box {{
    flex: 1;
    min-width: 95px;
    background: rgba(12,17,45,0.85);
    border: 1px solid rgba(0,212,255,0.12);
    border-radius: 16px;
    padding: 0.9rem 0.8rem;
    text-align: center;
    transition: transform 0.2s, border-color 0.2s;
}}

.stat-box:hover {{
    transform: translateY(-2px);
    border-color: rgba(0,212,255,0.28);
}}

.stat-num {{
    font-size: 1.5rem;
    font-weight: 900;
    color: #00d4ff;
    direction: ltr !important;
    text-align: center !important;
    font-family: 'Space Grotesk', sans-serif;
    display: block;
}}

.stat-lbl {{
    font-size: 0.7rem;
    color: rgba(140,165,210,0.6);
    margin-top: 0.25rem;
    direction: rtl;
}}

div[data-testid="stTabs"] {{
    direction: rtl;
}}

div[data-testid="stTabs"] button[role="tab"] {{
    border-radius: 12px !important;
    background: rgba(12,17,45,0.8) !important;
    border: 1px solid rgba(0,212,255,0.12) !important;
    color: rgba(140,180,220,0.8) !important;
    font-weight: 700 !important;
    font-size: 0.9rem !important;
    font-family: 'Cairo', sans-serif !important;
}}

div[data-testid="stTabs"] button[aria-selected="true"] {{
    background: linear-gradient(135deg, rgba(0,212,255,0.15), rgba(123,97,255,0.15)) !important;
    color: #00d4ff !important;
    border-color: rgba(0,212,255,0.3) !important;
}}

div[data-testid="stTextArea"] textarea {{
    background: rgba(7,9,26,0.8) !important;
    border-radius: 16px !important;
    border: 1.5px solid rgba(0,212,255,0.15) !important;
    color: #c8deff !important;
    direction: rtl;
    text-align: right;
    font-size: 0.95rem !important;
    font-family: 'Cairo', sans-serif !important;
}}

div[data-testid="stTextArea"] textarea::placeholder {{
    text-align: right;
    color: rgba(100,130,180,0.45) !important;
}}

div[data-testid="stFileUploader"] section {{
    background: rgba(7,9,26,0.6) !important;
    border-radius: 16px !important;
    border: 1.5px dashed rgba(0,212,255,0.2) !important;
    color: rgba(140,180,220,0.7) !important;
}}

.stButton > button,
.stDownloadButton > button {{
    border: none !important;
    border-radius: 14px !important;
    font-size: 0.97rem !important;
    font-weight: 800 !important;
    padding: 0.78rem 1.5rem !important;
    width: 100% !important;
    transition: all 0.25s cubic-bezier(.4, 0, .2, 1);
    font-family: 'Cairo', sans-serif !important;
}}

.stButton > button[kind="primary"],
.stDownloadButton > button {{
    background: linear-gradient(135deg, #00d4ff, #7b61ff) !important;
    color: white !important;
    box-shadow: 0 8px 25px rgba(0,212,255,0.25), 0 4px 10px rgba(0,0,0,0.3) !important;
}}

.stButton > button[kind="primary"]:hover,
.stDownloadButton > button:hover {{
    transform: translateY(-2px);
    box-shadow: 0 14px 35px rgba(0,212,255,0.35), 0 6px 15px rgba(0,0,0,0.3) !important;
}}

.stButton > button:not([kind="primary"]) {{
    background: rgba(12,17,45,0.85) !important;
    color: rgba(140,180,220,0.85) !important;
    border: 1.5px solid rgba(0,212,255,0.18) !important;
}}

.stButton > button:not([kind="primary"]):hover {{
    background: rgba(0,212,255,0.08) !important;
    border-color: rgba(0,212,255,0.35) !important;
    color: #00d4ff !important;
    transform: translateY(-1px);
}}

.stAlert,
div[data-testid="stAlert"] {{
    border-radius: 14px !important;
    direction: rtl !important;
    text-align: right !important;
}}

div[data-testid="stAlert"] {{
    background: rgba(0,212,255,0.06) !important;
    border: 1px solid rgba(0,212,255,0.18) !important;
}}

div[data-testid="stAlert"] p {{
    direction: rtl !important;
    text-align: right !important;
    color: rgba(140,200,240,0.85) !important;
}}

h3 {{
    text-align: right !important;
    direction: rtl !important;
    color: #7db8e0 !important;
    font-weight: 800 !important;
    margin: 1.2rem 0 0.6rem !important;
    font-family: 'Cairo', sans-serif !important;
}}

hr {{
    border: none !important;
    border-top: 1px solid rgba(0,212,255,0.12) !important;
    margin: 1.5rem 0 !important;
}}

.info-note {{
    text-align: center;
    font-size: 0.8rem;
    color: rgba(140,165,210,0.5);
    margin-top: 0.5rem;
    direction: rtl;
}}

code {{
    background: rgba(0,212,255,0.08) !important;
    border: 1px solid rgba(0,212,255,0.15) !important;
    color: #00d4ff !important;
    border-radius: 6px !important;
    padding: 2px 6px !important;
    direction: ltr !important;
    text-align: left !important;
    font-family: 'Space Grotesk', 'Courier New', monospace !important;
}}

[data-testid="stMarkdownContainer"] p,
[data-testid="stMarkdownContainer"] li {{
    color: rgba(160,185,225,0.85);
    direction: rtl;
    text-align: right;
    line-height: 1.85;
}}

[data-testid="stMarkdownContainer"] strong {{
    color: #c8deff;
    font-weight: 800;
}}

[data-testid="stSuccess"] {{
    background: rgba(0,200,100,0.08) !important;
    border-color: rgba(0,200,100,0.2) !important;
    border-radius: 14px !important;
    direction: rtl !important;
    text-align: right !important;
}}

[data-testid="stError"] {{
    background: rgba(255,80,80,0.08) !important;
    border-color: rgba(255,80,80,0.2) !important;
    border-radius: 14px !important;
    direction: rtl !important;
    text-align: right !important;
}}

[data-testid="stWarning"] {{
    background: rgba(255,170,0,0.08) !important;
    border-color: rgba(255,170,0,0.2) !important;
    border-radius: 14px !important;
    direction: rtl !important;
    text-align: right !important;
}}

@media (max-width: 720px) {{
    .hero-grid, .features-grid {{
        grid-template-columns: 1fr;
    }}
    .hero-logo {{
        margin: 0 auto;
    }}
    .hero-title {{
        font-size: 1.6rem;
        text-align: right;
    }}
}}
</style>
""", unsafe_allow_html=True)


# ── Hero Section ──────────────────────────────────────────────────────────────
st.markdown(f"""
<div class="hero-card">
    <div class="hero-card-topbar"></div>
    <div class="hero-grid">
        <div class="hero-logo">
            <img src="{logo_data_uri}" alt="AI Fingerprint Detector">
        </div>
        <div class="hero-text">
            <div class="hero-title">AI Fingerprint Detector</div>
            <div class="hero-subtitle">
                أداة متقدمة لكشف بصمات الذكاء الاصطناعي في النصوص الأكاديمية،
                مع تقرير مُظلَّل يعمل على ملفات PDF و Word مع الحفاظ التام على التنسيق والجداول والأشكال.
            </div>
            <div class="hero-badges">
                <span class="hero-badge">🎯 تظليل دقيق على الملف الأصلي</span>
                <span class="hero-badge">📄 PDF و Word (تنسيق كامل)</span>
                <span class="hero-badge">🚫 المراجع خارج التظليل</span>
                <span class="hero-badge">⚖️ مراجعة محافظة تقلل الإيجابيات الكاذبة</span>
            </div>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)


# ── Features ──────────────────────────────────────────────────────────────────
st.markdown("""
<div class="features-wrap">
    <div class="features-title">ماذا يُقدّم البرنامج؟</div>
    <div class="features-grid">
        <div class="feature-card">
            <div class="feature-icon">🔍</div>
            <strong>فحص دقيق بطبقتين</strong>
            <span>تحليل إشارات T1 وأنماط T2 مع منع احتساب الإشارة نفسها مرتين وتقييم الأسلوب والبنية.</span>
        </div>
        <div class="feature-card">
            <div class="feature-icon">📑</div>
            <strong>التقرير المظلل: PDF و Word (تنسيق كامل)</strong>
            <span>عند رفع PDF يتم التظليل فوق الصفحات الأصلية مباشرة. وعند رفع DOCX يتم تصدير الملف الأصلي أولاً إلى PDF بواسطة Microsoft Word عند توفره، ثم يُنشأ عليه التقرير نفسه تماماً مع الحفاظ على الجداول والصور والأشكال والتنسيق.</span>
        </div>
        <div class="feature-card">
            <div class="feature-icon">🚫</div>
            <strong>المراجع محمية من التظليل</strong>
            <span>يستبعد قسم المراجع والاستشهادات الرقمية والمؤلف–السنة من الحساب، ولا يطبّق على قسم المراجع أي تظليل.</span>
        </div>
        <div class="feature-card">
            <div class="feature-icon">📊</div>
            <strong>إحصاءات تفصيلية وغلاف احترافي</strong>
            <span>صفحة غلاف فنية + إحصاءات + صفحة ملخص الإشارات المكتشفة في نهاية التقرير.</span>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)


# ── Input Section ─────────────────────────────────────────────────────────────
st.markdown('<div class="section-card">', unsafe_allow_html=True)

st.markdown("### رفع الملف للفحص")
st.info("📌 ارفع ملف PDF أو Word (DOCX). ملف Word يُصدَّر أولاً إلى PDF من الأصل دون تعديل، ثم يُطبَّق عليه تقرير PDF نفسه. كل نتيجة من 0% إلى 20% شاملة تظهر *%، ولا يتم معها تظليل أي نص نهائياً.")

input_text         = ""
uploaded_name      = "Document"
original_pdf_bytes = None
original_docx_bytes = None

uploaded = st.file_uploader(
    "ارفع ملف PDF أو Word:",
    type=["pdf", "docx"],
    key="file_upload",
)

if uploaded:
    uploaded_name = uploaded.name
    raw = uploaded.read()
    ext = uploaded.name.lower().rsplit(".", 1)[-1]

    if ext == "pdf":
        original_pdf_bytes = raw
        input_text = extract_text_from_pdf(raw)
        if input_text:
            st.success(f"✅ تم استخراج النص من ملف PDF ({len(input_text.split()):,} كلمة)")
        else:
            st.error("تعذّر استخراج النص من ملف PDF. تأكد من أن الملف قابل للقراءة وأن PyMuPDF مثبتة.")
    elif ext == "docx":
        original_docx_bytes = raw
        input_text = extract_text_from_docx(raw)
        if input_text:
            st.success(f"✅ تم استخراج النص من ملف Word ({len(input_text.split()):,} كلمة)")
            with st.spinner("جارٍ تصدير ملف Word الأصلي إلى PDF مع الحفاظ على التنسيق والجداول والصور والأشكال..."):
                converted_pdf = convert_docx_to_pdf_bytes(raw)
                conversion_diag = get_last_docx_conversion_diagnostic()
            if converted_pdf:
                original_pdf_bytes = converted_pdf
                engine_name = conversion_diag.get("engine") or "محرك التحويل"
                st.info(
                    f"✅ تم تصدير Word الأصلي إلى PDF بواسطة {engine_name} دون أي تعديل مسبق. "
                    "سيُنشأ التقرير عليه بنفس غلاف وتظليل وملخص تقرير PDF تماماً."
                )
            else:
                st.error(
                    "تعذّر تحويل Word إلى PDF داخل بيئة الاستضافة، ولذلك لا يمكن إنشاء تقرير "
                    "مطابق لتقرير PDF مع الحفاظ على الجداول والصور والتنسيق."
                )
                st.warning(
                    conversion_diag.get("message")
                    or "على Streamlit Community Cloud أضف LibreOffice إلى packages.txt، وليس requirements.txt فقط."
                )
                details = (conversion_diag.get("details") or "").strip()
                if details:
                    with st.expander("التفاصيل الفنية للتحويل"):
                        st.code(details)
        else:
            st.error("تعذّر استخراج النص من ملف Word. تأكد من تثبيت python-docx.")


# ── Buttons ───────────────────────────────────────────────────────────────────
col_btn, col_clear = st.columns([3, 1])
with col_btn:
    analyze_btn = st.button("🔍 فحص الآن", use_container_width=True, type="primary")
with col_clear:
    if st.button("🗑️ مسح", use_container_width=True):
        st.session_state.pop("last_result", None)
        st.session_state.pop("last_text", None)
        st.session_state.pop("last_name", None)
        st.session_state.pop("last_pdf_bytes", None)
        st.session_state.pop("last_docx_bytes", None)
        st.rerun()

if analyze_btn and input_text and input_text.strip():
    with st.spinner("جارٍ الفحص..."):
        result = _ai_fingerprint_score(input_text.strip())
    st.session_state["last_result"]     = result
    st.session_state["last_text"]       = input_text.strip()
    st.session_state["last_name"]       = uploaded_name
    st.session_state["last_pdf_bytes"]  = original_pdf_bytes   # PDF أصلي أو PDF مُصدَّر من Word الأصلي
    st.session_state["last_docx_bytes"] = original_docx_bytes

st.markdown('</div>', unsafe_allow_html=True)


# ── Results ───────────────────────────────────────────────────────────────────
if "last_result" in st.session_state:
    st.markdown("### نتيجة التحليل")
    r   = st.session_state["last_result"]
    pct = r["percentage"]
    clr = r["color"]
    # لا نعتمد على قيمة مخزنة قديمة؛ نطبق سياسة الإخفاء في كل إعادة رسم للواجهة.
    pct_label = _format_percentage(pct)
    human_label = _format_human_percentage(pct, r["human_score"])
    visual_pct = _visual_percentage(pct)
    meter_width = 0.0 if visual_pct is None else visual_pct
    masked_note = (
        '<div class="meter-mask-note">0–20% مخفية بالكامل ولا ينتج عنها أي تظليل.</div>'
        if visual_pct is None else ''
    )

    # Meter
    st.markdown(f"""
    <div class="meter-wrap">
        <div class="meter-pct" style="color:{clr}">{pct_label}</div>
        <div class="meter-label" style="color:{clr}">{r['verdict']}</div>
        <div class="meter-human">بشري: {human_label}</div>
        <div class="bar-track">
            <div class="bar-fill" style="width:{meter_width}%;background:{clr}"></div>
        </div>
        {masked_note}
    </div>
    """, unsafe_allow_html=True)

    st.caption(
        "هذه النتيجة مؤشر أسلوبي للمراجعة وليست إثباتاً قاطعاً على استخدام الذكاء الاصطناعي. "
        "كفاية الأدلة تعكس حجم النص وتنوع الإشارات، وليست نسبة دقة مؤكدة."
    )
    if r.get("reference_words_excluded", 0) > 0:
        st.success(
            f"✅ تم استبعاد {r['reference_words_excluded']:,} كلمة من قسم المراجع "
            f"و{r.get('inline_citation_words_removed', 0):,} كلمة من علامات الاستشهاد داخل المتن قبل الحساب."
        )
    elif not r.get("reference_section_found", False):
        st.info("ℹ️ لم يُكتشف قسم مراجع مستقل. تم مع ذلك حذف الاستشهادات المضمنة التي أمكن التعرف عليها.")

    # Stats
    st.markdown(f"""
    <div class="stat-row">
        <div class="stat-box"><div class="stat-num">{r['n_words']:,}</div><div class="stat-lbl">كلمة محللة</div></div>
        <div class="stat-box"><div class="stat-num">{r['n_sents']}</div><div class="stat-lbl">جملة</div></div>
        <div class="stat-box"><div class="stat-num">{r['avg_sent_len']}</div><div class="stat-lbl">كلمة/جملة</div></div>
        <div class="stat-box"><div class="stat-num">{r['t1_count']}</div><div class="stat-lbl">عبارة T1</div></div>
        <div class="stat-box"><div class="stat-num">{r['t2_total']}</div><div class="stat-lbl">نمط T2</div></div>
        <div class="stat-box"><div class="stat-num">{r['reference_words_excluded']:,}</div><div class="stat-lbl">كلمة مراجع مستبعدة</div></div>
        <div class="stat-box"><div class="stat-num">{r['confidence_percentage']:.0f}%</div><div class="stat-lbl">كفاية الأدلة</div></div>
    </div>
    """, unsafe_allow_html=True)

    # Layers
    st.markdown("### تفصيل الطبقات")
    for title, raw_val, lcolor in [
        ("T1 — إشارات لغوية (×0.45)", r["t1_score"],    "#e74c3c"),
        ("T2 — أنماط جملة (×0.40)",   r["t2_score"],    "#e67e22"),
        ("Style — أسلوب (×0.15)",      r["style_score"], "#3498db"),
        ("Struct — بنية (boost)",       min(r["struct_boost"] * 10, 1.0), "#9b59b6"),
    ]:
        bar_pct = int(raw_val * 100)
        st.markdown(f"""
        <div class="layer-card">
            <span class="layer-title">{title}</span>
            <span class="layer-val" style="color:{lcolor}">{raw_val:.3f}</span>
            <div class="layer-bar">
                <div class="layer-bar-fill" style="width:{bar_pct}%;background:{lcolor}"></div>
            </div>
        </div>
        """, unsafe_allow_html=True)

    # T1 Hits
    if r["t1_hits"]:
        st.markdown("### 🎯 عبارات T1 المكتشفة")
        tags = "".join(f'<span class="hit-tag">{h}</span>' for h in r["t1_hits"])
        st.markdown(f'<div style="line-height:2.4;direction:rtl">{tags}</div>', unsafe_allow_html=True)

    # T2
    if r["t2_matched"]:
        st.markdown("### 🔎 أنماط T2 المكتشفة")
        for cnt, pat in r["t2_matched"]:
            st.markdown(
                f'<span class="hit-tag">×{cnt}</span> <code style="font-size:.8rem;direction:ltr">{pat}</code>',
                unsafe_allow_html=True
            )

    # Top keyword
    if r["top_kw"]:
        st.markdown(
            f"**الكلمة الأكثر تكراراً:** `{r['top_kw']}` "
            f"× {r['top_kw_count']} من {r['n_words']:,} كلمة "
            f"({r['top_kw_count']/r['n_words']*100:.2f}%)"
        )

    st.divider()

    # ── PDF / DOCX Report Download ─────────────────────────────────────────────
    if RLAB_OK:
        st.markdown("### 📥 تحميل التقرير الكامل")

        saved_text                 = st.session_state.get("last_text", "")
        saved_name                 = st.session_state.get("last_name", "Document")
        saved_pdf_bytes            = st.session_state.get("last_pdf_bytes", None)
        saved_docx_bytes           = st.session_state.get("last_docx_bytes", None)

        if saved_text and saved_pdf_bytes:
            # ── تقرير واحد موحد: PDF أصلي أو PDF مُصدَّر من Word الأصلي ───
            is_converted_docx = saved_docx_bytes is not None
            label_spinner = (
                "جارٍ توليد تقرير PDF الموحد من ملف Word المُصدَّر مع الحفاظ على تخطيطه..."
                if is_converted_docx
                else "جارٍ توليد التقرير من ملف PDF الأصلي..."
            )
            with st.spinner(label_spinner):
                pdf_bytes = generate_report_pdf_from_original(
                    saved_pdf_bytes, saved_text, saved_name
                )
            if pdf_bytes:
                report_filename = f"AI_Report_{saved_name.replace('.', '_')}.pdf"
                btn_label = (
                    "⬇️  تحميل تقرير Word بصيغة PDF (مطابق لمسار تقرير PDF)"
                    if is_converted_docx
                    else "⬇️  تحميل التقرير المظلل من PDF الأصلي"
                )
                st.download_button(
                    label=btn_label,
                    data=pdf_bytes,
                    file_name=report_filename,
                    mime="application/pdf",
                    use_container_width=True,
                )

                no_highlight_note = (
                    " · بدون أي تظليل لأن النتيجة *%"
                    if not _highlighting_allowed(r.get('percentage', 0))
                    else " · تظليل أزرق سماوي فوق المقاطع المختارة"
                )
                info_detail = (
                    "✅ تقرير Word خرج عبر نفس مسار تقرير PDF: غلاف احترافي · صفحات Word بعد تصديرها كما هي "
                    "بجداولها وصورها وأشكالها ورؤوسها وتذييلاتها · المراجع مستبعدة · صفحة ملخص"
                    + no_highlight_note
                    if is_converted_docx
                    else
                    "✅ التقرير يحتوي على: صفحة غلاف · الصفحات الأصلية كما هي · المراجع بدون تظليل · ملخص الإشارات"
                    + no_highlight_note
                )
                st.markdown(
                    f"<div class='info-note'>{info_detail}</div>",
                    unsafe_allow_html=True,
                )
            else:
                st.warning("تعذّر توليد التقرير. تأكد من تثبيت PyMuPDF و reportlab.")

        elif saved_text and saved_docx_bytes:
            # ── حل أخير فقط إذا تعذرت كل مسارات الحفاظ على الأصل ────────────
            with st.spinner("جارٍ توليد التقرير النصي الاحتياطي..."):
                pdf_bytes = generate_report_pdf_text_only(saved_text, saved_name)
            if pdf_bytes:
                report_filename = f"AI_Report_{saved_name.replace('.', '_')}.pdf"
                st.download_button(
                    label="⬇️  تحميل التقرير النصي الاحتياطي",
                    data=pdf_bytes,
                    file_name=report_filename,
                    mime="application/pdf",
                    use_container_width=True,
                )
                st.markdown(
                    "<div class='info-note'>"
                    "⚠️ هذا المسار احتياطي فقط عند تعذّر الحفاظ على الأصل."
                    "</div>",
                    unsafe_allow_html=True,
                )
            else:
                st.warning("تعذّر توليد التقرير. تأكد من تثبيت reportlab.")
    else:
        st.info("💡 لتفعيل تصدير PDF: `pip install reportlab`")

    st.markdown(
        f"<div style='text-align:center;font-size:.82rem;opacity:.5;direction:ltr'>"
        f"Displayed score: {pct_label} | Decision: {r['decision']} | Evidence sufficiency: {r['confidence_percentage']:.0f}%"
        f"</div>",
        unsafe_allow_html=True,
    )

elif analyze_btn and not (input_text and input_text.strip()):
    st.warning("الرجاء رفع ملف PDF أو Word أولاً.")

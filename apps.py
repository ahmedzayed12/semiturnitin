"""
AI Fingerprint Detector v5.5.1 — Strict Academic Fingerprint Analysis
تقرير PDF يعمل على ملفات PDF الأصلية فقط مع تظليل النقاط المشكوك بها فوق الأصل نفسه
تقرير فوق الصفحات الأصلية مع تظليل شفاف، واستبعاد محافظ للمراجع والاستشهادات من الحساب
"""
import re
import math
import collections
import io
import base64
import json
import hashlib
import random
import zlib
import unicodedata
import os
import time
import textwrap
import urllib.request
import urllib.error
import urllib.parse
import ssl
from pathlib import Path
from datetime import datetime

import streamlit as st

# ── مكتبات اختيارية ──────────────────────────────────────────────────────────
try:
    import fitz  # PyMuPDF
    FITZ_OK = True
except Exception:
    FITZ_OK = False

try:
    import docx
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.text.paragraph import Paragraph
    from docx.table import Table, _Cell
    from docx.document import Document as _Document
    DOCX_OK = True
except Exception:
    DOCX_OK = False

try:
    from reportlab.pdfgen import canvas as rl_canvas
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.colors import HexColor
    RLAB_OK = True
except Exception:
    RLAB_OK = False

try:
    from PIL import Image
    PIL_OK = True
except Exception:
    PIL_OK = False


# ══════════════════════════════════════════════════════════════════════════════
# GEMINI-ASSISTED ACADEMIC REVIEW
# ══════════════════════════════════════════════════════════════════════════════
# The API key is never stored in this file. It is read only from Streamlit
# Secrets or a server environment variable. End users never enter or see it.
# يُحسم نموذج واحد متاح للمشروع مرة واحدة عند تشغيل التطبيق، ثم يظل ثابتًا طوال الجلسة.
# لا ينتقل البرنامج بين النماذج أثناء الفحص، ولا يصدر التقرير إلا بعد نجاح المراجعة الذكية.
GEMINI_PREFERRED_MODELS = (
    "gemini-2.5-flash",
    "gemini-2.5-flash-lite",
    "gemini-2.0-flash",
)
GEMINI_DEFAULT_MODEL = "gemini-2.5-flash"
GEMINI_MAX_PASSAGES = 8
GEMINI_MAX_WORDS_PER_PASSAGE = 180
GEMINI_API_TIMEOUT_SECONDS = 120
GEMINI_MAX_ATTEMPTS = 4
GEMINI_RETRY_BASE_SECONDS = 2.5
GEMINI_RETRY_CAP_SECONDS = 18.0


def _gemini_ssl_context():
    try:
        import certifi
        return ssl.create_default_context(cafile=certifi.where())
    except Exception:
        return ssl.create_default_context()


class GeminiAPIError(RuntimeError):
    """A user-safe Gemini API failure."""


def _secret_or_env(name: str, default=""):
    """Read a server-side setting without exposing it in the user interface."""
    try:
        value = st.secrets.get(name, default)
    except Exception:
        value = os.environ.get(name, default)
    return value


def _get_gemini_api_key() -> str:
    """Read the owner-managed API key; users never enter it."""
    return str(_secret_or_env("GEMINI_API_KEY", "") or "").strip()


def _gemini_service_enabled() -> bool:
    raw = str(_secret_or_env("GEMINI_ENABLED", "true") or "true").strip().lower()
    return raw not in {"0", "false", "no", "off", "disabled"}


_RESOLVED_GEMINI_MODEL = None


def _list_generate_content_models(api_key: str) -> list[str]:
    """استعلام إداري خفيف يُنفذ مرة واحدة لمعرفة نماذج المشروع المتاحة."""
    endpoint = "https://generativelanguage.googleapis.com/v1beta/models"
    request = urllib.request.Request(
        endpoint,
        headers={
            "x-goog-api-key": api_key,
            "User-Agent": "Academic-AI-Reviewer/1.0",
        },
        method="GET",
    )
    try:
        with urllib.request.urlopen(
            request, timeout=30, context=_gemini_ssl_context()
        ) as response:
            payload = json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as exc:
        if exc.code == 403:
            raise GeminiAPIError("بيانات اتصال خدمة المراجعة الذكية غير صالحة أو غير مخولة.")
        raise GeminiAPIError(f"تعذر التحقق من نماذج المراجعة المتاحة (HTTP {exc.code}).")
    except Exception as exc:
        raise GeminiAPIError(f"تعذر التحقق من نموذج المراجعة المتاح: {exc}")

    available = []
    for item in payload.get("models") or []:
        if not isinstance(item, dict):
            continue
        methods = item.get("supportedGenerationMethods") or []
        if "generateContent" not in methods:
            continue
        name = str(item.get("name") or "").strip()
        if name.startswith("models/"):
            name = name.split("/", 1)[1]
        if name:
            available.append(name)
    return available


def _resolve_fixed_gemini_model(api_key: str) -> str:
    """اختيار نموذج واحد مرة واحدة ثم تثبيته طوال تشغيل التطبيق."""
    global _RESOLVED_GEMINI_MODEL
    if _RESOLVED_GEMINI_MODEL:
        return _RESOLVED_GEMINI_MODEL

    configured = str(_secret_or_env("GEMINI_MODEL", "") or "").strip()
    available = _list_generate_content_models(api_key)
    if configured and configured in available:
        _RESOLVED_GEMINI_MODEL = configured
        return configured

    for candidate in GEMINI_PREFERRED_MODELS:
        if candidate in available:
            _RESOLVED_GEMINI_MODEL = candidate
            return candidate

    flash_models = [m for m in available if "flash" in m.lower()]
    if flash_models:
        _RESOLVED_GEMINI_MODEL = sorted(flash_models)[0]
        return _RESOLVED_GEMINI_MODEL

    if available:
        _RESOLVED_GEMINI_MODEL = sorted(available)[0]
        return _RESOLVED_GEMINI_MODEL

    raise GeminiAPIError(
        "لا يوجد في هذا المشروع نموذج يدعم المراجعة النصية. راجع المفتاح والمشروع والفوترة وصلاحية Gemini API."
    )


def _configured_gemini_model(api_key: str = "") -> str:
    """إرجاع النموذج المثبت للمشروع، مع حسمه مرة واحدة عند الحاجة."""
    if not api_key:
        api_key = _get_gemini_api_key()
    if not api_key:
        return GEMINI_DEFAULT_MODEL
    return _resolve_fixed_gemini_model(api_key)


def _gemini_review_color(score: float) -> str:
    score = max(0.0, min(100.0, float(score or 0.0)))
    if score >= 80:
        return "#c0392b"
    if score >= 60:
        return "#e67e22"
    if score >= 40:
        return "#f39c12"
    if score >= 20:
        return "#27ae60"
    return "#2ecc71"


def _gemini_review_label(score: float) -> str:
    score = max(0.0, min(100.0, float(score or 0.0)))
    if score >= 80:
        return "أولوية مراجعة مرتفعة جدًا"
    if score >= 60:
        return "أولوية مراجعة مرتفعة"
    if score >= 40:
        return "أولوية مراجعة متوسطة"
    if score >= 20:
        return "أولوية مراجعة محدودة"
    return "أولوية مراجعة منخفضة"


def _truncate_words(text: str, limit: int) -> str:
    words = re.findall(r"\S+", text or "")
    if len(words) <= limit:
        return " ".join(words)
    return " ".join(words[:limit]) + " …"


def _select_gemini_passages(analysis_text: str, local_result: dict) -> list[dict]:
    """Select representative and locally elevated passages without sending the full paper."""
    sentences = _split_analysis_sentences(analysis_text)
    chunks = _build_analysis_chunks(
        sentences,
        target_words=180,
        min_words=75,
        max_words=GEMINI_MAX_WORDS_PER_PASSAGE,
        overlap_sentences=1,
    )
    if not chunks:
        chunks = _split_analysis_paragraphs(analysis_text)
    chunks = [c.strip() for c in chunks if len(_style_tokens(c)) >= 45]
    if not chunks:
        chunks = [_truncate_words(analysis_text, GEMINI_MAX_WORDS_PER_PASSAGE)]

    raw_scores = list(local_result.get("chunk_scores") or [])
    scores = []
    for idx in range(len(chunks)):
        try:
            scores.append(float(raw_scores[idx]))
        except Exception:
            scores.append(0.0)

    max_items = min(GEMINI_MAX_PASSAGES, len(chunks))
    top_n = min(max_items // 2 + 1, len(chunks))
    top_indices = sorted(range(len(chunks)), key=lambda i: scores[i], reverse=True)[:top_n]

    remaining = max_items - len(top_indices)
    evenly_spaced = []
    if remaining > 0:
        if remaining == 1:
            candidates = [len(chunks) // 2]
        else:
            candidates = [round(i * (len(chunks) - 1) / (remaining - 1)) for i in range(remaining)]
        evenly_spaced.extend(candidates)

    selected_indices = []
    for idx in top_indices + evenly_spaced:
        if idx not in selected_indices:
            selected_indices.append(idx)
    for idx in range(len(chunks)):
        if len(selected_indices) >= max_items:
            break
        if idx not in selected_indices:
            selected_indices.append(idx)
    selected_indices = sorted(selected_indices[:max_items])

    passages = []
    for number, idx in enumerate(selected_indices, start=1):
        passage_id = f"P{number:02d}"
        passages.append({
            "passage_id": passage_id,
            "source_index": idx,
            "local_signal": round(scores[idx], 1),
            "text": _truncate_words(chunks[idx], GEMINI_MAX_WORDS_PER_PASSAGE),
        })
    return passages


_GEMINI_RESPONSE_SCHEMA = {
    "type": "object",
    "properties": {
        "ai_assistance_likelihood_score": {"type": "integer"},
        "review_priority_score": {"type": "integer"},
        "citation_integrity_risk_score": {"type": "integer"},
        "confidence": {"type": "string", "enum": ["low", "medium", "high"]},
        "summary_ar": {"type": "string"},
        "summary_en": {"type": "string"},
        "key_reasons_ar": {"type": "array", "items": {"type": "string"}},
        "key_reasons_en": {"type": "array", "items": {"type": "string"}},
        "counter_evidence_ar": {"type": "array", "items": {"type": "string"}},
        "recommended_actions_ar": {"type": "array", "items": {"type": "string"}},
        "recommended_actions_en": {"type": "array", "items": {"type": "string"}},
        "flagged_passages": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "passage_id": {"type": "string"},
                    "category": {
                        "type": "string",
                        "enum": [
                            "formulaic_language", "unsupported_claim", "citation_problem",
                            "style_discontinuity", "overgeneralization", "possible_ai_assistance",
                            "other"
                        ],
                    },
                    "severity": {"type": "integer"},
                    "reason_ar": {"type": "string"},
                    "reason_en": {"type": "string"},
                    "quote": {"type": "string"},
                },
                "required": ["passage_id", "category", "severity", "reason_ar", "reason_en", "quote"],
            },
        },
    },
    "required": [
        "ai_assistance_likelihood_score", "review_priority_score", "citation_integrity_risk_score", "confidence",
        "summary_ar", "summary_en", "key_reasons_ar", "key_reasons_en",
        "counter_evidence_ar", "recommended_actions_ar", "recommended_actions_en",
        "flagged_passages"
    ],
}


def _gemini_prompt(passages: list[dict], local_result: dict) -> str:
    local_summary = {
        "internal_screening_score": local_result.get("percentage"),
        "evidence_sufficiency": local_result.get("confidence_percentage"),
        "evidence_families": local_result.get("evidence_families"),
        "references_excluded_words": local_result.get("reference_words_excluded"),
        "document_words": local_result.get("n_words"),
    }
    passage_text = "\n\n".join(
        f"[{p['passage_id']}] local_signal={p['local_signal']}\n{p['text']}" for p in passages
    )
    return f"""
You are an academic integrity reviewer. Review the supplied English academic passages and the weak internal screening summary.

Important constraints:
1. Do not claim that AI authorship or plagiarism can be proven from style alone.
2. The internal score is a weak screening signal, not ground truth. Correct it when the prose provides contrary evidence.
3. Distinguish polished academic writing from genuinely repetitive, generic, unsupported, or discontinuous writing.
4. The citation-integrity score concerns missing/weak/implausible citation support only. It is NOT a plagiarism or similarity percentage.
5. Estimate ai_assistance_likelihood_score from 0-100 as a conservative screening likelihood that substantial AI assistance shaped the supplied prose. It is not proof, must not be based on polish alone, and high values require repeated evidence across passages.
6. The review-priority score is a separate triage score (0-100), not a plagiarism percentage.
7. Flag only supplied passage IDs. Quotes must be copied exactly from the supplied passage and kept under 180 characters.
8. Use conservative judgments. A high score requires repeated, concrete evidence across several passages.
9. Return Arabic explanations for the interface and English explanations for the PDF report.

Internal screening summary:
{json.dumps(local_summary, ensure_ascii=False)}

Passages:
{passage_text}
""".strip()


def _extract_json_object(raw_text: str) -> dict:
    text = (raw_text or "").strip()
    text = re.sub(r"^```(?:json)?\s*", "", text, flags=re.I)
    text = re.sub(r"\s*```$", "", text)
    try:
        value = json.loads(text)
        if isinstance(value, dict):
            return value
    except Exception:
        pass
    start, end = text.find("{"), text.rfind("}")
    if start >= 0 and end > start:
        try:
            value = json.loads(text[start:end + 1])
            if isinstance(value, dict):
                return value
        except Exception:
            pass
    raise GeminiAPIError("أعادت خدمة المراجعة الذكية استجابة غير قابلة للقراءة.")


def _retry_after_seconds(exc: urllib.error.HTTPError) -> float | None:
    """قراءة Retry-After عند إرساله من الخدمة، دون تجاوز حد انتظار معقول."""
    try:
        raw = (exc.headers or {}).get("Retry-After")
        if raw is None:
            return None
        value = float(str(raw).strip())
        return max(0.0, min(GEMINI_RETRY_CAP_SECONDS, value))
    except Exception:
        return None


def _smart_review_retry_delay(attempt_index: int, retry_after: float | None = None) -> float:
    """تأخير تصاعدي مع jitter لتجنب إعادة الضغط الفوري على النموذج نفسه."""
    exponential = min(
        GEMINI_RETRY_CAP_SECONDS,
        GEMINI_RETRY_BASE_SECONDS * (2 ** max(0, attempt_index)),
    )
    delay = max(exponential, retry_after or 0.0)
    return min(GEMINI_RETRY_CAP_SECONDS, delay + random.uniform(0.25, 1.10))


def _gemini_rest_generate(api_key: str, model: str, prompt: str,
                          response_schema: dict | None = None) -> tuple[dict, str]:
    """إرسال طلب واحد إلى نموذج ثابت، مع إعادة المحاولة لنفس الطلب عند الأعطال المؤقتة فقط."""
    model = str(model or "").strip()
    if not model:
        raise GeminiAPIError("تعذر تحديد نموذج المراجعة الذكية المتاح لهذا المشروع.")

    endpoint = (
        "https://generativelanguage.googleapis.com/v1beta/models/"
        + urllib.parse.quote(model, safe="-._")
        + ":generateContent"
    )
    payload = {
        "system_instruction": {
            "parts": [{"text": "Follow the requested JSON schema exactly. Be conservative and evidence-based."}]
        },
        "contents": [{"role": "user", "parts": [{"text": prompt}]}],
        "generationConfig": {
            "responseMimeType": "application/json",
            "responseSchema": response_schema or _GEMINI_RESPONSE_SCHEMA,
            "maxOutputTokens": 3072,
        },
    }
    body = json.dumps(payload, ensure_ascii=False).encode("utf-8")

    last_error = None
    last_code = None
    for attempt in range(GEMINI_MAX_ATTEMPTS):
        # إنشاء Request جديد لكل محاولة يضمن عدم إعادة استخدام كائن اتصال مستهلك.
        request = urllib.request.Request(
            endpoint,
            data=body,
            headers={
                "Content-Type": "application/json",
                "x-goog-api-key": api_key,
                "User-Agent": "Academic-AI-Reviewer/1.0",
            },
            method="POST",
        )
        try:
            with urllib.request.urlopen(
                request,
                timeout=GEMINI_API_TIMEOUT_SECONDS,
                context=_gemini_ssl_context(),
            ) as response:
                response_data = json.loads(response.read().decode("utf-8"))

            candidates = response_data.get("candidates") or []
            if not candidates:
                feedback = response_data.get("promptFeedback") or {}
                raise GeminiAPIError(f"لم تُرجع خدمة المراجعة الذكية نتيجة صالحة. {feedback}")

            parts = (((candidates[0] or {}).get("content") or {}).get("parts") or [])
            raw_text = "\n".join(
                str(part.get("text") or "")
                for part in parts
                if isinstance(part, dict)
            )
            parsed = _extract_json_object(raw_text)
            return parsed, model

        except urllib.error.HTTPError as exc:
            last_code = exc.code
            try:
                detail = exc.read().decode("utf-8", errors="replace")
            except Exception:
                detail = str(exc)
            last_error = detail

            # لا نعيد المحاولة في أخطاء الإعداد أو الصلاحيات لأنها لن تتحسن بالانتظار.
            if exc.code == 400 and "free tier is not available" in detail.lower():
                raise GeminiAPIError("خدمة المراجعة الذكية غير متاحة ضمن إعدادات المشروع الحالية.")
            if exc.code == 400:
                raise GeminiAPIError("رفضت خدمة المراجعة الذكية الطلب بسبب إعداد غير صالح.")
            if exc.code == 401:
                raise GeminiAPIError("بيانات اتصال خدمة المراجعة الذكية غير صحيحة.")
            if exc.code == 403:
                raise GeminiAPIError("بيانات اتصال خدمة المراجعة الذكية غير مخوّلة لاستخدام النموذج الثابت.")
            if exc.code == 404:
                raise GeminiAPIError(
                    "نموذج المراجعة الذكية الثابت غير متاح لهذا المشروع. "
                    "يجب على إدارة التطبيق مراجعة إعدادات الخدمة."
                )

            # فقط الأخطاء المؤقتة/السعة/الحصة تعاد على النموذج نفسه.
            if exc.code in (429, 500, 502, 503, 504) and attempt < GEMINI_MAX_ATTEMPTS - 1:
                time.sleep(_smart_review_retry_delay(attempt, _retry_after_seconds(exc)))
                continue
            if exc.code == 429:
                raise GeminiAPIError(
                    "تعذر إكمال التحليل المشترك بعد عدة محاولات بسبب حد الاستخدام. لم يُصدر أي تقرير."
                )
            if exc.code in (500, 502, 503, 504):
                raise GeminiAPIError(
                    "تعذر إكمال التحليل المشترك بعد إعادة المحاولة على النموذج الثابت بسبب ضغط مؤقت في الخدمة. "
                    "لم يُصدر أي تقرير؛ أعد الفحص لاحقًا."
                )
            raise GeminiAPIError(f"فشل اتصال خدمة المراجعة الذكية (HTTP {exc.code}).")

        except urllib.error.URLError as exc:
            last_error = str(exc.reason)
            if attempt < GEMINI_MAX_ATTEMPTS - 1:
                time.sleep(_smart_review_retry_delay(attempt))
                continue

        except TimeoutError:
            last_error = "timeout"
            if attempt < GEMINI_MAX_ATTEMPTS - 1:
                time.sleep(_smart_review_retry_delay(attempt))
                continue

        except GeminiAPIError:
            # أخطاء الاستجابة غير الصالحة لا تستدعي نموذجًا آخر؛ يسمح بمحاولة جديدة لنفس النموذج فقط.
            if attempt < GEMINI_MAX_ATTEMPTS - 1:
                time.sleep(_smart_review_retry_delay(attempt))
                continue
            raise

    suffix = f" (HTTP {last_code})" if last_code else ""
    raise GeminiAPIError(
        f"تعذر إكمال التحليل المشترك عبر النموذج الثابت{suffix}: {last_error or 'unknown error'}. "
        "لم يُصدر أي تقرير."
    )

def _normalize_gemini_review(raw: dict, passages: list[dict], model: str) -> dict:
    passage_map = {p["passage_id"]: p for p in passages}

    def clamp_int(value, default=0):
        try:
            return int(max(0, min(100, round(float(value)))))
        except Exception:
            return default

    def clean_list(value, limit):
        if not isinstance(value, list):
            return []
        return [str(item).strip()[:700] for item in value if str(item).strip()][:limit]

    normalized_flags = []
    for flag in raw.get("flagged_passages") or []:
        if not isinstance(flag, dict):
            continue
        pid = str(flag.get("passage_id") or "").strip().upper()
        source = passage_map.get(pid)
        if not source:
            continue
        quote = str(flag.get("quote") or "").strip()
        if quote and quote not in source["text"]:
            quote = ""
        if not quote:
            quote = _truncate_words(source["text"], 28)
        try:
            severity = max(1, min(5, int(flag.get("severity", 1))))
        except Exception:
            severity = 1
        normalized_flags.append({
            "passage_id": pid,
            "category": str(flag.get("category") or "other"),
            "severity": severity,
            "reason_ar": str(flag.get("reason_ar") or "يحتاج هذا المقطع إلى مراجعة بشرية.").strip()[:700],
            "reason_en": str(flag.get("reason_en") or "This passage requires human review.").strip()[:700],
            "quote": quote[:500],
            "source_text": source["text"],
        })

    ai_likelihood_score = clamp_int(raw.get("ai_assistance_likelihood_score"))
    score = clamp_int(raw.get("review_priority_score"))
    citation_score = clamp_int(raw.get("citation_integrity_risk_score"))
    confidence = str(raw.get("confidence") or "low").lower()
    if confidence not in {"low", "medium", "high"}:
        confidence = "low"
    return {
        "available": True,
        "model": model,
        "ai_assistance_likelihood_score": ai_likelihood_score,
        "review_priority_score": score,
        "citation_integrity_risk_score": citation_score,
        "confidence": confidence,
        "verdict_ar": _gemini_review_label(score),
        "color": _gemini_review_color(score),
        "summary_ar": str(raw.get("summary_ar") or "لم تقدم خدمة المراجعة الذكية ملخصًا.").strip()[:1800],
        "summary_en": str(raw.get("summary_en") or "No summary was provided.").strip()[:1800],
        "key_reasons_ar": clean_list(raw.get("key_reasons_ar"), 6),
        "key_reasons_en": clean_list(raw.get("key_reasons_en"), 6),
        "counter_evidence_ar": clean_list(raw.get("counter_evidence_ar"), 5),
        "recommended_actions_ar": clean_list(raw.get("recommended_actions_ar"), 6),
        "recommended_actions_en": clean_list(raw.get("recommended_actions_en"), 6),
        "flagged_passages": normalized_flags[:20],
        "passages_reviewed": len(passages),
        "disclaimer_ar": "هذه درجة أولوية للمراجعة وليست نسبة اقتباس أو احتمالًا لإثبات أن الذكاء الاصطناعي كتب النص.",
    }


def _combined_probability(local_result: dict, gemini_review: dict) -> dict:
    """Combine independent screening signals with reliability-aware weights.

    This is an operational screening probability, not a scientifically calibrated
    probability of authorship. Disagreement lowers the displayed confidence rather
    than being hidden by an arbitrary threshold.
    """
    local_score = max(0.0, min(100.0, float(local_result.get("percentage", 0.0))))
    gemini_score = max(0.0, min(100.0, float(gemini_review.get("ai_assistance_likelihood_score", 0.0))))
    local_reliability = max(0.20, min(1.0, float(local_result.get("confidence_percentage", 0.0)) / 100.0))
    gemini_reliability = {"low": 0.45, "medium": 0.70, "high": 0.90}.get(
        str(gemini_review.get("confidence", "low")).lower(), 0.45
    )
    local_weight = 0.45 * local_reliability
    gemini_weight = 0.55 * gemini_reliability
    total_weight = max(local_weight + gemini_weight, 1e-9)
    combined = (local_score * local_weight + gemini_score * gemini_weight) / total_weight
    disagreement = abs(local_score - gemini_score)
    confidence_value = 100.0 * ((local_reliability + gemini_reliability) / 2.0) * (1.0 - 0.55 * disagreement / 100.0)
    confidence_value = max(0.0, min(100.0, confidence_value))
    if confidence_value >= 75:
        confidence_label = "مرتفعة"
    elif confidence_value >= 50:
        confidence_label = "متوسطة"
    else:
        confidence_label = "منخفضة"
    if combined >= 80:
        verdict = "احتمالية فحص مرتفعة جدًا"
    elif combined >= 60:
        verdict = "احتمالية فحص مرتفعة"
    elif combined >= 40:
        verdict = "احتمالية فحص متوسطة"
    elif combined >= 20:
        verdict = "احتمالية فحص محدودة"
    else:
        verdict = "احتمالية فحص منخفضة"
    return {
        "score": round(combined, 2),
        "local_score": round(local_score, 2),
        "gemini_score": round(gemini_score, 2),
        "local_weight": round(local_weight / total_weight, 4),
        "gemini_weight": round(gemini_weight / total_weight, 4),
        "disagreement": round(disagreement, 2),
        "confidence_score": round(confidence_value, 2),
        "confidence_label": confidence_label,
        "verdict_ar": verdict,
        "color": _gemini_review_color(combined),
        "disclaimer_ar": "هذه نسبة احتمالية فحص مركبة وليست دليلًا قاطعًا على التأليف بالذكاء الاصطناعي أو نسبة اقتباس.",
    }


def _gemini_assisted_analysis(analysis_text: str, local_result: dict,
                               api_key: str, preferred_model: str) -> dict:
    """تنفيذ مراجعة إلزامية بنموذج واحد مثبت طوال تشغيل التطبيق."""
    fixed_model = _resolve_fixed_gemini_model(api_key)
    passages = _select_gemini_passages(analysis_text, local_result)
    prompt = _gemini_prompt(passages, local_result)
    raw, used_model = _gemini_rest_generate(
        api_key,
        fixed_model,
        prompt,
    )
    review = _normalize_gemini_review(raw, passages, used_model)
    if not review.get("available"):
        raise GeminiAPIError("لم تكتمل المراجعة الذكية، ولذلك لن يُصدر التقرير.")
    return review


# ══════════════════════════════════════════════════════════════════════════════
# CUSTOMER RATINGS — تخزين بسيط دائم على القرص
# ══════════════════════════════════════════════════════════════════════════════
RATINGS_FILE = Path(__file__).resolve().parent / "customer_ratings.json"


def _seed_ratings():
    """تقييمات افتراضية تظهر أول مرة قبل ورود تقييمات حقيقية."""
    return [
        {"name": "أحمد سالم", "stars": 5,
         "comment": "أداة دقيقة جدًا ووفّرت عليّ وقتًا كبيرًا في مراجعة الأبحاث.",
         "date": "2026-06-14"},
        {"name": "منى عبد الله", "stars": 4,
         "comment": "التقرير المظلل فوق ملف PDF ممتاز، أتمنى دعم صيغ أخرى مستقبلاً.",
         "date": "2026-06-20"},
        {"name": "خالد يوسف", "stars": 5,
         "comment": "التحديث الجديد حسّن سرعة التحليل بشكل ملحوظ.",
         "date": "2026-07-05"},
        {"name": "سارة إبراهيم", "stars": 4,
         "comment": "واجهة واضحة وسهلة الاستخدام، شكرًا على الجهد المبذول.",
         "date": "2026-07-14"},
        {"name": "محمد الطاهر", "stars": 5,
         "comment": "نتائج موثوقة وتقرير احترافي جاهز للطباعة مباشرة.",
         "date": "2026-07-22"},
    ]


def _load_ratings():
    try:
        if RATINGS_FILE.exists():
            data = json.loads(RATINGS_FILE.read_text(encoding="utf-8"))
            if isinstance(data, list) and data:
                return data
    except Exception:
        pass
    return _seed_ratings()


def _save_rating(entry: dict):
    ratings = _load_ratings()
    ratings.append(entry)
    try:
        RATINGS_FILE.write_text(
            json.dumps(ratings, ensure_ascii=False, indent=2), encoding="utf-8"
        )
    except Exception:
        pass  # في بيئات القراءة فقط لن يُحفظ التقييم بشكل دائم، لكن الواجهة تستمر بالعمل
    return ratings


# ══════════════════════════════════════════════════════════════════════════════
# FINGERPRINT ENGINE
# ══════════════════════════════════════════════════════════════════════════════

# Internal implementation details are intentionally not exposed in the user
# interface. This build remains local, CPU-friendly, and English-only.


# تجهيز النص قبل التحليل: استبعاد المراجع والاستشهادات داخل المتن
# ══════════════════════════════════════════════════════════════════════════════

def _count_text_words(text: str) -> int:
    return len(re.findall(r"[A-Za-zÀ-ÖØ-öø-ÿ]+(?:['’-][A-Za-zÀ-ÖØ-öø-ÿ]+)?", text or ''))


def _normalize_reference_heading(text: str) -> str:
    """Normalize an English reference-section heading before matching."""
    s = re.sub(r'\s+', ' ', text or '').strip()
    s = re.sub(r'^[\s\-–—:;,.()\[\]]+', '', s)
    s = re.sub(r'[\s\-–—:;,.()\[\]]+$', '', s)
    return s


REF_SECTION_HEADERS = re.compile(
    r'''^(?:(?:(?:chapter|section|appendix)\s+)?(?:\d+(?:\.\d+)*)?\s*[:.\-–—]?\s*)?(?:references?|reference\s+list|bibliography|works\s+cited|literature\s+cited|sources?)$''',
    re.I | re.X
)


def _is_reference_header(text: str) -> bool:
    """
    يكتشف عنوان قسم المراجع حتى لو سبقه رقم فصل أو تبعته نقطتان.
    لا يعتمد على وجود كلمة References داخل فقرة طويلة.
    """
    for raw_line in re.split(r'[\r\n]+', text or ''):
        line = _normalize_reference_heading(raw_line)
        if not line:
            continue
        if len(line.split()) <= 8 and REF_SECTION_HEADERS.fullmatch(line):
            return True
    return False


REF_LINE_PATTERNS = [
    r'^\s*\[\s*\d+(?:\s*[-–,]\s*\d+)*\s*\]',
    r'^\s*\d{1,4}\s*[.)]\s+',
    r'\bdoi\s*:\s*10\.\d{4,9}/\S+',
    r'\bhttps?://\S+',
    r'\b(?:ISBN|ISSN)\b',
    r'\b(?:Vol\.?|Volume)\s*\d+',
    r'\b(?:No\.?|Issue)\s*\d+',
    r'\bpp?\.\s*\d+(?:\s*[-–]\s*\d+)?',
    r'\b(?:Journal|Proceedings|Transactions|Conference|Press|Publisher)\b',
    r'\bet\s+al\.?',
]


def _reference_line_score(line: str) -> float:
    """درجة احتمالية أن السطر مدخل ببليوجرافي وليس فقرة من المتن."""
    s = re.sub(r'\s+', ' ', (line or '').strip())
    if not s:
        return 0.0

    score = 0.0
    if re.search(r'^\s*\[\s*\d+(?:\s*[-–,]\s*\d+)*\s*\]', s):
        score += 2.5
    if re.search(r'^\s*\d{1,4}\s*[.)]\s+', s):
        score += 1.8
    if re.search(r'\b(?:18|19|20)\d{2}[a-z]?\b', s, re.I):
        score += 0.8
    if re.search(r'\bdoi\s*:|doi\.org/10\.|\b10\.\d{4,9}/\S+', s, re.I):
        score += 2.2
    if re.search(r'https?://', s, re.I):
        score += 1.7
    if re.search(r'\bet\s+al\.?\b', s, re.I):
        score += 1.0
    if re.search(r'\b(?:Journal|Proceedings|Transactions|Conference|Press|Publisher|Springer|Elsevier|Wiley|IEEE|ACM)\b', s, re.I):
        score += 1.1
    if re.search(r'\b(?:Vol\.?|Volume|No\.?|Issue|pp?\.)\s*\d+', s, re.I):
        score += 1.0
    if re.search(r'\b(?:ISBN|ISSN)\b', s, re.I):
        score += 1.5
    if re.search(r'^[A-ZÀ-ÖØ-Ý][A-Za-zÀ-ÖØ-öø-ÿ\'’-]+,\s*(?:[A-Z]\.?\s*){1,4}', s):
        score += 1.2
    if re.search(r'\([12]\d{3}[a-z]?\)', s, re.I):
        score += 0.8
    if s.count(',') >= 2 and s.count('.') >= 2:
        score += 0.5
    if len(s) >= 45:
        score += 0.2
    return score


def _is_reference_line(line: str) -> bool:
    """
    تصنيف محافظ للسطر المرجعي. الأنماط القوية تكفي منفردة،
    وإلا نحتاج اجتماع أكثر من علامة ببليوجرافية لتجنب حذف متن الباحث.
    """
    s = re.sub(r'\s+', ' ', (line or '').strip())
    if not s:
        return False
    if _is_reference_header(s):
        return True
    if re.search(r'^\s*\[\s*\d+(?:\s*[-–,]\s*\d+)*\s*\]', s):
        return True
    if re.search(r'\bdoi\s*:|doi\.org/10\.|\b10\.\d{4,9}/\S+', s, re.I):
        return True
    return _reference_line_score(s) >= 2.6


_INLINE_NUMERIC_CITATION_RE = re.compile(
    r'\[(?:\s*\d+\s*(?:[-–]\s*\d+)?\s*)(?:[,;]\s*\d+\s*(?:[-–]\s*\d+)?\s*)*\]'
)
_INLINE_AUTHOR_YEAR_RE = re.compile(
    r'''\(
        (?=[^()]{0,180}\b(?:18|19|20)\d{2}[a-z]?\b)
        (?:[^()]*?[A-ZÀ-ÖØ-Ý][^()]*?)
        (?:18|19|20)\d{2}[a-z]?
        (?:\s*[,;]\s*[^()]*?(?:18|19|20)\d{2}[a-z]?)*
    \)''',
    re.I | re.X
)
_INLINE_YEAR_ONLY_RE = re.compile(r'\((?:18|19|20)\d{2}[a-z]?\)')
_INLINE_URL_RE = re.compile(r'https?://\S+|www\.\S+', re.I)
_INLINE_DOI_RE = re.compile(r'\b(?:doi\s*:\s*)?10\.\d{4,9}/\S+', re.I)


def _strip_inline_citations(text: str) -> str:
    """يحذف علامات الاستشهاد فقط ويُبقي الجملة العلمية نفسها للتحليل."""
    s = text or ''
    s = _INLINE_NUMERIC_CITATION_RE.sub(' ', s)
    s = _INLINE_AUTHOR_YEAR_RE.sub(' ', s)
    s = _INLINE_YEAR_ONLY_RE.sub(' ', s)
    s = _INLINE_DOI_RE.sub(' ', s)
    s = _INLINE_URL_RE.sub(' ', s)
    s = re.sub(r'[ \t]+', ' ', s)
    s = re.sub(r' *\n *', '\n', s)
    s = re.sub(r'\n{3,}', '\n\n', s)
    return s.strip()


def _find_unheaded_reference_tail(lines) -> int | None:
    """
    إذا غاب عنوان المراجع، يبحث فقط في الجزء الأخير من المستند عن كتلة
    كثيفة من المداخل الببليوجرافية. هذا يمنع إسقاط فقرات وسط البحث.
    """
    nonempty = [(i, line.strip()) for i, line in enumerate(lines) if line.strip()]
    if len(nonempty) < 6:
        return None

    earliest_pos = max(0, int(len(nonempty) * 0.50))
    for pos in range(earliest_pos, len(nonempty)):
        window = nonempty[pos:pos + 6]
        if len(window) < 3:
            continue
        ref_hits = sum(1 for _, line in window if _is_reference_line(line))
        if ref_hits < 3:
            continue
        first_ref_offset = next(
            (j for j, (_, line) in enumerate(window) if _is_reference_line(line)),
            None
        )
        if first_ref_offset is None:
            continue
        tail = nonempty[pos + first_ref_offset:]
        density = sum(1 for _, line in tail if _is_reference_line(line)) / max(len(tail), 1)
        if density >= 0.55:
            return tail[0][0]
    return None


def _prepare_analysis_text(text: str) -> dict:
    """
    يفصل متن البحث عن قسم المراجع ويزيل الاستشهادات داخل المتن.
    يعيد إحصاءات الاستبعاد لتسهيل المراجعة والاختبار.
    """
    raw_text = text or ''
    lines = raw_text.replace('\r\n', '\n').replace('\r', '\n').split('\n')

    ref_start = None
    for idx, line in enumerate(lines):
        if _is_reference_header(line):
            ref_start = idx
            break

    if ref_start is None:
        ref_start = _find_unheaded_reference_tail(lines)

    if ref_start is None:
        main_lines = lines
        ref_lines = []
        header_found = False
    else:
        main_lines = lines[:ref_start]
        ref_lines = lines[ref_start:]
        header_found = any(_is_reference_header(line) for line in ref_lines[:3])

    main_text_raw = '\n'.join(main_lines)
    analysis_text = _strip_inline_citations(main_text_raw)
    references_text = '\n'.join(ref_lines)

    return {
        'analysis_text': analysis_text,
        'main_text_raw': main_text_raw,
        'references_text': references_text,
        'reference_section_found': ref_start is not None,
        'reference_header_found': header_found,
        'original_words': _count_text_words(raw_text),
        'analyzed_words': _count_text_words(analysis_text),
        'reference_words_excluded': _count_text_words(references_text),
        'inline_citation_words_removed': max(
            0,
            _count_text_words(main_text_raw) - _count_text_words(analysis_text)
        ),
    }


LOW_SCORE_HIGHLIGHT_THRESHOLD = 20


def _public_whole_percentage(value: float) -> int | None:
    """تحويل النتيجة إلى رقم صحيح واحد باستخدام تقريب نصف الدرجة إلى أعلى."""
    try:
        pct = float(value)
    except (TypeError, ValueError):
        return None
    if not math.isfinite(pct):
        return None
    pct = max(0.0, min(100.0, pct))
    return int(math.floor(pct + 0.5))


def _low_score_is_masked(value: float) -> bool:
    """إخفاء أي نتيجة يكون رقمها الظاهر بين 0 و20 شاملًا."""
    public_pct = _public_whole_percentage(value)
    return public_pct is None or 0 <= public_pct <= LOW_SCORE_HIGHLIGHT_THRESHOLD


def _format_percentage(value: float, mask_low: bool = True) -> str:
    """عرض *% من 0 إلى 20، وإظهار الرقم الحقيقي ابتداءً من 21%."""
    public_pct = _public_whole_percentage(value)
    if public_pct is None:
        return '*%' if mask_low else '0%'
    if mask_low and _low_score_is_masked(value):
        return '*%'
    return f'{public_pct}%'


def _format_human_percentage(ai_value: float, human_value: float) -> str:
    """منع كشف النسبة المخفية بصورة عكسية عن طريق نسبة الكتابة البشرية."""
    try:
        human_pct = float(human_value)
    except (TypeError, ValueError):
        return 'غير محدد'
    if _low_score_is_masked(ai_value):
        return 'مرتفع'
    public_human = _public_whole_percentage(human_pct)
    return f'{public_human}%' if public_human is not None else 'غير محدد'


def _highlighting_allowed(ai_percentage: float) -> bool:
    """بدء التظليل فقط عندما يصبح الرقم النهائي الظاهر 21% فأعلى."""
    public_pct = _public_whole_percentage(ai_percentage)
    return public_pct is not None and public_pct > LOW_SCORE_HIGHLIGHT_THRESHOLD


# ══════════════════════════════════════════════════════════════════════════════
# FINGERPRINT ENGINE
# ══════════════════════════════════════════════════════════════════════════════
# ENGLISH-ONLY INTERNAL MULTI-LAYER FINGERPRINT ENGINE
# ══════════════════════════════════════════════════════════════════════════════

AI_ENGINE_VERSION = "5.5.1 Internal Calibrated Hierarchical Academic Engine"
PUBLIC_ENGINE_LABEL = "محرك الفحص الداخلي متعدد المؤشرات"
MIN_ENGLISH_WORDS = 80
STRICT_FINGERPRINT_POLICY = True
ADVERSARIAL_RESISTANCE_POLICY = True
HIERARCHICAL_ACADEMIC_POLICY = True

_STYLE_WORD_RE = re.compile(r"[A-Za-zÀ-ÖØ-öø-ÿ]+(?:['’-][A-Za-zÀ-ÖØ-öø-ÿ]+)?")
_SENTENCE_END_RE = re.compile(r"(?<=[.!?])\s+|\n\s*\n+")

_ENGLISH_FUNCTION_WORDS = {
    'a','about','above','after','again','against','all','am','an','and','any','are','as','at',
    'be','because','been','before','being','below','between','both','but','by','can','could',
    'did','do','does','doing','down','during','each','few','for','from','further','had','has',
    'have','having','he','her','here','hers','herself','him','himself','his','how','i','if',
    'in','into','is','it','its','itself','just','me','more','most','my','myself','no','nor',
    'not','now','of','off','on','once','only','or','other','our','ours','ourselves','out',
    'over','own','same','she','should','so','some','such','than','that','the','their','theirs',
    'them','themselves','then','there','these','they','this','those','through','to','too','under',
    'until','up','very','was','we','were','what','when','where','which','while','who','whom',
    'why','will','with','would','you','your','yours','yourself','yourselves'
}

_PUNCTUATION_MARKS = ['.', ',', ';', ':', '!', '?', '—', '-', '(', ')', '"', "'"]


class EnglishOnlyError(ValueError):
    """Raised when the uploaded text is not sufficiently English for this engine."""


class InsufficientEnglishTextError(ValueError):
    """Raised when the English text is too short for a defensible fingerprint."""


def _clamp(value: float, low: float = 0.0, high: float = 1.0) -> float:
    return max(low, min(high, float(value)))


def _safe_mean(values) -> float:
    vals = [float(v) for v in values if v is not None]
    return sum(vals) / len(vals) if vals else 0.0


def _safe_std(values) -> float:
    vals = [float(v) for v in values if v is not None]
    if len(vals) < 2:
        return 0.0
    mean = sum(vals) / len(vals)
    return math.sqrt(sum((v - mean) ** 2 for v in vals) / len(vals))


def _percentile(values, q: float) -> float:
    vals = sorted(float(v) for v in values)
    if not vals:
        return 0.0
    if len(vals) == 1:
        return vals[0]
    pos = _clamp(q) * (len(vals) - 1)
    lo = int(math.floor(pos))
    hi = int(math.ceil(pos))
    if lo == hi:
        return vals[lo]
    frac = pos - lo
    return vals[lo] * (1.0 - frac) + vals[hi] * frac


def _median_absolute_deviation(values) -> float:
    vals = [float(v) for v in values]
    if not vals:
        return 0.0
    med = _percentile(vals, 0.50)
    return _percentile([abs(v - med) for v in vals], 0.50)


def _sigmoid(x: float) -> float:
    if x >= 35:
        return 1.0
    if x <= -35:
        return 0.0
    return 1.0 / (1.0 + math.exp(-x))


def _normalized_entropy(counter: collections.Counter) -> float:
    total = sum(counter.values())
    if total <= 1 or len(counter) <= 1:
        return 0.0
    entropy = -sum((c / total) * math.log(c / total, 2) for c in counter.values() if c > 0)
    return _clamp(entropy / math.log(len(counter), 2))


def _normalize_style_text(text: str) -> str:
    s = (text or '').replace('\u00ad', '')
    s = re.sub(r'(?<=\w)-\s*\n\s*(?=\w)', '', s)
    s = re.sub(r'[ \t]+', ' ', s)
    return s.strip()


def _style_tokens(text: str) -> list[str]:
    return [m.group(0).lower().replace('’', "'") for m in _STYLE_WORD_RE.finditer(_normalize_style_text(text))]


def _english_language_profile(text: str) -> dict:
    raw = text or ''
    tokens = _style_tokens(raw)
    alpha_chars = [ch for ch in raw if ch.isalpha()]
    latin_chars = sum('LATIN' in unicodedata.name(ch, '') for ch in alpha_chars)
    arabic_chars = sum('ARABIC' in unicodedata.name(ch, '') for ch in alpha_chars)
    script_chars = latin_chars + arabic_chars
    latin_ratio = latin_chars / max(script_chars, 1)
    function_ratio = sum(t in _ENGLISH_FUNCTION_WORDS for t in tokens) / max(len(tokens), 1)
    long_word_ratio = sum(len(t) >= 4 for t in tokens) / max(len(tokens), 1)

    # The engine intentionally accepts only predominantly Latin-script English.
    confidence = _clamp(
        0.62 * _clamp((latin_ratio - 0.82) / 0.18) +
        0.30 * _clamp((function_ratio - 0.035) / 0.15) +
        0.08 * _clamp((long_word_ratio - 0.35) / 0.45)
    )
    supported = (
        len(tokens) >= MIN_ENGLISH_WORDS and
        latin_ratio >= 0.965 and
        arabic_chars <= max(2, int(max(script_chars, 1) * 0.005)) and
        function_ratio >= 0.055
    )
    return {
        'supported': supported,
        'confidence': confidence,
        'word_count': len(tokens),
        'latin_ratio': latin_ratio,
        'arabic_alpha_chars': arabic_chars,
        'function_word_ratio': function_ratio,
    }


def _validate_english_text(text: str) -> dict:
    profile = _english_language_profile(text)
    script_mismatch = (
        profile['arabic_alpha_chars'] > 2 or
        (profile['word_count'] >= 10 and profile['latin_ratio'] < 0.965) or
        (profile['word_count'] >= 30 and profile['function_word_ratio'] < 0.055)
    )
    if script_mismatch:
        raise EnglishOnlyError(
            "This build analyzes English text only. Arabic, mixed-language, and non-English documents are rejected."
        )
    if profile['word_count'] < MIN_ENGLISH_WORDS:
        raise InsufficientEnglishTextError(
            f"The analysis requires at least {MIN_ENGLISH_WORDS} English words."
        )
    if not profile['supported']:
        raise EnglishOnlyError(
            "This build analyzes English text only. Arabic, mixed-language, and non-English documents are rejected."
        )
    return profile


def _split_analysis_sentences(text: str) -> list[str]:
    s = (text or '').replace('\r\n', '\n').replace('\r', '\n')
    s = re.sub(r'(?<![.!?:;])\n(?!\s*\n)', ' ', s)
    raw = _SENTENCE_END_RE.split(s)
    out = []
    for item in raw:
        item = re.sub(r'\s+', ' ', item).strip()
        if len(_style_tokens(item)) >= 3:
            out.append(item)
    return out


def _split_analysis_paragraphs(text: str) -> list[str]:
    paragraphs = []
    for p in re.split(r'\n\s*\n+', text or ''):
        p = re.sub(r'\s+', ' ', p).strip()
        if len(_style_tokens(p)) >= 12:
            paragraphs.append(p)
    return paragraphs


def _build_analysis_chunks(sentences: list[str], target_words: int = 190,
                           min_words: int = 90, max_words: int = 285,
                           overlap_sentences: int = 2) -> list[str]:
    if not sentences:
        return []
    chunks = []
    start = 0
    while start < len(sentences):
        current = []
        current_words = 0
        end = start
        while end < len(sentences):
            sw = len(_style_tokens(sentences[end]))
            if current and current_words + sw > max_words:
                break
            current.append(sentences[end])
            current_words += sw
            end += 1
            if current_words >= target_words:
                break
        if current_words >= min_words or (not chunks and current_words > 0):
            chunks.append(' '.join(current))
        if end >= len(sentences):
            break
        start = max(start + 1, end - max(0, overlap_sentences))
    return chunks


def _moving_ttr(tokens: list[str], window: int = 50) -> tuple[float, float]:
    if not tokens:
        return 0.0, 0.0
    if len(tokens) <= window:
        return len(set(tokens)) / max(len(tokens), 1), 0.0
    vals = []
    step = max(1, window // 5)
    for i in range(0, len(tokens) - window + 1, step):
        seg = tokens[i:i + window]
        vals.append(len(set(seg)) / window)
    mean = _safe_mean(vals)
    cv = _safe_std(vals) / mean if mean else 0.0
    return mean, cv


def _lexical_profile(tokens: list[str]) -> dict:
    if not tokens:
        return {
            'hapax_ratio': 0.0, 'simpson': 0.0, 'entropy': 0.0,
            'mattr40': 0.0, 'mattr80': 0.0, 'mattr_cv': 0.0,
            'yule_k': 0.0, 'zipf_r2': 0.0, 'zipf_slope': 0.0,
        }
    counts = collections.Counter(tokens)
    n = len(tokens)
    v = len(counts)
    hapax_ratio = sum(1 for c in counts.values() if c == 1) / max(v, 1)
    simpson = sum(c * (c - 1) for c in counts.values()) / max(n * (n - 1), 1)
    entropy = _normalized_entropy(counts)
    mattr40, cv40 = _moving_ttr(tokens, 40)
    mattr80, cv80 = _moving_ttr(tokens, 80)
    mattr_cv = _safe_mean([cv40, cv80])
    m2 = sum(c * c for c in counts.values())
    yule_k = 10000.0 * (m2 - n) / max(n * n, 1)

    freqs = sorted(counts.values(), reverse=True)
    freqs = freqs[:min(300, len(freqs))]
    if len(freqs) >= 8:
        xs = [math.log(i + 1) for i in range(len(freqs))]
        ys = [math.log(f) for f in freqs]
        xm, ym = _safe_mean(xs), _safe_mean(ys)
        denom = sum((x - xm) ** 2 for x in xs)
        slope = sum((x - xm) * (y - ym) for x, y in zip(xs, ys)) / max(denom, 1e-9)
        intercept = ym - slope * xm
        ss_res = sum((y - (intercept + slope * x)) ** 2 for x, y in zip(xs, ys))
        ss_tot = sum((y - ym) ** 2 for y in ys)
        r2 = 1.0 - ss_res / max(ss_tot, 1e-9)
    else:
        slope, r2 = 0.0, 0.0
    return {
        'hapax_ratio': hapax_ratio, 'simpson': simpson, 'entropy': entropy,
        'mattr40': mattr40, 'mattr80': mattr80, 'mattr_cv': mattr_cv,
        'yule_k': yule_k, 'zipf_r2': _clamp(r2), 'zipf_slope': slope,
    }


def _word_ngram_repetition(tokens: list[str]) -> tuple[float, float, float]:
    ratios = []
    for n in (2, 3, 4):
        grams = [tuple(tokens[i:i+n]) for i in range(max(0, len(tokens) - n + 1))]
        if not grams:
            ratios.append(0.0)
            continue
        counts = collections.Counter(grams)
        repeated = sum(max(0, c - 1) for c in counts.values()) / len(grams)
        ratios.append(repeated)
    weighted = ratios[0] * 0.20 + ratios[1] * 0.35 + ratios[2] * 0.45
    score = _clamp((weighted - 0.006) / 0.055)
    return weighted, score, max(ratios)


def _character_ngram_profile(text: str) -> tuple[float, float, float]:
    compact = re.sub(r'\s+', ' ', (text or '').lower()).strip()
    if len(compact) < 250:
        return 0.0, 0.0, 0.0
    repeat_values = []
    entropy_values = []
    for n in (3, 4, 5):
        grams = [compact[i:i+n] for i in range(len(compact) - n + 1)]
        counts = collections.Counter(grams)
        repeat_values.append(sum(max(0, c - 1) for c in counts.values()) / max(len(grams), 1))
        entropy_values.append(_normalized_entropy(counts))
    repeat_ratio = _safe_mean(repeat_values)
    entropy = _safe_mean(entropy_values)
    score = _clamp((repeat_ratio - 0.18) / 0.30)
    if entropy > 0.965:
        score *= 0.82
    return repeat_ratio, entropy, score


def _opener_repetition(sentences: list[str]) -> tuple[float, float]:
    openers = []
    for sentence in sentences:
        toks = _style_tokens(sentence)
        if len(toks) >= 3:
            openers.append(tuple(toks[:3]))
    if len(openers) < 5:
        return 0.0, 0.0
    counts = collections.Counter(openers)
    ratio = sum(max(0, c - 1) for c in counts.values()) / len(openers)
    return ratio, _clamp((ratio - 0.025) / 0.20)


def _sentence_template_repetition(sentences: list[str]) -> tuple[float, float]:
    templates = []
    for sentence in sentences:
        toks = _style_tokens(sentence)
        if len(toks) < 6:
            continue
        skeleton = []
        for token in toks[:24]:
            if token in _ENGLISH_FUNCTION_WORDS:
                skeleton.append('F')
            elif len(token) <= 3:
                skeleton.append('S')
            elif len(token) >= 9:
                skeleton.append('L')
            else:
                skeleton.append('C')
        compressed = []
        for symbol in skeleton:
            if not compressed or compressed[-1] != symbol:
                compressed.append(symbol)
        templates.append(tuple(compressed[:12]))
    if len(templates) < 6:
        return 0.0, 0.0
    counts = collections.Counter(templates)
    ratio = sum(max(0, c - 1) for c in counts.values()) / len(templates)
    return ratio, _clamp((ratio - 0.04) / 0.30)


def _lag1_autocorrelation(values) -> float:
    vals = [float(v) for v in values]
    if len(vals) < 5:
        return 0.0
    mean = _safe_mean(vals)
    denom = sum((v - mean) ** 2 for v in vals)
    if denom <= 1e-9:
        return 1.0
    return sum((vals[i] - mean) * (vals[i-1] - mean) for i in range(1, len(vals))) / denom


def _sentence_rhythm_profile(sentences: list[str]) -> dict:
    lengths = [len(_style_tokens(s)) for s in sentences]
    if not lengths:
        return {
            'mean': 0.0, 'cv': 0.0, 'mad_ratio': 0.0, 'qspread': 0.0,
            'autocorr': 0.0, 'outlier_share': 0.0, 'score': 0.0,
        }
    mean = _safe_mean(lengths)
    cv = _safe_std(lengths) / mean if mean else 0.0
    med = _percentile(lengths, 0.50)
    mad_ratio = _median_absolute_deviation(lengths) / max(med, 1.0)
    qspread = (_percentile(lengths, 0.90) - _percentile(lengths, 0.10)) / max(med, 1.0)
    autocorr = _lag1_autocorrelation(lengths)
    outlier_share = sum(abs(v - med) > max(2.5 * _median_absolute_deviation(lengths), 8.0) for v in lengths) / len(lengths)
    regularity = (
        _clamp((0.64 - cv) / 0.42) * 0.36 +
        _clamp((0.48 - mad_ratio) / 0.34) * 0.26 +
        _clamp((2.2 - qspread) / 1.55) * 0.20 +
        _clamp((0.12 - outlier_share) / 0.12) * 0.12 +
        _clamp((autocorr + 0.12) / 0.55) * 0.06
    )
    reliability = min(1.0, len(lengths) / 12.0)
    return {
        'mean': mean, 'cv': cv, 'mad_ratio': mad_ratio, 'qspread': qspread,
        'autocorr': autocorr, 'outlier_share': outlier_share,
        'score': _clamp(regularity * reliability),
    }


def _paragraph_uniformity(paragraphs: list[str]) -> tuple[float, float]:
    lengths = [len(_style_tokens(p)) for p in paragraphs]
    if len(lengths) < 3:
        return 0.0, 0.0
    mean = _safe_mean(lengths)
    cv = _safe_std(lengths) / mean if mean else 0.0
    score = _clamp((0.78 - cv) / 0.58) * min(1.0, len(lengths) / 7.0)
    return cv, score


def _count_syllables(word: str) -> int:
    w = re.sub(r'[^a-z]', '', word.lower())
    if not w:
        return 0
    groups = re.findall(r'[aeiouy]+', w)
    count = len(groups)
    if w.endswith('e') and not w.endswith(('le', 'ye')) and count > 1:
        count -= 1
    return max(1, count)


def _readability_values(texts: list[str]) -> tuple[list[float], list[float]]:
    flesch = []
    word_lengths = []
    for text in texts:
        toks = _style_tokens(text)
        sents = _split_analysis_sentences(text)
        if len(toks) < 20 or not sents:
            continue
        syllables = sum(_count_syllables(t) for t in toks)
        fre = 206.835 - 1.015 * (len(toks) / len(sents)) - 84.6 * (syllables / len(toks))
        flesch.append(fre)
        word_lengths.append(_safe_mean([len(t) for t in toks]))
    return flesch, word_lengths


def _js_divergence(p, q) -> float:
    p = [max(0.0, float(x)) for x in p]
    q = [max(0.0, float(x)) for x in q]
    ps, qs = sum(p), sum(q)
    if ps <= 0 or qs <= 0:
        return 0.0
    p = [x / ps for x in p]
    q = [x / qs for x in q]
    m = [(a + b) / 2.0 for a, b in zip(p, q)]
    def _kl(a, b):
        return sum(x * math.log(x / y, 2) for x, y in zip(a, b) if x > 0 and y > 0)
    return 0.5 * _kl(p, m) + 0.5 * _kl(q, m)


def _segment_tokens(tokens: list[str], min_size: int = 70, max_segments: int = 8) -> list[list[str]]:
    if len(tokens) < min_size * 2:
        return []
    n_segments = min(max_segments, max(2, len(tokens) // min_size))
    boundaries = [round(i * len(tokens) / n_segments) for i in range(n_segments + 1)]
    return [tokens[boundaries[i]:boundaries[i+1]] for i in range(n_segments) if boundaries[i+1] > boundaries[i]]


def _function_word_stability(tokens: list[str]) -> tuple[float, float]:
    segments = _segment_tokens(tokens)
    if len(segments) < 2:
        return 0.0, 0.0
    vocab = sorted(_ENGLISH_FUNCTION_WORDS)
    vectors = []
    for seg in segments:
        counts = collections.Counter(seg)
        vectors.append([counts[w] + 0.15 for w in vocab])
    divergences = []
    for i in range(len(vectors) - 1):
        divergences.append(_js_divergence(vectors[i], vectors[i+1]))
    mean_js = _safe_mean(divergences)
    stability = _clamp((0.085 - mean_js) / 0.075)
    return mean_js, stability


def _punctuation_stability(sentences: list[str]) -> tuple[float, float, int]:
    if len(sentences) < 8:
        return 0.0, 0.0, 0
    n_groups = min(6, max(2, len(sentences) // 4))
    bounds = [round(i * len(sentences) / n_groups) for i in range(n_groups + 1)]
    vectors = []
    observed = set()
    for i in range(n_groups):
        group = ' '.join(sentences[bounds[i]:bounds[i+1]])
        vec = [group.count(p) + 0.05 for p in _PUNCTUATION_MARKS]
        vectors.append(vec)
        observed.update(p for p in _PUNCTUATION_MARKS if group.count(p))
    divergences = [_js_divergence(vectors[i], vectors[i+1]) for i in range(len(vectors) - 1)]
    mean_js = _safe_mean(divergences)
    score = _clamp((0.12 - mean_js) / 0.105)
    if len(observed) >= 8:
        score *= 0.86
    return mean_js, _clamp(score), len(observed)


def _lexical_stability(tokens: list[str]) -> tuple[float, float, float]:
    segments = _segment_tokens(tokens)
    if len(segments) < 2:
        return 0.0, 0.0, 0.0
    mattrs = []
    entropies = []
    for seg in segments:
        prof = _lexical_profile(seg)
        mattrs.append(prof['mattr40'])
        entropies.append(prof['entropy'])
    mattr_mean = _safe_mean(mattrs)
    mattr_cv = _safe_std(mattrs) / mattr_mean if mattr_mean else 0.0
    entropy_cv = _safe_std(entropies) / max(_safe_mean(entropies), 1e-9)
    score = _clamp((0.11 - mattr_cv) / 0.095) * 0.62 + _clamp((0.045 - entropy_cv) / 0.04) * 0.38
    return mattr_cv, entropy_cv, _clamp(score)


def _compression_profile(text: str) -> tuple[float, float]:
    raw = (text or '').encode('utf-8', errors='ignore')
    if len(raw) < 500:
        return 1.0, 0.0
    compressed = zlib.compress(raw, level=9)
    ratio = len(compressed) / len(raw)
    # Lower ratios mean more repeated/predictable surface structure.
    score = _clamp((0.57 - ratio) / 0.22)
    return ratio, score


def _chunk_vector(text: str) -> list[float]:
    tokens = _style_tokens(text)
    sentences = _split_analysis_sentences(text)
    lex = _lexical_profile(tokens)
    rhythm = _sentence_rhythm_profile(sentences)
    wrep, _, _ = _word_ngram_repetition(tokens)
    crep, cent, _ = _character_ngram_profile(text)
    punct_total = max(len(text), 1)
    return [
        rhythm['mean'] / 35.0,
        rhythm['cv'],
        lex['mattr40'],
        lex['hapax_ratio'],
        lex['entropy'],
        min(1.0, lex['yule_k'] / 250.0),
        min(1.0, wrep * 10.0),
        min(1.0, crep),
        cent,
        text.count(',') / punct_total * 100.0,
        text.count(';') / punct_total * 100.0,
        text.count(':') / punct_total * 100.0,
    ]


def _cosine_similarity(a, b) -> float:
    dot = sum(x * y for x, y in zip(a, b))
    na = math.sqrt(sum(x * x for x in a))
    nb = math.sqrt(sum(y * y for y in b))
    return dot / (na * nb) if na and nb else 0.0


def _chunk_style_consistency(chunks: list[str]) -> tuple[float, float]:
    if len(chunks) < 2:
        return 0.0, 0.0
    vectors = [_chunk_vector(c) for c in chunks]
    sims = []
    for i in range(len(vectors)):
        for j in range(i + 1, len(vectors)):
            sims.append(_cosine_similarity(vectors[i], vectors[j]))
    similarity = _safe_mean(sims)
    score = _clamp((similarity - 0.88) / 0.105)
    return similarity, score


def _stylometric_symbol_sequence(text: str) -> list[str]:
    """
    Converts English prose into a topic-reduced structural sequence.

    Function words keep their identity because they are strong stylometric
    markers. Content words are reduced to length and suffix classes so the
    verifier measures writing structure rather than subject terminology.
    """
    symbols = []
    suffixes = ('tion', 'sion', 'ment', 'ness', 'ity', 'ive', 'ous', 'able',
                'ible', 'ally', 'ing', 'ed', 'ly', 'al')
    for sentence in _split_analysis_sentences(text):
        tokens = _style_tokens(sentence)
        if not tokens:
            continue
        n = len(tokens)
        if n <= 9:
            symbols.append('<SENT_SHORT>')
        elif n <= 18:
            symbols.append('<SENT_MEDIUM>')
        elif n <= 30:
            symbols.append('<SENT_LONG>')
        else:
            symbols.append('<SENT_XLONG>')

        for token in tokens:
            if token in _ENGLISH_FUNCTION_WORDS:
                symbols.append('F:' + token)
                continue
            if len(token) <= 4:
                length_class = 'S'
            elif len(token) <= 7:
                length_class = 'M'
            elif len(token) <= 10:
                length_class = 'L'
            else:
                length_class = 'X'
            suffix_class = next((sf for sf in suffixes if token.endswith(sf) and len(token) > len(sf) + 2), 'other')
            symbols.append(f'C:{length_class}:{suffix_class}')

        stripped = sentence.rstrip()
        if stripped.endswith('?'):
            symbols.append('<END_Q>')
        elif stripped.endswith('!'):
            symbols.append('<END_E>')
        else:
            symbols.append('<END_D>')
    return symbols


def _interpolated_markov_metrics(train: list[str], heldout: list[str]) -> tuple[float, float, float]:
    """Return interpolated bigram entropy, unigram entropy, and seen-transition coverage."""
    if len(train) < 30 or len(heldout) < 12:
        return 0.0, 0.0, 0.0
    unigrams = collections.Counter(train)
    bigrams = collections.Counter(zip(train, train[1:]))
    contexts = collections.Counter(train[:-1])
    vocab_size = max(len(unigrams) + 1, 2)
    total = max(len(train), 1)
    alpha = 0.12
    model_surprisal = []
    unigram_surprisal = []
    seen = 0

    for previous, current in zip(heldout, heldout[1:]):
        p_uni = (unigrams.get(current, 0) + alpha) / (total + alpha * vocab_size)
        context_count = contexts.get(previous, 0)
        transition_count = bigrams.get((previous, current), 0)
        if transition_count:
            seen += 1
        # Context-adaptive interpolation prevents unseen bigrams from receiving
        # an artificial infinite penalty and makes the verifier training-free.
        lam = context_count / (context_count + 7.0) if context_count else 0.0
        p_ml = transition_count / context_count if context_count else 0.0
        p_model = max(lam * p_ml + (1.0 - lam) * p_uni, 1e-12)
        model_surprisal.append(-math.log(p_model, 2))
        unigram_surprisal.append(-math.log(max(p_uni, 1e-12), 2))

    return (
        _safe_mean(model_surprisal),
        _safe_mean(unigram_surprisal),
        seen / max(len(heldout) - 1, 1),
    )


def _cross_validated_markov_verifier(text: str) -> dict:
    """
    Independent, non-neural corroboration layer.

    It performs leave-one-segment-out validation on a topic-reduced structural
    symbol sequence. Each held-out segment is scored by a Markov model built
    only from the remaining segments. A deterministic shuffled control is also
    scored. Real cross-segment predictability must exceed the shuffled control
    before the verifier can raise the final fingerprint.
    """
    symbols = _stylometric_symbol_sequence(text)
    segments = _segment_tokens(symbols, min_size=105, max_segments=8)
    empty = {
        'score': 0.0, 'information_gain': 0.0, 'shuffle_contrast': 0.0,
        'transition_coverage': 0.0, 'surprisal_cv': 1.0,
        'folds': len(segments), 'reliability': 0.0,
    }
    if len(segments) < 3 or len(symbols) < 315:
        return empty

    observed_gains = []
    shuffled_gains = []
    coverages = []
    model_entropies = []
    seed_material = hashlib.sha256('|'.join(symbols[:6000]).encode('utf-8')).hexdigest()

    for idx, heldout in enumerate(segments):
        train = [symbol for j, segment in enumerate(segments) if j != idx for symbol in segment]
        model_h, unigram_h, coverage = _interpolated_markov_metrics(train, heldout)
        if unigram_h <= 1e-9:
            continue
        observed_gain = (unigram_h - model_h) / unigram_h

        shuffled = list(heldout)
        fold_seed = int(hashlib.sha256(f'{seed_material}:{idx}'.encode('utf-8')).hexdigest()[:16], 16)
        random.Random(fold_seed).shuffle(shuffled)
        shuffled_h, shuffled_uni_h, _ = _interpolated_markov_metrics(train, shuffled)
        shuffled_gain = ((shuffled_uni_h - shuffled_h) / shuffled_uni_h) if shuffled_uni_h > 1e-9 else 0.0

        observed_gains.append(observed_gain)
        shuffled_gains.append(shuffled_gain)
        coverages.append(coverage)
        model_entropies.append(model_h)

    if len(observed_gains) < 3:
        return empty

    information_gain = _safe_mean(observed_gains)
    shuffle_contrast = information_gain - _safe_mean(shuffled_gains)
    transition_coverage = _safe_mean(coverages)
    entropy_mean = _safe_mean(model_entropies)
    surprisal_cv = _safe_std(model_entropies) / max(entropy_mean, 1e-9)
    gain_cv = _safe_std(observed_gains) / max(abs(information_gain), 0.035)

    gain_signal = _clamp((information_gain - 0.018) / 0.155)
    contrast_signal = _clamp((shuffle_contrast - 0.020) / 0.180)
    coverage_signal = _clamp((transition_coverage - 0.10) / 0.52)
    entropy_stability = _clamp((0.22 - surprisal_cv) / 0.18)
    gain_stability = _clamp((0.90 - gain_cv) / 0.72)

    score = _clamp(
        gain_signal * 0.28 + contrast_signal * 0.31 + coverage_signal * 0.16 +
        entropy_stability * 0.15 + gain_stability * 0.10
    )
    reliability = _clamp(
        min(1.0, len(symbols) / 1200.0) * 0.58 +
        min(1.0, len(observed_gains) / 6.0) * 0.42
    )
    # Low-data verification is allowed to corroborate weakly, never dominate.
    score *= 0.70 + 0.30 * reliability

    return {
        'score': _clamp(score),
        'information_gain': information_gain,
        'shuffle_contrast': shuffle_contrast,
        'transition_coverage': transition_coverage,
        'surprisal_cv': surprisal_cv,
        'folds': len(observed_gains),
        'reliability': reliability,
    }



def _leave_one_sentence_surprisal_profile(sentences: list[str]) -> dict:
    """
    Measures how smoothly internal predictability changes from sentence to
    sentence while subtracting each sentence from its own reference counts.
    This remains useful after superficial rewriting because it relies on
    distributional movement rather than a fixed phrase list.
    """
    tokenized = [_style_tokens(sentence) for sentence in sentences]
    tokenized = [tokens for tokens in tokenized if len(tokens) >= 5]
    empty = {
        'score': 0.0, 'mean': 0.0, 'cv': 1.0, 'delta_mean': 0.0,
        'delta_cv': 1.0, 'outlier_share': 0.0, 'reliability': 0.0,
    }
    if len(tokenized) < 5:
        return empty

    global_uni = collections.Counter(token for sent in tokenized for token in sent)
    global_bi = collections.Counter(
        pair for sent in tokenized for pair in zip(sent, sent[1:])
    )
    global_ctx = collections.Counter(
        token for sent in tokenized for token in sent[:-1]
    )
    vocab_size = max(len(global_uni) + 1, 2)
    total_tokens = sum(global_uni.values())
    alpha = 0.18
    values = []

    for sent in tokenized:
        local_uni = collections.Counter(sent)
        local_bi = collections.Counter(zip(sent, sent[1:]))
        local_ctx = collections.Counter(sent[:-1])
        train_total = max(total_tokens - len(sent), 1)
        surprisals = []
        previous = '<START>'
        for idx, current in enumerate(sent):
            uni_count = max(0, global_uni.get(current, 0) - local_uni.get(current, 0))
            p_uni = (uni_count + alpha) / (train_total + alpha * vocab_size)
            if idx == 0:
                p = p_uni
            else:
                context_count = max(0, global_ctx.get(previous, 0) - local_ctx.get(previous, 0))
                transition_count = max(
                    0,
                    global_bi.get((previous, current), 0) - local_bi.get((previous, current), 0),
                )
                lam = context_count / (context_count + 6.0) if context_count else 0.0
                p_ml = transition_count / context_count if context_count else 0.0
                p = lam * p_ml + (1.0 - lam) * p_uni
            surprisals.append(-math.log(max(p, 1e-12), 2))
            previous = current
        values.append(_safe_mean(surprisals))

    mean_value = _safe_mean(values)
    cv = _safe_std(values) / max(mean_value, 1e-9)
    deltas = [abs(values[i] - values[i - 1]) for i in range(1, len(values))]
    delta_mean = _safe_mean(deltas)
    delta_cv = _safe_std(deltas) / max(delta_mean, 0.08)
    med = _percentile(values, 0.50)
    mad = _median_absolute_deviation(values)
    outlier_share = sum(abs(value - med) > max(2.8 * mad, 0.75) for value in values) / len(values)

    smooth_level = _clamp((0.24 - cv) / 0.18)
    smooth_changes = _clamp((0.78 - delta_cv) / 0.62)
    low_jumps = _clamp((0.72 - delta_mean) / 0.58)
    low_outliers = _clamp((0.18 - outlier_share) / 0.16)
    reliability = _clamp(min(1.0, len(values) / 14.0) * min(1.0, sum(map(len, tokenized)) / 420.0))
    score = _clamp(
        smooth_level * 0.38 + smooth_changes * 0.24 +
        low_jumps * 0.23 + low_outliers * 0.15
    ) * (0.62 + 0.38 * reliability)
    return {
        'score': _clamp(score), 'mean': mean_value, 'cv': cv,
        'delta_mean': delta_mean, 'delta_cv': delta_cv,
        'outlier_share': outlier_share, 'reliability': reliability,
    }


def _symbol_view_signature(sequence: list[str], mode: str) -> list[str]:
    view = []
    for symbol in sequence:
        if symbol.startswith('F:'):
            if mode == 'fine':
                view.append(symbol)
            elif mode == 'coarse':
                view.append('F')
            else:
                view.append('W')
        elif symbol.startswith('C:'):
            parts = symbol.split(':')
            if mode == 'fine':
                view.append(symbol)
            elif mode == 'coarse':
                view.append('C:' + (parts[1] if len(parts) > 1 else 'M'))
            else:
                view.append('W')
        else:
            view.append(symbol)
    return view


def _sequence_recurrence_metrics(sequence: list[str], baseline: float) -> dict:
    if len(sequence) < 80:
        return {'score': 0.0, 'repeat': 0.0, 'entropy': 1.0, 'top_share': 0.0}
    repeats = []
    entropies = []
    top_shares = []
    for n in (3, 4, 5):
        grams = [tuple(sequence[i:i + n]) for i in range(len(sequence) - n + 1)]
        counts = collections.Counter(grams)
        total = max(len(grams), 1)
        repeats.append(sum(max(0, count - 1) for count in counts.values()) / total)
        entropies.append(_normalized_entropy(counts))
        top_shares.append(sum(count for _, count in counts.most_common(8)) / total)
    repeat = _safe_mean(repeats)
    entropy = _safe_mean(entropies)
    top_share = _safe_mean(top_shares)
    recurrence = _clamp((repeat - baseline) / max(0.18, 0.52 - baseline))
    concentration = _clamp((top_share - 0.055) / 0.24)
    low_entropy = _clamp((0.965 - entropy) / 0.16)
    score = _clamp(recurrence * 0.52 + concentration * 0.28 + low_entropy * 0.20)
    return {'score': score, 'repeat': repeat, 'entropy': entropy, 'top_share': top_share}


def _paraphrase_invariant_structure_profile(text: str) -> dict:
    """
    Uses several progressively coarser structural views. Word substitutions
    and many humanizer edits alter the fine view but leave coarse sentence and
    function/content patterns comparatively stable.
    """
    symbols = _stylometric_symbol_sequence(text)
    empty = {
        'score': 0.0, 'agreement': 0.0, 'fine': 0.0, 'coarse': 0.0,
        'shape': 0.0, 'reliability': 0.0,
    }
    if len(symbols) < 150:
        return empty
    fine = _sequence_recurrence_metrics(_symbol_view_signature(symbols, 'fine'), 0.020)
    coarse = _sequence_recurrence_metrics(_symbol_view_signature(symbols, 'coarse'), 0.090)
    shape = _sequence_recurrence_metrics(_symbol_view_signature(symbols, 'shape'), 0.180)
    scores = [fine['score'], coarse['score'], shape['score']]
    agreement = 1.0 - _clamp(_safe_std(scores) / 0.42)
    floor_support = _percentile(scores, 0.35)
    mean_support = _safe_mean(scores)
    reliability = _clamp(min(1.0, len(symbols) / 1500.0))
    score = _clamp(mean_support * 0.62 + floor_support * 0.23 + agreement * mean_support * 0.15)
    score *= 0.66 + 0.34 * reliability
    return {
        'score': _clamp(score), 'agreement': agreement,
        'fine': fine['score'], 'coarse': coarse['score'],
        'shape': shape['score'], 'reliability': reliability,
    }


def _style_trajectory_profile(chunks: list[str]) -> dict:
    """Detects an overly smooth style trajectory and isolated rewrite spikes."""
    empty = {
        'score': 0.0, 'smoothness': 0.0, 'rewrite_spike': 0.0,
        'adjacent_distance': 0.0, 'distance_cv': 1.0, 'reliability': 0.0,
    }
    if len(chunks) < 3:
        return empty
    vectors = [_chunk_vector(chunk) for chunk in chunks]
    distances = [
        max(0.0, 1.0 - _cosine_similarity(vectors[i - 1], vectors[i]))
        for i in range(1, len(vectors))
    ]
    mean_distance = _safe_mean(distances)
    distance_cv = _safe_std(distances) / max(mean_distance, 0.015)
    med = _percentile(distances, 0.50)
    max_distance = max(distances) if distances else 0.0
    spike_ratio = max_distance / max(med, 0.012)

    smooth_level = _clamp((0.095 - mean_distance) / 0.075)
    smooth_variation = _clamp((1.15 - distance_cv) / 0.90)
    smoothness = _clamp(smooth_level * 0.62 + smooth_variation * 0.38)
    rewrite_spike = _clamp((spike_ratio - 2.2) / 4.2) * _clamp((0.085 - med) / 0.070)
    reliability = _clamp(min(1.0, len(chunks) / 7.0))
    score = _clamp(smoothness * 0.70 + rewrite_spike * 0.30) * (0.58 + 0.42 * reliability)
    return {
        'score': _clamp(score), 'smoothness': smoothness,
        'rewrite_spike': rewrite_spike, 'adjacent_distance': mean_distance,
        'distance_cv': distance_cv, 'reliability': reliability,
    }


_ACADEMIC_TRANSITIONS = {
    'additionally', 'consequently', 'furthermore', 'however', 'moreover',
    'nevertheless', 'nonetheless', 'therefore', 'thus', 'overall', 'notably',
    'specifically', 'importantly', 'accordingly', 'subsequently', 'similarly',
}

_ACADEMIC_FORMULA_PATTERNS = [
    re.compile(pattern, re.I) for pattern in (
        r'\bit is (?:important|essential|crucial) to (?:note|recognize|acknowledge)\b',
        r'\bplays? a (?:vital|crucial|significant|pivotal) role\b',
        r'\bsheds? light on\b',
        r'\bunderscores? the (?:importance|need|significance)\b',
        r'\bhighlights? the (?:importance|need|significance)\b',
        r'\bthe findings (?:indicate|suggest|demonstrate|reveal)\b',
        r'\bthis (?:study|paper|research|analysis) (?:aims|seeks|examines|explores|investigates)\b',
        r'\ba comprehensive (?:understanding|analysis|framework|approach|overview)\b',
        r'\bin today[’\']s rapidly evolving\b',
        r'\bmultifaceted (?:nature|approach|challenge|framework)\b',
        r'\bsignificant implications for\b',
        r'\bfuture research should\b',
        r'\bit can be concluded that\b',
    )
]


def _academic_formulaicity_profile(text: str, sentences: list[str], tokens: list[str]) -> dict:
    """Weak corroborating signal for repeated academic boilerplate."""
    n_words = len(tokens)
    empty = {'score': 0.0, 'transition_density': 0.0, 'formula_hits': 0, 'reliability': 0.0}
    if n_words < 90 or len(sentences) < 5:
        return empty
    lower_sentences = [sentence.lower() for sentence in sentences]
    transition_hits = []
    for sentence in lower_sentences:
        sent_tokens = _style_tokens(sentence)
        transition_hits.append(sum(token in _ACADEMIC_TRANSITIONS for token in sent_tokens))
    transition_total = sum(transition_hits)
    transition_density = transition_total * 100.0 / max(n_words, 1)
    transition_sentence_share = sum(hit > 0 for hit in transition_hits) / len(transition_hits)
    repeated_transition_share = sum(hit >= 2 for hit in transition_hits) / len(transition_hits)
    formula_hits = sum(len(pattern.findall(text)) for pattern in _ACADEMIC_FORMULA_PATTERNS)
    formula_density = formula_hits * 200.0 / max(n_words, 1)

    score = _clamp(
        _clamp((transition_density - 0.70) / 2.20) * 0.30 +
        _clamp((transition_sentence_share - 0.22) / 0.48) * 0.24 +
        _clamp((repeated_transition_share - 0.025) / 0.20) * 0.16 +
        _clamp((formula_density - 0.18) / 1.30) * 0.30
    )
    reliability = _clamp(min(1.0, n_words / 500.0) * min(1.0, len(sentences) / 18.0))
    score *= 0.48 + 0.52 * reliability
    return {
        'score': _clamp(score), 'transition_density': transition_density,
        'formula_hits': formula_hits, 'reliability': reliability,
    }


def _anti_evasion_consensus(profiles: dict) -> dict:
    values = [
        profiles['diversity']['score'], profiles['invariance']['score'],
        profiles['trajectory']['score'], profiles['formulaicity']['score'],
    ]
    ordered = sorted(values)
    median = _percentile(values, 0.50)
    agreement = 1.0 - _clamp(_safe_std(values) / 0.44)
    support = sum(value >= 0.48 for value in values)
    strong = sum(value >= 0.66 for value in values)
    # Boilerplate is never allowed to create consensus by itself.
    non_formula_support = sum(value >= 0.48 for value in values[:3])
    score = _clamp(_safe_mean(ordered[1:]) * 0.72 + median * agreement * 0.28)
    if non_formula_support < 2:
        score = min(score, 0.42)
    return {
        'score': score, 'agreement': agreement, 'support': support,
        'strong': strong, 'non_formula_support': non_formula_support,
    }

def _compute_stylometric_features(text: str) -> dict:
    normalized = _normalize_style_text(text)
    tokens = _style_tokens(normalized)
    sentences = _split_analysis_sentences(normalized)
    paragraphs = _split_analysis_paragraphs(text)
    chunks = _build_analysis_chunks(sentences)

    n_words = len(tokens)
    n_sents = len(sentences)
    rhythm = _sentence_rhythm_profile(sentences)
    para_cv, para_score = _paragraph_uniformity(paragraphs)
    lexical = _lexical_profile(tokens)
    lexical_mattr_cv, lexical_entropy_cv, lexical_stability = _lexical_stability(tokens)
    word_repeat_ratio, word_repeat_score, max_ngram_repeat = _word_ngram_repetition(tokens)
    char_repeat_ratio, char_entropy, char_repeat_score = _character_ngram_profile(normalized)
    opener_ratio, opener_score = _opener_repetition(sentences)
    template_ratio, template_score = _sentence_template_repetition(sentences)
    function_js, function_stability = _function_word_stability(tokens)
    punct_js, punctuation_stability, punctuation_types = _punctuation_stability(sentences)
    compression_ratio, compression_score = _compression_profile(normalized)
    chunk_similarity, chunk_consistency = _chunk_style_consistency(chunks)
    markov_verifier = _cross_validated_markov_verifier(normalized)

    diversity_profile = _leave_one_sentence_surprisal_profile(sentences)
    invariance_profile = _paraphrase_invariant_structure_profile(normalized)
    trajectory_profile = _style_trajectory_profile(chunks)
    formulaicity_profile = _academic_formulaicity_profile(normalized, sentences, tokens)
    anti_evasion = _anti_evasion_consensus({
        'diversity': diversity_profile,
        'invariance': invariance_profile,
        'trajectory': trajectory_profile,
        'formulaicity': formulaicity_profile,
    })

    readability_texts = chunks if len(chunks) >= 2 else paragraphs
    flesch_values, word_length_values = _readability_values(readability_texts)
    flesch_mean = _safe_mean(flesch_values)
    flesch_cv = _safe_std(flesch_values) / max(abs(flesch_mean), 20.0) if flesch_values else 0.0
    word_length_cv = _safe_std(word_length_values) / max(_safe_mean(word_length_values), 1e-9) if word_length_values else 0.0
    readability_stability = _clamp((0.24 - flesch_cv) / 0.20) * 0.62 + _clamp((0.07 - word_length_cv) / 0.06) * 0.38

    repetition_score = _clamp(
        word_repeat_score * 0.34 + char_repeat_score * 0.22 +
        opener_score * 0.19 + template_score * 0.25
    )
    rhythm_score = _clamp(rhythm['score'] * 0.72 + para_score * 0.28)
    lexical_score = _clamp(
        lexical_stability * 0.42 +
        _clamp((0.58 - lexical['hapax_ratio']) / 0.34) * 0.18 +
        _clamp((lexical['simpson'] - 0.012) / 0.045) * 0.12 +
        _clamp((0.10 - lexical['mattr_cv']) / 0.085) * 0.18 +
        _clamp((lexical['zipf_r2'] - 0.86) / 0.13) * 0.10
    )
    distribution_score = _clamp(
        function_stability * 0.34 + punctuation_stability * 0.22 +
        readability_stability * 0.20 + chunk_consistency * 0.24
    )
    predictability_score = _clamp(
        compression_score * 0.42 + char_repeat_score * 0.22 +
        _clamp((0.94 - char_entropy) / 0.12) * 0.16 +
        _clamp((lexical['simpson'] - 0.010) / 0.055) * 0.20
    )

    verification_score = markov_verifier['score']
    diversity_score = diversity_profile['score']
    invariance_score = invariance_profile['score']
    trajectory_score = trajectory_profile['score']
    formulaicity_score = formulaicity_profile['score']
    anti_evasion_score = anti_evasion['score']

    head_scores = {
        'rhythm': rhythm_score,
        'repetition': repetition_score,
        'lexical': lexical_score,
        'distribution': distribution_score,
        'predictability': predictability_score,
        'verification': verification_score,
        'diversity': diversity_score,
        'invariance': invariance_score,
        'trajectory': trajectory_score,
        'formulaicity': formulaicity_score,
    }
    support_count = sum(v >= 0.52 for v in head_scores.values())
    strong_support_count = sum(v >= 0.68 for v in head_scores.values())
    evidence_families = sum(v >= 0.45 for v in head_scores.values())

    human_complexity = _clamp(
        (0.032 if punctuation_types >= 9 else 0.0) +
        (0.028 if rhythm['outlier_share'] >= 0.14 else 0.0) +
        (0.028 if lexical_mattr_cv >= 0.15 else 0.0) +
        (0.022 if lexical['hapax_ratio'] >= 0.68 else 0.0) +
        (0.020 if diversity_profile['outlier_share'] >= 0.24 else 0.0),
        0.0, 0.11
    )

    return {
        'normalized_text': normalized, 'tokens': tokens, 'sentences': sentences,
        'paragraphs': paragraphs, 'chunks': chunks,
        'n_words': n_words, 'n_sents': n_sents,
        'avg_sent_len': rhythm['mean'], 'sent_len_cv': rhythm['cv'],
        'sentence_mad_ratio': rhythm['mad_ratio'], 'sentence_qspread': rhythm['qspread'],
        'sentence_autocorr': rhythm['autocorr'], 'sentence_outlier_share': rhythm['outlier_share'],
        'paragraph_len_cv': para_cv,
        'word_ngram_repeat_ratio': word_repeat_ratio, 'max_ngram_repeat': max_ngram_repeat,
        'char_repeat_ratio': char_repeat_ratio, 'char_entropy': char_entropy,
        'opener_repeat_ratio': opener_ratio, 'template_repeat_ratio': template_ratio,
        'function_word_js': function_js, 'punctuation_js': punct_js,
        'compression_ratio': compression_ratio, 'chunk_similarity': chunk_similarity,
        'mattr40': lexical['mattr40'], 'mattr80': lexical['mattr80'],
        'mattr_cv': lexical['mattr_cv'], 'lexical_mattr_segment_cv': lexical_mattr_cv,
        'lexical_entropy_segment_cv': lexical_entropy_cv,
        'hapax_ratio': lexical['hapax_ratio'], 'simpson': lexical['simpson'],
        'lexical_entropy': lexical['entropy'], 'yule_k': lexical['yule_k'],
        'zipf_r2': lexical['zipf_r2'], 'zipf_slope': lexical['zipf_slope'],
        'flesch_mean': flesch_mean, 'flesch_cv': flesch_cv,
        'rhythm_score': rhythm_score, 'repetition_score': repetition_score,
        'lexical_score': lexical_score, 'distribution_score': distribution_score,
        'predictability_score': predictability_score,
        'verification_score': verification_score,
        'diversity_score': diversity_score,
        'invariance_score': invariance_score,
        'trajectory_score': trajectory_score,
        'formulaicity_score': formulaicity_score,
        'anti_evasion_score': anti_evasion_score,
        'anti_evasion_agreement': anti_evasion['agreement'],
        'anti_evasion_support': anti_evasion['support'],
        'anti_evasion_strong': anti_evasion['strong'],
        'anti_evasion_non_formula_support': anti_evasion['non_formula_support'],
        'surprisal_mean': diversity_profile['mean'],
        'surprisal_cv': diversity_profile['cv'],
        'surprisal_delta_mean': diversity_profile['delta_mean'],
        'surprisal_delta_cv': diversity_profile['delta_cv'],
        'surprisal_outlier_share': diversity_profile['outlier_share'],
        'surprisal_reliability': diversity_profile['reliability'],
        'invariance_agreement': invariance_profile['agreement'],
        'invariance_fine': invariance_profile['fine'],
        'invariance_coarse': invariance_profile['coarse'],
        'invariance_shape': invariance_profile['shape'],
        'invariance_reliability': invariance_profile['reliability'],
        'trajectory_smoothness': trajectory_profile['smoothness'],
        'trajectory_rewrite_spike': trajectory_profile['rewrite_spike'],
        'trajectory_adjacent_distance': trajectory_profile['adjacent_distance'],
        'trajectory_distance_cv': trajectory_profile['distance_cv'],
        'trajectory_reliability': trajectory_profile['reliability'],
        'academic_transition_density': formulaicity_profile['transition_density'],
        'academic_formula_hits': formulaicity_profile['formula_hits'],
        'academic_formula_reliability': formulaicity_profile['reliability'],
        'markov_information_gain': markov_verifier['information_gain'],
        'markov_shuffle_contrast': markov_verifier['shuffle_contrast'],
        'markov_transition_coverage': markov_verifier['transition_coverage'],
        'markov_surprisal_cv': markov_verifier['surprisal_cv'],
        'markov_folds': markov_verifier['folds'],
        'markov_reliability': markov_verifier['reliability'],
        'style_score': _safe_mean(head_scores.values()),
        'support_count': support_count, 'strong_support_count': strong_support_count,
        'evidence_families': evidence_families, 'human_complexity': human_complexity,
        'head_scores': head_scores,
        'top_kw': collections.Counter(t for t in tokens if t not in _ENGLISH_FUNCTION_WORDS and len(t) >= 5).most_common(1)[0][0]
                  if any(t not in _ENGLISH_FUNCTION_WORDS and len(t) >= 5 for t in tokens) else '',
        'top_kw_count': collections.Counter(t for t in tokens if t not in _ENGLISH_FUNCTION_WORDS and len(t) >= 5).most_common(1)[0][1]
                        if any(t not in _ENGLISH_FUNCTION_WORDS and len(t) >= 5 for t in tokens) else 0,
    }

def _score_feature_probability(features: dict, local: bool = False) -> tuple[float, dict]:
    head_probabilities = {
        'rhythm': _sigmoid(8.0 * (features['rhythm_score'] - 0.57)),
        'repetition': _sigmoid(8.6 * (features['repetition_score'] - 0.52)),
        'lexical': _sigmoid(7.6 * (features['lexical_score'] - 0.47)),
        'distribution': _sigmoid(8.0 * (features['distribution_score'] - 0.48)),
        'predictability': _sigmoid(8.4 * (features['predictability_score'] - 0.57)),
        'verification': _sigmoid(8.2 * (features['verification_score'] - 0.52)),
        'diversity': _sigmoid(8.0 * (features['diversity_score'] - 0.48)),
        'invariance': _sigmoid(8.3 * (features['invariance_score'] - 0.50)),
        'trajectory': _sigmoid(7.6 * (features['trajectory_score'] - 0.55)),
        'formulaicity': _sigmoid(7.6 * (features['formulaicity_score'] - 0.36)),
    }
    probs = list(head_probabilities.values())
    weights = {
        'rhythm': 0.13, 'repetition': 0.14, 'lexical': 0.11,
        'distribution': 0.13, 'predictability': 0.09, 'verification': 0.12,
        'diversity': 0.09, 'invariance': 0.10, 'trajectory': 0.05,
        'formulaicity': 0.04,
    }
    weighted_raw = sum(head_probabilities[k] * weights[k] for k in weights)
    ordered_probs = sorted(probs)
    trimmed_consensus = _safe_mean(ordered_probs[2:-2]) if len(ordered_probs) >= 7 else _safe_mean(ordered_probs)
    raw = weighted_raw * 0.87 + trimmed_consensus * 0.13
    agreement = 1.0 - _clamp(_safe_std(probs) / 0.43)
    median_head = _percentile(probs, 0.50)
    raw += median_head * agreement * 0.040
    raw -= features['human_complexity']

    verifier_probability = head_probabilities['verification']
    other_support = sum(
        probability >= 0.50 for name, probability in head_probabilities.items()
        if name != 'verification'
    )
    core_support = sum(
        head_probabilities[name] >= 0.50
        for name in ('rhythm', 'repetition', 'lexical', 'distribution', 'predictability')
    )
    anti_support = features.get('anti_evasion_non_formula_support', 0)
    anti_probability = _sigmoid(8.0 * (features.get('anti_evasion_score', 0.0) - 0.52))

    if verifier_probability >= 0.56 and other_support >= 3:
        raw += min(0.045, 0.010 + (verifier_probability - 0.56) * 0.10)
    elif verifier_probability < 0.24 and core_support <= 2:
        raw -= 0.020

    # Resistance boost requires agreement between at least two paraphrase-
    # resilient families and conventional evidence. Formulaic wording alone
    # cannot trigger it.
    if ADVERSARIAL_RESISTANCE_POLICY and anti_support >= 2 and core_support >= 2:
        raw += min(0.080, max(0.0, anti_probability - 0.52) * 0.15 + 0.012)
    elif features.get('anti_evasion_agreement', 0.0) < 0.30 and core_support <= 2:
        raw -= 0.018

    # Complementary bundles target two common evasion paths. The first catches
    # polished academic generation whose formulaic progression survives surface
    # edits. The second catches humanized text whose vocabulary was diversified
    # while sentence-level predictability remains unusually even. Neither bundle
    # can activate from a single signal.
    formulaic_bundle = min(
        head_probabilities['rhythm'],
        head_probabilities['distribution'],
        head_probabilities['formulaicity'],
    )
    humanizer_bundle = min(
        head_probabilities['lexical'],
        head_probabilities['diversity'],
    )
    if formulaic_bundle >= 0.54:
        raw += min(0.135, 0.085 + (formulaic_bundle - 0.54) * 0.22)
    if (
        humanizer_bundle >= 0.50 and
        head_probabilities['diversity'] >= 0.62 and
        features.get('n_words', 0) >= 140
    ):
        raw += min(0.090, 0.050 + (humanizer_bundle - 0.50) * 0.18)

    if (
        features['repetition_score'] < 0.15 and
        features['predictability_score'] < 0.15 and
        features['verification_score'] < 0.22 and
        features['invariance_score'] < 0.24 and
        features['distribution_score'] < 0.42
    ):
        raw *= 0.68

    support = sum(p >= 0.55 for p in probs)
    strong = sum(p >= 0.72 for p in probs)
    if support <= 1:
        raw = min(raw, 0.22 if local else 0.24)
    elif support == 2:
        raw = min(raw, 0.36 if local else 0.38)
    elif support == 3:
        raw = min(raw, 0.52 if local else 0.56)
    elif support == 4 and strong == 0:
        raw = min(raw, 0.62)
    elif features['evidence_families'] <= 3:
        raw = min(raw, 0.50)

    if STRICT_FINGERPRINT_POLICY:
        if agreement < 0.36:
            raw *= 0.86
        if verifier_probability < 0.22 and core_support < 3 and anti_support < 2:
            raw = min(raw, 0.43 if not local else 0.37)
        if head_probabilities['formulaicity'] >= 0.70 and core_support < 2 and anti_support < 2:
            raw = min(raw, 0.34)

    n_words = features['n_words']
    if n_words < 35:
        raw = min(raw * 0.70, 0.38)
    elif n_words < 70:
        raw = min(raw * 0.84, 0.55)
    elif n_words < 120 and not local:
        raw = min(raw * 0.92, 0.72)

    return _clamp(raw), {
        'head_probabilities': head_probabilities,
        'agreement': agreement,
        'support': support,
        'strong': strong,
        'other_support': other_support,
        'core_support': core_support,
        'anti_support': anti_support,
        'anti_probability': anti_probability,
        'verifier_probability': verifier_probability,
    }

def _bootstrap_chunk_score(chunk_scores: list[float], iterations: int = 240) -> dict:
    if not chunk_scores:
        return {'median': 0.0, 'low': 0.0, 'high': 0.0, 'width': 1.0}
    if len(chunk_scores) == 1:
        value = chunk_scores[0]
        return {'median': value, 'low': max(0.0, value - 0.18), 'high': min(1.0, value + 0.18), 'width': 0.36}
    seed_material = '|'.join(f'{s:.6f}' for s in chunk_scores).encode('utf-8')
    seed = int(hashlib.sha256(seed_material).hexdigest()[:16], 16)
    rng = random.Random(seed)
    means = []
    n = len(chunk_scores)
    for _ in range(iterations):
        sample = [chunk_scores[rng.randrange(n)] for _ in range(n)]
        means.append(_safe_mean(sample))
    low = _percentile(means, 0.05)
    high = _percentile(means, 0.95)
    return {'median': _percentile(means, 0.50), 'low': low, 'high': high, 'width': high - low}



def _adaptive_document_window_scales(n_words: int) -> list[tuple[int, int, int]]:
    """
    Build evidence windows that are deliberately much larger than the passages
    used for highlighting. Short passages are useful for location, but several
    structural families are statistically underpowered below a few hundred
    words. The scales therefore adapt to the document length.
    """
    n = max(int(n_words or 0), 1)
    scales = []
    if n >= 700:
        scales.append((min(800, max(520, n // 8)), 320, 1150))
    if n >= 1800:
        target = min(1700, max(900, n // 9))
        scales.append((target, max(560, int(target * 0.62)), int(target * 1.38)))
    if n >= 5000:
        target = min(3200, max(1700, n // 6))
        scales.append((target, max(1050, int(target * 0.65)), int(target * 1.32)))
    # Deduplicate close scales while keeping ascending order.
    out = []
    for spec in sorted(scales):
        if not out or abs(spec[0] - out[-1][0]) >= 280:
            out.append(spec)
    return out


def _robust_weighted_mean(values: list[float], weights: list[float]) -> float:
    pairs = [(float(v), max(0.0, float(w))) for v, w in zip(values, weights)]
    pairs = [(v, w) for v, w in pairs if w > 0]
    if not pairs:
        return _safe_mean(values)
    pairs.sort(key=lambda item: item[0])
    # Winsorize only extreme tails. This prevents one unusual section from
    # dominating a long research paper while retaining mixed-authorship clues.
    vals = [p[0] for p in pairs]
    lo = _percentile(vals, 0.08)
    hi = _percentile(vals, 0.92)
    return sum(max(lo, min(hi, v)) * w for v, w in pairs) / max(sum(w for _, w in pairs), 1e-9)


def _score_evidence_window(window_text: str) -> dict:
    features = _compute_stylometric_features(window_text)
    # At evidence-window sizes, use the document scoring path rather than the
    # short-passage caps. Reliability below prevents weak windows dominating.
    probability, meta = _score_feature_probability(features, local=False)
    n_words = max(features.get('n_words', 0), 1)
    evidence = features.get('evidence_families', 0)
    support = meta.get('support', 0)
    agreement = meta.get('agreement', 0.0)
    length_reliability = _clamp((n_words - 180.0) / 950.0, 0.18, 1.0)
    evidence_reliability = _clamp((evidence - 1.0) / 7.0, 0.16, 1.0)
    support_reliability = _clamp((support - 1.0) / 6.0, 0.18, 1.0)
    reliability = _clamp(
        0.38 * length_reliability + 0.27 * evidence_reliability +
        0.20 * support_reliability + 0.15 * agreement,
        0.12, 1.0
    )
    return {
        'score': probability,
        'reliability': reliability,
        'n_words': n_words,
        'support': support,
        'strong': meta.get('strong', 0),
        'evidence_families': evidence,
        'agreement': agreement,
    }


def _hierarchical_document_profile(analyzed_text: str, sentences: list[str],
                                   global_score: float, global_meta: dict,
                                   global_features: dict) -> dict:
    """
    Hierarchical evidence aggregation for long academic documents.

    The earlier engine used ~190-word windows both for locating passages and
    for estimating the whole-document percentage. Many independent structural
    tests cannot become reliable in such short windows, causing severe downward
    bias on long papers even when document-level evidence is strong. This layer
    keeps fine windows for highlighting, but estimates authorship evidence from
    adaptive medium and long windows and applies reliability-aware shrinkage.
    """
    empty = {
        'score': global_score, 'reliability': 0.0, 'scale_count': 0,
        'window_count': 0, 'cross_scale_agreement': 0.0,
        'scale_scores': [], 'scale_details': [], 'strong_window_share': 0.0,
    }
    if not HIERARCHICAL_ACADEMIC_POLICY:
        return empty
    n_words = max(global_features.get('n_words', 0), 1)
    scales = _adaptive_document_window_scales(n_words)
    if not scales:
        return empty

    scale_details = []
    all_scores = []
    all_reliabilities = []
    all_strong_flags = []

    for target, min_words, max_words in scales:
        windows = _build_analysis_chunks(
            sentences, target_words=target, min_words=min_words,
            max_words=max_words, overlap_sentences=1
        )
        records = [_score_evidence_window(window) for window in windows]
        if not records:
            continue
        scores = [r['score'] for r in records]
        reliabilities = [r['reliability'] for r in records]
        weighted = _robust_weighted_mean(scores, reliabilities)
        median = _percentile(scores, 0.50)
        q75 = _percentile(scores, 0.75)
        q90 = _percentile(scores, 0.90)
        robust_score = _clamp(
            0.36 * weighted + 0.24 * median + 0.25 * q75 + 0.15 * q90
        )
        coverage = sum(s >= 0.42 for s in scores) / max(len(scores), 1)
        strong_share = sum(s >= 0.50 for s in scores) / max(len(scores), 1)
        scale_reliability = _clamp(
            _safe_mean(reliabilities) * (0.72 + 0.28 * min(1.0, len(scores) / 5.0))
        )
        scale_details.append({
            'target_words': target, 'window_count': len(scores),
            'score': robust_score, 'weighted_mean': weighted,
            'median': median, 'q75': q75, 'q90': q90,
            'coverage': coverage, 'strong_share': strong_share,
            'reliability': scale_reliability,
        })
        all_scores.extend(scores)
        all_reliabilities.extend(reliabilities)
        all_strong_flags.extend(s >= 0.50 for s in scores)

    if not scale_details:
        return empty

    scale_scores = [d['score'] for d in scale_details]
    scale_weights = [d['reliability'] * (0.85 + 0.15 * i) for i, d in enumerate(scale_details)]
    scale_consensus = _robust_weighted_mean(scale_scores, scale_weights)
    scale_dispersion = _safe_std(scale_scores)
    cross_scale_agreement = _clamp(1.0 - scale_dispersion / 0.23)

    # Long academic papers should not have strong whole-document evidence
    # halved merely because fine windows are underpowered. The hierarchy still
    # requires broad independent support and scales its influence by agreement.
    evidence_gate = _clamp((global_features.get('evidence_families', 0) - 3.0) / 5.0)
    support_gate = _clamp((global_meta.get('support', 0) - 3.0) / 5.0)
    verifier_gate = _clamp((global_meta.get('verifier_probability', 0.0) - 0.38) / 0.52)
    reliability = _clamp(
        0.30 * _safe_mean([d['reliability'] for d in scale_details]) +
        0.22 * cross_scale_agreement + 0.20 * evidence_gate +
        0.16 * support_gate + 0.12 * verifier_gate
    )

    long_scale = scale_details[-1]['score']
    upper_scale = _percentile(scale_scores, 0.75)
    hierarchical_score = _clamp(
        0.43 * global_score + 0.25 * scale_consensus +
        0.22 * long_scale + 0.10 * upper_scale
    )

    # Conservative floor is activated only on long documents with unusually
    # broad corroboration. It corrects statistical downward bias; it is not a
    # blanket percentage increase and cannot activate from wording alone.
    if (
        n_words >= 2500 and
        global_score >= 0.54 and
        global_features.get('evidence_families', 0) >= 7 and
        global_meta.get('support', 0) >= 6 and
        global_meta.get('verifier_probability', 0.0) >= 0.76
    ):
        corroborated_floor = global_score * (0.72 + 0.06 * cross_scale_agreement)
        hierarchical_score = max(hierarchical_score, corroborated_floor)

    return {
        'score': _clamp(hierarchical_score),
        'reliability': reliability,
        'scale_count': len(scale_details),
        'window_count': len(all_scores),
        'cross_scale_agreement': cross_scale_agreement,
        'scale_scores': scale_scores,
        'scale_details': scale_details,
        'strong_window_share': sum(all_strong_flags) / max(len(all_strong_flags), 1),
    }

def _ai_fingerprint_score(text: str) -> dict:
    prep = _prepare_analysis_text(text)
    analyzed_text = prep['analysis_text']
    language = _validate_english_text(analyzed_text)
    features = _compute_stylometric_features(analyzed_text)
    global_score, global_meta = _score_feature_probability(features)
    hierarchical = _hierarchical_document_profile(
        analyzed_text, features['sentences'], global_score, global_meta, features
    )

    chunk_scores = []
    chunk_words = []
    for chunk in features['chunks']:
        cf = _compute_stylometric_features(chunk)
        score, _ = _score_feature_probability(cf, local=True)
        chunk_scores.append(score)
        chunk_words.append(max(cf['n_words'], 1))

    if len(chunk_scores) >= 2:
        weighted_mean = sum(s * w for s, w in zip(chunk_scores, chunk_words)) / max(sum(chunk_words), 1)
        median_score = _percentile(chunk_scores, 0.50)
        q75_score = _percentile(chunk_scores, 0.75)
        q90_score = _percentile(chunk_scores, 0.90)
        dispersion = _safe_std(chunk_scores)
        high_share = sum(s >= 0.62 for s in chunk_scores) / len(chunk_scores)
        medium_share = sum(s >= 0.42 for s in chunk_scores) / len(chunk_scores)
        bootstrap = _bootstrap_chunk_score(chunk_scores)
        final = (
            global_score * 0.36 + weighted_mean * 0.24 + median_score * 0.14 +
            q75_score * 0.12 + q90_score * 0.05 + bootstrap['median'] * 0.09
        )
        if high_share >= 0.50 and global_meta['support'] >= 3:
            final += min(0.06, (high_share - 0.50) * 0.18)
        if dispersion > 0.27 and high_share < 0.35:
            final -= min(0.05, (dispersion - 0.27) * 0.20)
    else:
        weighted_mean = median_score = q75_score = q90_score = global_score
        dispersion = 0.0
        high_share = float(global_score >= 0.62)
        medium_share = float(global_score >= 0.42)
        bootstrap = _bootstrap_chunk_score([global_score])
        final = global_score

    if features['evidence_families'] <= 2:
        final = min(final, 0.28)
    elif features['evidence_families'] == 3:
        final = min(final, 0.44)
    elif global_meta['support'] < 4:
        final = min(final, 0.58)
    if global_meta['strong'] < 2:
        final = min(final, 0.74)

    if (
        ADVERSARIAL_RESISTANCE_POLICY and
        features.get('anti_evasion_non_formula_support', 0) >= 2 and
        global_meta.get('core_support', 0) >= 2 and
        features.get('anti_evasion_score', 0.0) >= 0.58
    ):
        final += min(0.075, 0.018 + (features['anti_evasion_score'] - 0.58) * 0.16)

    if STRICT_FINGERPRINT_POLICY:
        if bootstrap['width'] > 0.30:
            final -= min(0.06, (bootstrap['width'] - 0.30) * 0.25)
        if len(chunk_scores) >= 3 and high_share < 0.15:
            final = min(final, 0.50)
        if global_meta['agreement'] < 0.35:
            final = min(final, 0.55)
        if global_meta.get('verifier_probability', 0.0) < 0.25 and global_meta.get('support', 0) < 4:
            final = min(final, 0.52)

    legacy_final = _clamp(final)
    if HIERARCHICAL_ACADEMIC_POLICY and hierarchical.get('scale_count', 0) > 0:
        # The hierarchy is influential only when the whole document already has
        # unusually broad independent corroboration. This prevents long, formal
        # human papers from being inflated merely because their prose is stable.
        evidence_gate = _clamp((features.get('evidence_families', 0) - 6.5) / 1.5)
        support_gate = _clamp((global_meta.get('support', 0) - 5.0) / 3.0)
        global_gate = _clamp((global_score - 0.48) / 0.14)
        hierarchy_gate = evidence_gate * support_gate * global_gate
        hierarchy_weight = _clamp(
            0.08 + 0.64 * hierarchical.get('reliability', 0.0) * hierarchy_gate,
            0.08, 0.70
        )
        hierarchical_blend = (
            legacy_final * (1.0 - hierarchy_weight) +
            hierarchical.get('score', legacy_final) * hierarchy_weight
        )
        final = max(legacy_final, hierarchical_blend)
        # A document-level floor is reserved for exceptionally broad evidence.
        # The stricter gate was calibrated against pre-LLM academic papers so
        # ordinary technical regularity cannot activate it.
        if (
            features.get('n_words', 0) >= 2500 and
            global_score >= 0.56 and
            features.get('evidence_families', 0) >= 8 and
            global_meta.get('support', 0) >= 7 and
            global_meta.get('strong', 0) >= 3 and
            hierarchical.get('cross_scale_agreement', 0.0) >= 0.58
        ):
            document_floor = global_score * (
                0.70 + 0.08 * hierarchical.get('cross_scale_agreement', 0.0)
            )
            final = max(final, document_floor)

    n_words = max(features['n_words'], 1)
    if n_words < 150:
        final = min(final * 0.95, 0.78)
    final = _clamp(final)

    length_factor = min(1.0, n_words / 1100.0)
    chunk_factor = min(1.0, len(chunk_scores) / 7.0)
    evidence_factor = min(1.0, features['evidence_families'] / 10.0)
    agreement_factor = global_meta['agreement']
    bootstrap_factor = 1.0 - _clamp(bootstrap['width'] / 0.38)
    confidence = _clamp(
        0.08 + 0.28 * length_factor + 0.16 * chunk_factor +
        0.20 * evidence_factor + 0.10 * agreement_factor +
        0.06 * bootstrap_factor + 0.08 * features['markov_reliability'] +
        0.04 * hierarchical.get('reliability', 0.0),
        0.05, 0.96
    )

    if final >= 0.68:
        decision, verdict, color = 'STRONG_FINGERPRINT', 'بصمة أسلوبية مرتفعة ومتعددة المصادر — يلزم تحقق بشري', '#c0392b'
    elif final >= 0.48:
        decision, verdict, color = 'ELEVATED_FINGERPRINT', 'بصمة أسلوبية مرتفعة — تحتاج مراجعة دقيقة', '#e67e22'
    elif final >= 0.30:
        decision, verdict, color = 'MIXED_FINGERPRINT', 'بصمة مختلطة — لا تكفي للحكم منفردة', '#f39c12'
    elif final >= 0.12:
        decision, verdict, color = 'LIMITED_FINGERPRINT', 'بصمة محدودة أو غير مستقرة', '#27ae60'
    else:
        decision, verdict, color = 'NO_CLEAR_FINGERPRINT', 'لا توجد بصمة أسلوبية كافية للحكم', '#2ecc71'

    percentage = round(final * 100.0, 1)
    if features['verification_score'] >= 0.58 and global_meta.get('other_support', 0) >= 2:
        verification_status = 'CORROBORATED'
        verification_message = 'التحقق المتقاطع دعم النتيجة من خلال تنبؤ بنيوي أعلى بوضوح من الضابط العشوائي.'
    elif features['verification_score'] <= 0.28 and final >= 0.30:
        verification_status = 'CONFLICTING'
        verification_message = 'طبقة التحقق المستقلة لم تؤيد الارتفاع؛ يجب التعامل مع النتيجة بحذر أكبر.'
    else:
        verification_status = 'PARTIAL'
        verification_message = 'التحقق المتقاطع قدّم دعمًا جزئيًا أو كانت كمية النص غير كافية لتأييد قوي.'

    signal_names = {
        'rhythm': 'Internal evidence family 1',
        'repetition': 'Internal evidence family 2',
        'lexical': 'Internal evidence family 3',
        'distribution': 'Internal evidence family 4',
        'predictability': 'Internal evidence family 5',
        'verification': 'Internal evidence family 6',
        'diversity': 'Internal evidence family 7',
        'invariance': 'Internal evidence family 8',
        'trajectory': 'Internal evidence family 9',
        'formulaicity': 'Internal evidence family 10',
    }
    top_signals = sorted(
        [(signal_names[k], features['head_scores'][k]) for k in features['head_scores']],
        key=lambda x: x[1], reverse=True
    )

    return {
        'engine_version': AI_ENGINE_VERSION,
        'engine_mode': 'english_only_non_neural_fingerprint',
        'score': round(final, 4),
        'raw_score_before_safety_gates': round(global_score, 4),
        'legacy_document_score': round(legacy_final, 4),
        'hierarchical_document_score': round(hierarchical.get('score', global_score), 4),
        'hierarchical_reliability': round(hierarchical.get('reliability', 0.0), 4),
        'hierarchical_scale_count': hierarchical.get('scale_count', 0),
        'hierarchical_window_count': hierarchical.get('window_count', 0),
        'hierarchical_cross_scale_agreement': round(hierarchical.get('cross_scale_agreement', 0.0), 4),
        'hierarchical_scale_scores': [round(v, 4) for v in hierarchical.get('scale_scores', [])],
        'hierarchical_strong_window_share': round(hierarchical.get('strong_window_share', 0.0), 4),
        'percentage': percentage,
        'public_percentage': _public_whole_percentage(percentage),
        'low_score_masked': _low_score_is_masked(percentage),
        'highlighting_allowed': _highlighting_allowed(percentage),
        'display_percentage': _format_percentage(percentage),
        'human_score': round((1.0 - final) * 100.0, 1),
        'confidence': round(confidence, 4),
        'confidence_percentage': round(confidence * 100.0, 1),
        'decision': decision, 'verdict': verdict, 'color': color,
        'style_score': round(features['style_score'], 4),
        'rhythm_score': round(features['rhythm_score'], 4),
        'repetition_score': round(features['repetition_score'], 4),
        'lexical_score': round(features['lexical_score'], 4),
        'distribution_score': round(features['distribution_score'], 4),
        'predictability_score': round(features['predictability_score'], 4),
        'verification_score': round(features['verification_score'], 4),
        'diversity_score': round(features['diversity_score'], 4),
        'invariance_score': round(features['invariance_score'], 4),
        'trajectory_score': round(features['trajectory_score'], 4),
        'formulaicity_score': round(features['formulaicity_score'], 4),
        'anti_evasion_score': round(features['anti_evasion_score'], 4),
        'anti_evasion_agreement': round(features['anti_evasion_agreement'], 4),
        'anti_evasion_support': features['anti_evasion_support'],
        'verification_status': verification_status,
        'verification_message': verification_message,
        'markov_information_gain': round(features['markov_information_gain'], 4),
        'markov_shuffle_contrast': round(features['markov_shuffle_contrast'], 4),
        'markov_transition_coverage': round(features['markov_transition_coverage'], 4),
        'markov_surprisal_cv': round(features['markov_surprisal_cv'], 4),
        'markov_folds': features['markov_folds'],
        'markov_reliability': round(features['markov_reliability'], 4),
        'top_signals': [(name, round(value, 4)) for name, value in top_signals],
        'avg_sent_len': round(features['avg_sent_len'], 1),
        'sent_len_cv': round(features['sent_len_cv'], 4),
        'sentence_mad_ratio': round(features['sentence_mad_ratio'], 4),
        'sentence_qspread': round(features['sentence_qspread'], 4),
        'sentence_autocorr': round(features['sentence_autocorr'], 4),
        'sentence_outlier_share': round(features['sentence_outlier_share'], 4),
        'word_ngram_repeat_ratio': round(features['word_ngram_repeat_ratio'], 4),
        'char_repeat_ratio': round(features['char_repeat_ratio'], 4),
        'char_entropy': round(features['char_entropy'], 4),
        'opener_repeat_ratio': round(features['opener_repeat_ratio'], 4),
        'template_repeat_ratio': round(features['template_repeat_ratio'], 4),
        'function_word_js': round(features['function_word_js'], 4),
        'punctuation_js': round(features['punctuation_js'], 4),
        'compression_ratio': round(features['compression_ratio'], 4),
        'chunk_similarity': round(features['chunk_similarity'], 4),
        'mattr40': round(features['mattr40'], 4),
        'mattr80': round(features['mattr80'], 4),
        'lexical_hapax_ratio': round(features['hapax_ratio'], 4),
        'lexical_entropy': round(features['lexical_entropy'], 4),
        'yule_k': round(features['yule_k'], 3),
        'zipf_r2': round(features['zipf_r2'], 4),
        'flesch_mean': round(features['flesch_mean'], 2),
        'top_kw': features['top_kw'], 'top_kw_count': features['top_kw_count'],
        'n_words': n_words, 'n_sents': max(features['n_sents'], 1),
        'chunk_count': len(chunk_scores),
        'chunk_scores': [round(s * 100.0, 1) for s in chunk_scores],
        'chunk_median': round(median_score * 100.0, 1),
        'chunk_q75': round(q75_score * 100.0, 1),
        'chunk_q90': round(q90_score * 100.0, 1),
        'high_risk_chunk_share': round(high_share * 100.0, 1),
        'medium_risk_chunk_share': round(medium_share * 100.0, 1),
        'chunk_dispersion': round(dispersion, 4),
        'bootstrap_low': round(bootstrap['low'] * 100.0, 1),
        'bootstrap_high': round(bootstrap['high'] * 100.0, 1),
        'bootstrap_width': round(bootstrap['width'], 4),
        'evidence_families': features['evidence_families'],
        'style_support_count': features['support_count'],
        'strong_style_count': features['strong_support_count'],
        'english_confidence': round(language['confidence'], 4),
        'english_latin_ratio': round(language['latin_ratio'], 4),
        'original_words': prep['original_words'],
        'reference_words_excluded': prep['reference_words_excluded'],
        'inline_citation_words_removed': prep['inline_citation_words_removed'],
        'reference_section_found': prep['reference_section_found'],
        'reference_header_found': prep['reference_header_found'],
        'analysis_text': analyzed_text,
        'breakdown': {
            'Rhythm fingerprint': round(features['rhythm_score'], 4),
            'Repetition fingerprint': round(features['repetition_score'], 4),
            'Lexical fingerprint': round(features['lexical_score'], 4),
            'Distribution fingerprint': round(features['distribution_score'], 4),
            'Predictability fingerprint': round(features['predictability_score'], 4),
            'Independent verification': round(features['verification_score'], 4),
            'Diversity consistency': round(features['diversity_score'], 4),
            'Rewrite-invariant structure': round(features['invariance_score'], 4),
            'Style trajectory': round(features['trajectory_score'], 4),
            'Academic formulaicity': round(features['formulaicity_score'], 4),
            'Anti-evasion consensus': round(features['anti_evasion_score'], 4),
        },
    }


def _score_paragraph_for_highlight(para: str) -> dict:
    """
    يحسب قوة الإشارة المحلية داخلياً دون كشف تفاصيل المحرك للمستخدم.
    يحتاج التظليل الأساسي إلى اتفاق عدة مجموعات مستقلة من الأدلة.
    """
    raw_text = (para or '').strip()
    empty = {
        'text': '', 'word_count': 0, 'self_ref': 0,
        'style_score': 0.0, 'evidence_score': 0.0,
        'raw_probability': 0.0, 'density': 0.0,
        'explicit_signal': False, 'signal_tier': 'none',
        'evidence_families': 0, 'support_count': 0,
    }
    if not raw_text or _is_reference_header(raw_text) or _is_reference_line(raw_text):
        return empty

    text = _strip_inline_citations(raw_text)
    profile = _english_language_profile(text)
    if profile['latin_ratio'] < 0.965 or profile['arabic_alpha_chars'] > max(2, int(len(text) * 0.002)):
        return empty

    features = _compute_stylometric_features(text)
    probability, meta = _score_feature_probability(features, local=True)
    raw_probability = float(probability)
    n_words = features['n_words']

    resistant_local = (
        features.get('anti_evasion_non_formula_support', 0) >= 2 and
        features.get('anti_evasion_score', 0.0) >= 0.52
    )
    strong_signal = (
        n_words >= 34 and raw_probability >= 0.44 and
        meta['support'] >= 4 and features['evidence_families'] >= 4
    ) or (
        n_words >= 28 and raw_probability >= 0.40 and resistant_local and
        meta.get('core_support', 0) >= 2
    )
    weak_signal = (
        n_words >= 16 and raw_probability >= 0.21 and
        meta['support'] >= 3 and features['evidence_families'] >= 3
    ) or (
        n_words >= 18 and raw_probability >= 0.19 and resistant_local
    )
    if strong_signal:
        tier = 'strong'
        evidence_score = raw_probability
    elif weak_signal:
        tier = 'weak'
        evidence_score = min(raw_probability, 0.44)
    else:
        tier = 'none'
        evidence_score = 0.0

    density = raw_probability / max(math.sqrt(max(n_words, 1)), 1.0)
    return {
        'text': text, 'word_count': n_words, 'self_ref': 0,
        'style_score': round(features['style_score'], 4),
        'evidence_score': round(evidence_score, 6),
        'raw_probability': round(raw_probability, 6),
        'signal_tier': tier, 'density': round(density, 6),
        'explicit_signal': tier != 'none',
        'evidence_families': features['evidence_families'],
        'support_count': meta.get('support', 0),
        'rhythm_score': round(features['rhythm_score'], 4),
        'repetition_score': round(features['repetition_score'], 4),
        'lexical_score': round(features['lexical_score'], 4),
        'anti_evasion_score': round(features.get('anti_evasion_score', 0.0), 4),
        'invariance_score': round(features.get('invariance_score', 0.0), 4),
        'diversity_score': round(features.get('diversity_score', 0.0), 4),
    }


def _merge_adjacent_blocks(blocks, gap_threshold: float = 18.0):
    """
    يدمج الـ blocks النصية المتجاورة رأسياً في فقرة واحدة حقيقية.
    يعيد قائمة من الفقرات، كل فقرة = قائمة line_data مع bbox كل سطر على حدة.
    """
    text_blocks = [b for b in blocks if b.get("type", 0) == 0]
    text_blocks.sort(key=lambda b: b["bbox"][1])

    paragraphs = []

    for block in text_blocks:
        lines = block.get("lines", [])
        line_data = []
        for line in lines:
            spans = line.get("spans", [])
            txt = "".join((sp.get("text", "") or "") for sp in spans).strip()
            if txt:
                line_data.append({
                    'text':  txt,
                    'bbox':  fitz.Rect(line["bbox"]),
                    'y0':    line["bbox"][1],
                    'y1':    line["bbox"][3],
                    # نحفظ span-level data لتحديد نقطة بدء/نهاية دقيقة
                    'spans': line.get("spans", []),
                })
        if not line_data:
            continue

        if paragraphs:
            last_para = paragraphs[-1]
            last_y1   = last_para['lines'][-1]['y1']
            this_y0   = line_data[0]['y0']
            gap       = this_y0 - last_y1
            last_x0   = last_para['lines'][0]['bbox'].x0
            this_x0   = line_data[0]['bbox'].x0
            if gap <= gap_threshold and abs(last_x0 - this_x0) < 40.0:
                last_para['lines'].extend(line_data)
                last_para['full_text'] = " ".join(d['text'] for d in last_para['lines'])
                continue

        paragraphs.append({
            'lines':     line_data,
            'full_text': " ".join(d['text'] for d in line_data),
        })

    return paragraphs


def _sentence_start_x(line_data_item) -> float:
    """
    يجد الـ x الذي تبدأ منه أول جملة في السطر.
    إذا كان السطر يبدأ بحرف كبير بعد نقطة سابقة → نقطة البداية الفعلية للجملة.
    في حالات أخرى نأخذ x0 السطر كما هو.
    """
    return line_data_item['bbox'].x0


def _punct_end_x(line_data_item) -> float:
    """
    يجد الـ x الذي تنتهي عنده آخر علامة ترقيم (. أو ،) في السطر.
    نبحث في الـ spans عن آخر حرف قبل المسافات النهائية.
    إذا لم يوجد → نأخذ x1 السطر كاملاً.
    """
    txt = line_data_item['text'].rstrip()
    if not txt:
        return line_data_item['bbox'].x1
    # آخر حرف هو علامة ترقيم؟
    if re.search(r'[.,;?!]\s*$', txt):
        # نأخذ x1 كاملاً (الترقيم هو الحرف الأخير)
        return line_data_item['bbox'].x1
    return line_data_item['bbox'].x1


def _build_line_rects_turnitin(chunk_lines):
    """
    يبني قائمة من المستطيلات بأسلوب Turnitin:
    - كل سطر له مستطيله المستقل (سطر تلو الآخر)
    - السطر الأول: من بداية الجملة (x0 السطر)
    - السطر الأخير: حتى آخر علامة ترقيم (x1 أو موضع الترقيم)
    - الأسطر الوسطى: من x0 إلى x1 كاملاً
    """
    rects = []
    n = len(chunk_lines)
    for i, ld in enumerate(chunk_lines):
        b = ld['bbox']
        x0 = b.x0
        x1 = b.x1
        y0 = b.y0
        y1 = b.y1

        if i == 0:
            # السطر الأول: من بداية الجملة
            x0 = _sentence_start_x(ld)
        if i == n - 1:
            # السطر الأخير: حتى الترقيم
            x1 = _punct_end_x(ld)

        rects.append(fitz.Rect(x0, y0, x1, y1))
    return rects


def _collect_page_block_candidates(page, in_refs_flag: bool):
    """
    يجمع مرشحاً واحداً كحد أقصى من كل فقرة. يبدأ بمسار صارم، ثم يحتفظ
    بمرشح نسبي احتياطي لا يستخدم إلا عند تجاوز النتيجة العامة 20%.
    """
    candidates = []
    in_refs = in_refs_flag
    analyzable_words = 0

    MIN_LINES = 2
    MAX_LINES = 7
    MIN_LOCAL_EVIDENCE = 0.24
    WEAK_LOCAL_EVIDENCE = 0.18
    FALLBACK_RAW_FLOOR = 0.10

    def _meets_local_threshold(local_res: dict) -> bool:
        tier = local_res.get('signal_tier', 'none')
        if tier == 'strong':
            return local_res.get('evidence_score', 0.0) >= MIN_LOCAL_EVIDENCE
        if tier == 'weak':
            return local_res.get('evidence_score', 0.0) >= WEAK_LOCAL_EVIDENCE
        return False

    def _make_candidate(lines, text_value, local_res, fallback_only=False, score_boost=1.0):
        rects = _build_line_rects_turnitin(lines)
        if not rects:
            return None
        combined = rects[0]
        for rect in rects[1:]:
            combined = combined | rect
        base_score = local_res.get('raw_probability', 0.0) if fallback_only else local_res.get('evidence_score', 0.0)
        return {
            'rects': rects,
            'rect': combined,
            'text': text_value,
            'word_count': local_res.get('word_count', 0),
            'score': float(base_score) * float(score_boost),
            'raw_probability': local_res.get('raw_probability', 0.0),
            'density': local_res.get('density', 0.0),
            'tier': local_res.get('signal_tier', 'none'),
            'evidence_families': local_res.get('evidence_families', 0),
            'support_count': local_res.get('support_count', 0),
            'self_ref': local_res.get('self_ref', 0),
            'explicit_signal': local_res.get('explicit_signal', False),
            'line_count': len(lines),
            'fallback_only': bool(fallback_only),
        }

    try:
        page_dict = page.get_text("dict")
        raw_blocks = page_dict.get("blocks", [])
    except Exception:
        raw_blocks = []

    paragraphs = _merge_adjacent_blocks(raw_blocks)

    for para in paragraphs:
        original_lines = para['lines']
        if not original_lines:
            continue

        header_idx = next(
            (i for i, ld in enumerate(original_lines) if _is_reference_header(ld.get('text', ''))),
            None
        )
        set_refs_after = header_idx is not None

        if in_refs:
            continue

        if header_idx == 0:
            in_refs = True
            continue
        elif header_idx is not None:
            line_data = original_lines[:header_idx]
        else:
            line_data = original_lines

        btext = " ".join(d['text'] for d in line_data).strip()
        if not btext:
            if set_refs_after:
                in_refs = True
            continue

        if _is_reference_line(btext):
            if set_refs_after:
                in_refs = True
            continue

        clean_btext = _strip_inline_citations(btext)
        analyzable_words += _count_text_words(clean_btext)
        total_lines = len(line_data)

        tested = []
        full_local = _score_paragraph_for_highlight(btext)
        tested.append((line_data, btext, full_local, 1.0))

        if total_lines >= MIN_LINES:
            for chunk_len in range(MIN_LINES, min(MAX_LINES, total_lines) + 1):
                for start_i in range(0, total_lines - chunk_len + 1):
                    chunk_lines = line_data[start_i:start_i + chunk_len]
                    chunk_text = " ".join(d['text'] for d in chunk_lines).strip()
                    if not chunk_text:
                        continue
                    local = _score_paragraph_for_highlight(chunk_text)
                    ends_with_punct = bool(re.search(r'[.,;?!]\s*$', chunk_lines[-1]['text'].rstrip()))
                    tested.append((chunk_lines, chunk_text, local, 1.06 if ends_with_punct else 1.0))

        primary = []
        fallback = []
        for lines, txt, local, boost in tested:
            if local.get('word_count', 0) < 8:
                continue
            if _meets_local_threshold(local):
                cand = _make_candidate(lines, txt, local, False, boost)
                if cand:
                    primary.append(cand)
            elif (
                local.get('word_count', 0) >= 12 and
                local.get('raw_probability', 0.0) >= FALLBACK_RAW_FLOOR and
                local.get('evidence_families', 0) >= 2 and
                local.get('support_count', 0) >= 1
            ):
                cand = _make_candidate(lines, txt, local, True, boost)
                if cand:
                    fallback.append(cand)

        if primary:
            candidates.append(max(primary, key=lambda x: (
                1 if x.get('tier') == 'strong' else 0,
                x.get('score', 0.0), x.get('evidence_families', 0),
                x.get('density', 0.0), -x.get('word_count', 0)
            )))
        elif fallback:
            candidates.append(max(fallback, key=lambda x: (
                x.get('score', 0.0), x.get('evidence_families', 0),
                x.get('support_count', 0), x.get('density', 0.0),
                -x.get('word_count', 0)
            )))
        elif full_local.get('word_count', 0) >= 12:
            # يحتفظ بالفقرة لمسار المقارنة السياقية على مستوى المستند كله.
            context_candidate = _make_candidate(line_data, btext, full_local, True, 1.0)
            if context_candidate:
                context_candidate['context_only'] = True
                candidates.append(context_candidate)

        if set_refs_after:
            in_refs = True

    return candidates, in_refs, analyzable_words



def _review_flag_texts(analysis_result: dict | None) -> list[str]:
    review = (analysis_result or {}).get("gemini_review") or {}
    texts = []
    for flag in review.get("flagged_passages") or []:
        for key in ("quote", "source_text"):
            value = re.sub(r"\s+", " ", str(flag.get(key) or "")).strip()
            if len(_style_tokens(value)) >= 5 and value not in texts:
                texts.append(value)
    return texts


def _text_matches_review_flag(text: str, analysis_result: dict | None) -> bool:
    normalized = re.sub(r"\s+", " ", text or "").strip().lower()
    if not normalized:
        return False
    normalized_tokens = _style_tokens(normalized)
    for flag_text in _review_flag_texts(analysis_result):
        flag_normalized = re.sub(r"\s+", " ", flag_text).strip().lower()
        if flag_normalized in normalized or normalized in flag_normalized:
            return True
        flag_tokens = _style_tokens(flag_normalized)
        if len(flag_tokens) >= 7 and len(normalized_tokens) >= 7:
            flag_grams = {tuple(flag_tokens[i:i+5]) for i in range(len(flag_tokens) - 4)}
            text_grams = {tuple(normalized_tokens[i:i+5]) for i in range(len(normalized_tokens) - 4)}
            if flag_grams and len(flag_grams & text_grams) / len(flag_grams) >= 0.35:
                return True
    return False


def _dedupe_rectangles(rects: list) -> list:
    unique = []
    for rect in rects:
        if rect is None or rect.is_empty:
            continue
        duplicate = False
        for existing in unique:
            intersection = rect & existing
            min_area = max(1e-6, min(rect.get_area(), existing.get_area()))
            if intersection.get_area() / min_area >= 0.70:
                duplicate = True
                break
        if not duplicate:
            unique.append(rect)
    return unique


def _smart_review_highlight_rects(orig_doc, analysis_result: dict | None) -> dict:
    """Locate every service-flagged quote on original PDF pages and return exact rectangles."""
    flag_texts = _review_flag_texts(analysis_result)
    if not flag_texts:
        return {"plan_by_page": {}, "matched_flags": 0, "matched_rects": 0, "matched_words": 0}

    reference_start_page = len(orig_doc)
    reference_cutoff_y = None
    for page_idx in range(len(orig_doc)):
        try:
            page_blocks = orig_doc[page_idx].get_text("blocks") or []
        except Exception:
            page_blocks = []
        header_blocks = [
            block for block in page_blocks
            if len(block) >= 5 and _is_reference_header(str(block[4] or ""))
        ]
        if header_blocks:
            reference_start_page = page_idx
            reference_cutoff_y = min(float(block[1]) for block in header_blocks)
            break

    plan_by_page = {}
    matched_flags = 0
    matched_words = 0
    for flag_text in flag_texts:
        words = flag_text.split()
        search_variants = [flag_text]
        for length in (28, 20, 14, 10):
            if len(words) >= length:
                search_variants.append(" ".join(words[:length]))
                search_variants.append(" ".join(words[-length:]))
        # Keep order while removing duplicate variants.
        search_variants = list(dict.fromkeys(v for v in search_variants if v.strip()))
        found_for_flag = False
        last_page = reference_start_page if reference_start_page < len(orig_doc) else len(orig_doc) - 1
        for page_idx in range(max(0, last_page) + 1):
            page = orig_doc[page_idx]
            page_rects = []
            for variant in search_variants:
                try:
                    page_rects = page.search_for(variant, quads=False)
                except Exception:
                    page_rects = []
                if page_rects:
                    break
            if page_idx == reference_start_page and reference_cutoff_y is not None:
                page_rects = [rect for rect in page_rects if rect.y1 < reference_cutoff_y]
            if page_rects:
                plan_by_page.setdefault(page_idx, []).extend(page_rects)
                found_for_flag = True
        if found_for_flag:
            matched_flags += 1
            matched_words += _count_text_words(flag_text)

    for page_idx, rects in list(plan_by_page.items()):
        plan_by_page[page_idx] = _dedupe_rectangles(rects)
    return {
        "plan_by_page": plan_by_page,
        "matched_flags": matched_flags,
        "matched_rects": sum(len(v) for v in plan_by_page.values()),
        "matched_words": matched_words,
    }


def _select_highlight_plan(orig_doc, target_pct: float):
    """
    كل نتيجة يقع رقمها الظاهر بين 0 و20 لا يوجد لها أي تظليل وتظهر *%. يبدأ الاختيار عندما يظهر 21% فأعلى بالمرشحات
    الصارمة، ثم يستخدم أعلى المرشحات النسبية فقط عند نقص التغطية.
    """
    pages_candidates = []
    in_refs_global = False
    total_words = 0

    for page_idx in range(len(orig_doc)):
        page = orig_doc[page_idx]
        candidates, in_refs_global, page_words = _collect_page_block_candidates(page, in_refs_global)
        total_words += page_words
        for item in candidates:
            item['page_idx'] = page_idx
        pages_candidates.extend(candidates)

    total_words = max(total_words, 1)

    if not _highlighting_allowed(target_pct):
        return {
            'plan_by_page': {}, 'total_words': total_words, 'target_words': 0,
            'covered_words': 0, 'achieved_pct': 0.0, 'selected_blocks': 0,
            'available_candidate_words': sum(i['word_count'] for i in pages_candidates),
            'suppressed_low_score': True, 'fallback_blocks_used': 0,
        }

    target_words = max(1, int(round(total_words * max(0.0, min(100.0, target_pct)) / 100.0)))

    # إشارات سياقية صارمة على مستوى المستند: تكرار مقاطع، تشابه بنيوي،
    # واتساق الفقرة مع النمط العام. هذه المرحلة تعالج حالة ارتفاع النتيجة
    # العامة بينما تكون كل فقرة قصيرة فلا تكفي وحدها لإظهار الإشارة.
    normalized_texts = []
    exact_counts = collections.Counter()
    document_text_parts = []
    for item in pages_candidates:
        norm = ' '.join(_style_tokens(item.get('text', '')))
        normalized_texts.append(norm)
        if norm:
            exact_counts[norm] += 1
            document_text_parts.append(item.get('text', ''))

    gram_document_frequency = collections.Counter()
    candidate_grams = []
    for norm in normalized_texts:
        toks = norm.split()
        grams = {tuple(toks[i:i+4]) for i in range(max(0, len(toks) - 3))}
        candidate_grams.append(grams)
        gram_document_frequency.update(grams)

    document_vector = _chunk_vector(' '.join(document_text_parts)) if document_text_parts else []
    for item, norm, grams in zip(pages_candidates, normalized_texts, candidate_grams):
        exact_signal = 1.0 if norm and exact_counts[norm] >= 2 else 0.0
        repeated_share = (
            sum(1 for gram in grams if gram_document_frequency.get(gram, 0) >= 2) / len(grams)
            if grams else 0.0
        )
        repetition_context = _clamp((repeated_share - 0.10) / 0.65)
        item_vector = _chunk_vector(item.get('text', ''))
        style_similarity = _cosine_similarity(item_vector, document_vector) if document_vector else 0.0
        style_context = _clamp((style_similarity - 0.90) / 0.09)
        context_score = max(exact_signal, 0.62 * repetition_context + 0.38 * style_context)
        local_score = item.get('raw_probability', 0.0)
        item['context_score'] = context_score
        item['combined_review_score'] = max(local_score, 0.58 * local_score + 0.42 * context_score)

    primary = sorted(
        [i for i in pages_candidates if not i.get('fallback_only', False)],
        key=lambda x: (
            1 if x.get('tier') == 'strong' else 0,
            x.get('density', 0.0), x.get('score', 0.0),
            x.get('evidence_families', 0), -x.get('word_count', 0)
        ), reverse=True
    )

    fallback_floor = max(0.10, min(0.28, (float(target_pct) / 100.0) * 0.45))
    fallback = sorted(
        [
            i for i in pages_candidates
            if i.get('fallback_only', False)
            and i.get('combined_review_score', 0.0) >= fallback_floor
            and (
                i.get('evidence_families', 0) >= 2
                or i.get('context_score', 0.0) >= 0.42
            )
        ],
        key=lambda x: (
            x.get('combined_review_score', 0.0), x.get('context_score', 0.0),
            x.get('raw_probability', 0.0), x.get('evidence_families', 0),
            x.get('support_count', 0), x.get('density', 0.0),
            -x.get('word_count', 0)
        ), reverse=True
    )

    # Highlight every passage that independently satisfies the suspicious-passage rules.
    # The final percentage is a document-level estimate; it is no longer used as a quota
    # that can hide qualifying passages or force weak passages into the report.
    selected = list(primary)
    selected_ids = {id(item) for item in selected}
    selected.extend(item for item in fallback if id(item) not in selected_ids)
    covered_words = sum(item['word_count'] for item in selected)
    fallback_blocks_used = sum(1 for item in selected if item.get('fallback_only', False))

    plan_by_page = {}
    for item in selected:
        plan_by_page.setdefault(item['page_idx'], []).extend(item.get('rects') or [item['rect']])

    achieved_pct = round((covered_words / total_words) * 100.0, 1)
    return {
        'plan_by_page': plan_by_page, 'total_words': total_words,
        'target_words': target_words, 'covered_words': covered_words,
        'achieved_pct': achieved_pct, 'selected_blocks': len(selected),
        'available_candidate_words': sum(i['word_count'] for i in pages_candidates),
        'suppressed_low_score': False, 'fallback_blocks_used': fallback_blocks_used,
    }


def generate_report_pdf_from_original(original_pdf_bytes: bytes,
                                       text: str,
                                       doc_name: str = "Document",
                                       analysis_result: dict | None = None) -> bytes:
    """
    يستخدم ملف PDF الأصلي مباشرة:
    - يحتفظ بكل صفحة كما هي (نص + صور + جداول + تنسيق)
    - يُضيف تظليل أصفر/برتقالي شفاف فوق الفقرات المشبوهة كاملة فقط
    - يظلل جميع الفقرات التي اجتازت قواعد الاشتباه من المحرك الداخلي أو المراجعة الذكية
    - وحدة التظليل هي الفقرة كاملة بكل أسطرها، وليس جزءاً من الجملة أو بعض الأسطر
    - لا يُظلّل قسم المراجع
    - يُضيف صفحة غلاف احترافية في البداية
    - يُضيف صفحة ملخص في النهاية
    - يعرض *% دون تظليل من 0 إلى 20، ويظهر الرقم ويظلل المقاطع المؤهلة من 21 فأعلى
    """
    if not FITZ_OK or not RLAB_OK:
        return b""

    result = analysis_result or _ai_fingerprint_score(text)
    _combined_report = result.get('combined_analysis') or {}
    pct = float(_combined_report.get('score', result['percentage']))
    report_color = _combined_report.get('color', result['color'])
    today  = datetime.now().strftime("%Y-%m-%d  %H:%M")

    try:
        orig_doc = fitz.open(stream=original_pdf_bytes, filetype="pdf")
    except Exception:
        return b""

    allow_highlighting = _highlighting_allowed(pct)
    highlight_plan = _select_highlight_plan(orig_doc, pct)
    smart_plan = _smart_review_highlight_rects(orig_doc, result)
    if allow_highlighting:
        for page_idx, rects in smart_plan.get('plan_by_page', {}).items():
            merged = list(highlight_plan['plan_by_page'].get(page_idx, [])) + list(rects)
            highlight_plan['plan_by_page'][page_idx] = _dedupe_rectangles(merged)
        highlight_plan['smart_flags_matched'] = smart_plan.get('matched_flags', 0)
        highlight_plan['smart_rects_added'] = smart_plan.get('matched_rects', 0)
        highlight_plan['selected_blocks'] += smart_plan.get('matched_flags', 0)
        highlight_plan['covered_words'] = min(
            highlight_plan.get('total_words', 0),
            highlight_plan.get('covered_words', 0) + smart_plan.get('matched_words', 0)
        )
        highlight_plan['achieved_pct'] = round(
            100.0 * highlight_plan['covered_words'] / max(highlight_plan.get('total_words', 1), 1), 1
        )
    # حاجز أمان ثانٍ مستقل: لا تمر أي مستطيلات عندما لا توجد نتيجة قابلة للتظليل.
    # مستطيلات تظليل عندما تكون النتيجة الظاهرة 20% أو أقل.
    if not allow_highlighting:
        highlight_plan['plan_by_page'] = {}
        highlight_plan['target_words'] = 0
        highlight_plan['covered_words'] = 0
        highlight_plan['achieved_pct'] = 0.0
        highlight_plan['selected_blocks'] = 0
        highlight_plan['suppressed_low_score'] = True
        highlight_plan['fallback_blocks_used'] = 0

    result['highlight_target_words'] = highlight_plan['target_words']
    result['highlight_covered_words'] = highlight_plan['covered_words']
    result['highlight_total_words'] = highlight_plan['total_words']
    result['highlight_achieved_pct'] = highlight_plan['achieved_pct']
    result['highlight_selected_blocks'] = highlight_plan['selected_blocks']
    result['highlight_suppressed_low_score'] = highlight_plan.get('suppressed_low_score', False)
    result['highlight_fallback_blocks_used'] = highlight_plan.get('fallback_blocks_used', 0)
    result['smart_flags_matched'] = highlight_plan.get('smart_flags_matched', 0)
    result['smart_rects_added'] = highlight_plan.get('smart_rects_added', 0)

    cover_buf = io.BytesIO()
    W_pt, H_pt = A4
    cv = rl_canvas.Canvas(cover_buf, pagesize=A4)
    _draw_cover_rl(cv, W_pt, H_pt, result, doc_name, today)
    cv.save()
    cover_buf.seek(0)

    cover_doc = fitz.open(stream=cover_buf.read(), filetype="pdf")

    out_doc = fitz.open()
    out_doc.insert_pdf(cover_doc)

    for page_idx in range(len(orig_doc)):
        out_doc.insert_pdf(orig_doc, from_page=page_idx, to_page=page_idx)
        new_page = out_doc[-1]

        highlights = (
            highlight_plan['plan_by_page'].get(page_idx, [])
            if allow_highlighting else []
        )
        for line_rect in highlights:
            # كل line_rect هو مستطيل سطر مستقل → annotation منفصل (أسلوب Turnitin)
            annot = new_page.add_highlight_annot(line_rect)
            annot.set_colors(stroke=[0.0, 0.75, 1.0])   # أزرق سماوي
            annot.set_opacity(0.40)
            annot.update()

        _stamp_page_header(new_page, page_idx + 1, len(orig_doc), pct, report_color)

    summary_buf = io.BytesIO()
    cv2 = rl_canvas.Canvas(summary_buf, pagesize=A4)
    _draw_summary_page(cv2, W_pt, H_pt, result, doc_name, today)
    cv2.save()
    summary_buf.seek(0)
    summary_doc = fitz.open(stream=summary_buf.read(), filetype="pdf")
    out_doc.insert_pdf(summary_doc)

    gemini_doc = None
    if (result.get("gemini_review") or {}).get("available"):
        gemini_buf = io.BytesIO()
        cv3 = rl_canvas.Canvas(gemini_buf, pagesize=A4)
        _draw_gemini_review_page(cv3, W_pt, H_pt, result, doc_name, today)
        cv3.save()
        gemini_buf.seek(0)
        gemini_doc = fitz.open(stream=gemini_buf.read(), filetype="pdf")
        out_doc.insert_pdf(gemini_doc)

    out_buf = io.BytesIO()
    out_doc.save(out_buf, garbage=4, deflate=True)
    orig_doc.close()
    cover_doc.close()
    summary_doc.close()
    if gemini_doc is not None:
        gemini_doc.close()
    out_doc.close()
    return out_buf.getvalue()


def _stamp_page_header(page, page_num: int, total_pages: int, pct: float, col_hex: str):
    """يُضيف شريط معلومات خفي في أعلى كل صفحة"""
    pw = page.rect.width
    ph = page.rect.height

    # شريط علوي شفاف
    strip_rect = fitz.Rect(0, 0, pw, 18)
    page.draw_rect(strip_rect, color=None, fill=(0.04, 0.08, 0.16), fill_opacity=0.82)

    # نص على اليسار
    page.insert_text(
        (6, 13), f"Document Review Report  |  Score: {_format_percentage(pct)}",
        fontsize=6.5, color=(0.7, 0.85, 1.0)
    )
    # رقم الصفحة على اليمين
    page.insert_text(
        (pw - 55, 13), f"Page {page_num}/{total_pages}",
        fontsize=6.5, color=(0.7, 0.85, 1.0)
    )


def _draw_cover_rl(c, W, H, result, doc_name, doc_date):
    """صفحة الغلاف الاحترافية"""
    gemini_review = result.get('gemini_review') or {}
    combined_review = result.get('combined_analysis') or {}
    local_pct = result['percentage']
    pct = float(combined_review.get('score', local_pct))
    display_color = combined_review.get('color') or result['color']

    score_colors = {
        '#c0392b': (0.75, 0.22, 0.17),
        '#e67e22': (0.90, 0.49, 0.13),
        '#f39c12': (0.95, 0.61, 0.07),
        '#27ae60': (0.15, 0.68, 0.38),
        '#2ecc71': (0.18, 0.80, 0.44),
    }
    score_color_rl = score_colors.get(display_color, (0.75, 0.22, 0.17))

    # خلفية داكنة
    c.setFillColor(HexColor('#060d1a'))
    c.rect(0, 0, W, H, fill=1, stroke=0)

    # زخارف دائرية
    for (cx, cy, r, cr, ca) in [
        (W*0.1, H*0.88, 200, '#00bcd4', 0.05),
        (W*0.9, H*0.12, 170, '#f4a300', 0.04),
        (W*0.78, H*0.70, 120, '#7c3aed', 0.035),
    ]:
        rgb = tuple(int(cr[i:i+2], 16)/255 for i in (1, 3, 5))
        c.setFillColorRGB(*rgb, alpha=ca)
        c.circle(cx, cy, r, fill=1, stroke=0)

    # شبكة نقاط خفيفة
    c.setFillColorRGB(0.2, 0.5, 0.7, alpha=0.08)
    for gx in range(0, int(W), 30):
        for gy in range(0, int(H), 30):
            c.circle(gx, gy, 1.0, fill=1, stroke=0)

    # شرائط علوية
    for i, (col, th) in enumerate([('#00bcd4', 7), ('#f4a300', 4), ('#7c3aed', 2)]):
        offset = sum(t for _, t in [('#00bcd4', 7), ('#f4a300', 4), ('#7c3aed', 2)][:i])
        c.setFillColor(HexColor(col))
        c.rect(0, H - offset - th, W, th, fill=1, stroke=0)

    # شعار دائري
    lcx, lcy = W/2, H - 105
    c.setFillColorRGB(0.0, 0.74, 0.83, alpha=0.12)
    c.circle(lcx, lcy, 52, fill=1, stroke=0)
    c.setFillColor(HexColor('#00bcd4'))
    c.circle(lcx, lcy, 38, fill=1, stroke=0)
    c.setFillColor(HexColor('#060d1a'))
    c.setFont("Helvetica-Bold", 20)
    c.drawCentredString(lcx, lcy - 7, "AI")

    # عنوان
    c.setFillColor(HexColor('#ffffff'))
    c.setFont("Helvetica-Bold", 26)
    c.drawCentredString(W/2, H - 170, "AI-Assisted Academic Review")
    c.setFillColor(HexColor('#7fb3c8'))
    c.setFont("Helvetica", 11)
    c.drawCentredString(W/2, H - 190, "Unified Screening, Highlighting and Academic Risk Map")

    # فاصل
    c.setStrokeColor(HexColor('#1e3a5f'))
    c.setLineWidth(0.8)
    c.line(60, H - 210, W - 60, H - 210)

    # دائرة النتيجة
    cx, cy = W/2, H - 340
    c.setFillColorRGB(0.06, 0.10, 0.18, alpha=1.0)
    c.circle(cx, cy, 78, fill=1, stroke=0)
    c.setStrokeColorRGB(*score_color_rl)
    c.setLineWidth(4)
    c.circle(cx, cy, 78, fill=0, stroke=1)
    c.setFillColorRGB(*score_color_rl)
    c.setFont("Helvetica-Bold", 46)
    c.drawCentredString(cx, cy - 16, _format_percentage(pct))
    c.setFillColor(HexColor('#aaaaaa'))
    c.setFont("Helvetica", 9)
    c.drawCentredString(cx, cy + 30, "Final Screening Score / 100")

    # الحكم
    c.setFillColorRGB(*score_color_rl)
    c.setFont("Helvetica-Bold", 17)
    _cover_label = (
        "Very high review priority" if pct >= 80 else
        "High review priority" if pct >= 60 else
        "Moderate review priority" if pct >= 40 else
        "Limited review priority" if pct >= 20 else
        "Low review priority"
    )
    c.drawCentredString(W/2, H - 466, _cover_label)
    c.setFillColor(HexColor('#7fb3c8'))
    c.setFont("Helvetica", 10)
    c.drawCentredString(
        W/2, H - 484,
        f"Citation-support risk: {gemini_review.get('citation_integrity_risk_score', 0)}/100"
        f"    |    Local supporting signal: {local_pct:.0f}/100"
    )

    # شريط اللون
    bx, by = 55, H - 538
    bw, bh = W - 110, 18
    for sc, s0, s1 in [
        ('#2ecc71', 0.0, 0.22), ('#27ae60', 0.22, 0.42),
        ('#f39c12', 0.42, 0.62), ('#e67e22', 0.62, 0.80),
        ('#c0392b', 0.80, 1.00),
    ]:
        sx = bx + bw * s0
        sw = bw * (s1 - s0) - 1
        c.setFillColor(HexColor(sc))
        c.roundRect(sx, by, sw, bh, 3, fill=1, stroke=0)

    ind_x = bx + bw * (pct / 100.0)
    c.setFillColor(HexColor('#ffffff'))
    c.setStrokeColor(HexColor('#000000'))
    c.setLineWidth(1.0)
    p = c.beginPath()
    p.moveTo(ind_x, by - 2)
    p.lineTo(ind_x - 6, by - 13)
    p.lineTo(ind_x + 6, by - 13)
    p.close()
    c.drawPath(p, fill=1, stroke=1)

    # صناديق الإحصائيات
    sy = H - 628
    sw_each = (W - 110) / 4
    for i, (lbl, val, vc) in enumerate([
        ("English words", f"{result['n_words']:,}", '#00bcd4'),
        ("Passages reviewed", f"{gemini_review.get('passages_reviewed', 0)}", '#7c3aed'),
        ("Smart-review flags", f"{len(gemini_review.get('flagged_passages') or [])}", '#f4a300'),
        ("Local evidence", f"{result.get('confidence_percentage', 0):.0f}%", '#e74c3c'),
    ]):
        bx2 = 55 + i * sw_each
        c.setFillColor(HexColor('#0b1829'))
        c.roundRect(bx2 + 3, sy - 32, sw_each - 6, 52, 7, fill=1, stroke=0)
        c.setStrokeColor(HexColor(vc))
        c.setLineWidth(0.8)
        c.roundRect(bx2 + 3, sy - 32, sw_each - 6, 52, 7, fill=0, stroke=1)
        c.setFillColor(HexColor(vc))
        c.setFont("Helvetica-Bold", 17)
        c.drawCentredString(bx2 + sw_each/2, sy + 4, val)
        c.setFillColor(HexColor('#888888'))
        c.setFont("Helvetica", 8)
        c.drawCentredString(bx2 + sw_each/2, sy - 20, lbl)

    # معلومات المستند
    iy = H - 708
    c.setFillColor(HexColor('#0b1829'))
    c.roundRect(55, iy - 18, W - 110, 52, 8, fill=1, stroke=0)
    c.setStrokeColor(HexColor('#1e3a5f'))
    c.setLineWidth(0.5)
    c.roundRect(55, iy - 18, W - 110, 52, 8, fill=0, stroke=1)
    c.setFillColor(HexColor('#aaaaaa'))
    c.setFont("Helvetica", 8.5)
    c.drawString(70, iy + 20, f"Document: {doc_name[:55]}")
    c.drawString(70, iy + 7,  f"Analysis Date: {doc_date}")
    h_pct = result.get('highlight_achieved_pct')
    if result.get('highlight_suppressed_low_score'):
        c.drawString(70, iy - 6, "Analysis status: no qualifying passages were found")
    elif h_pct is not None:
        c.drawString(70, iy - 6,  f"Analysis status: highlighted coverage {_format_percentage(h_pct, mask_low=False)}")
    else:
        c.drawString(70, iy - 6,  "Analysis status: original format preserved")

    # شريط أسفل
    c.setFillColor(HexColor('#060d1a'))
    c.rect(0, 0, W, 26, fill=1, stroke=0)
    c.setFillColor(HexColor('#00bcd4'))
    c.rect(0, 25, W, 1.5, fill=1, stroke=0)
    c.setFillColor(HexColor('#8aa0b8'))
    c.setFont("Helvetica", 7)
    c.drawCentredString(W/2, 8,
        "AI Fingerprint Detector  •  Confidential Analysis Report  •  Academic & Research Use Only")


def _draw_summary_page(c, W, H, result, doc_name, doc_date):
    """صفحة ملخص عامة لا تكشف أسماء الخوارزميات أو النماذج الداخلية."""
    c.setFillColor(HexColor('#f8fbff'))
    c.rect(0, 0, W, H, fill=1, stroke=0)
    c.setFillColor(HexColor('#060d1a'))
    c.rect(0, H - 40, W, 40, fill=1, stroke=0)
    c.setFillColor(HexColor('#00bcd4'))
    c.rect(0, H - 41, W, 1.5, fill=1, stroke=0)
    c.setFillColor(HexColor('#aaaaaa'))
    c.setFont("Helvetica", 8)
    c.drawString(18, H - 24, "Document Review Report — English Text")
    _summary_review = result.get('gemini_review') or {}
    _summary_combined = result.get('combined_analysis') or {}
    _summary_score = _summary_combined.get('score', result['percentage'])
    c.drawRightString(W - 18, H - 24, f"Summary  |  FINAL SCORE: {_format_percentage(_summary_score)}")

    y = H - 76
    c.setFillColor(HexColor('#00bcd4'))
    c.setFont("Helvetica-Bold", 15)
    c.drawString(52, y, "Analysis Summary")
    y -= 10
    c.setStrokeColor(HexColor('#1e3a5f'))
    c.line(52, y, W - 52, y)
    y -= 28

    low_score_no_highlight = bool(result.get("highlight_suppressed_low_score")) or not bool(result.get("highlight_selected_blocks", 0))
    if low_score_no_highlight:
        c.setFillColor(HexColor('#e8f5e9'))
        c.roundRect(52, y - 5, W - 104, 28, 5, fill=1, stroke=0)
        c.setStrokeColor(HexColor('#27ae60'))
        c.roundRect(52, y - 5, W - 104, 28, 5, fill=0, stroke=1)
        c.setFillColor(HexColor('#1b5e20'))
        c.setFont("Helvetica-Bold", 9)
        c.drawString(60, y + 8, "No qualifying suspicious passages were found for highlighting.")
        y -= 46
    else:
        h_pct = result.get("highlight_achieved_pct", 0.0)
        h_cov = result.get("highlight_covered_words", 0)
        h_tot = result.get("highlight_total_words", 0)
        c.setFillColor(HexColor('#e3f7ff'))
        c.roundRect(52, y - 5, W - 104, 34, 5, fill=1, stroke=0)
        c.setStrokeColor(HexColor('#00a8cc'))
        c.roundRect(52, y - 5, W - 104, 34, 5, fill=0, stroke=1)
        c.setFillColor(HexColor('#0b3d4d'))
        c.setFont("Helvetica-Bold", 9)
        c.drawString(60, y + 13, "Highlighted passages require manual review; references are excluded.")
        c.setFont("Helvetica", 8.5)
        c.drawString(60, y + 1, f"Highlighted coverage: {_format_percentage(h_pct, mask_low=False)}  |  {h_cov:,} / {h_tot:,} words")
        y -= 54

    metrics = [
        ("FINAL SCREENING SCORE", _format_percentage(_summary_score)),
        ("English words analyzed", f"{result.get('n_words', 0):,}"),
        ("Evidence sufficiency", f"{result.get('confidence_percentage', 0):.0f}/100"),
        ("Flagged passages", f"{len(_summary_review.get('flagged_passages') or []):,}"),
        ("Citation-support risk", f"{_summary_review.get('citation_integrity_risk_score', 0)}/100"),
        ("AI-reviewed passages", f"{_summary_review.get('passages_reviewed', 0):,}"),
    ]
    box_gap = 10
    box_w = (W - 104 - 2 * box_gap) / 3
    box_h = 58
    for idx, (label, value) in enumerate(metrics):
        row, col = divmod(idx, 3)
        bx = 52 + col * (box_w + box_gap)
        by = y - row * (box_h + 12)
        c.setFillColor(HexColor('#ffffff'))
        c.roundRect(bx, by - box_h, box_w, box_h, 6, fill=1, stroke=0)
        c.setStrokeColor(HexColor('#d6e5ef'))
        c.roundRect(bx, by - box_h, box_w, box_h, 6, fill=0, stroke=1)
        c.setFillColor(HexColor('#00a8cc'))
        value_text = str(value)
        value_size = 18 if len(value_text) <= 10 else 13 if len(value_text) <= 20 else 10
        c.setFont("Helvetica-Bold", value_size)
        c.drawCentredString(bx + box_w/2, by - 25, value_text)
        c.setFillColor(HexColor('#455a64'))
        c.setFont("Helvetica", 7.5)
        c.drawCentredString(bx + box_w/2, by - 43, label)

    y -= 2 * (box_h + 12) + 8
    c.setFillColor(HexColor('#0b1829'))
    c.roundRect(52, y - 78, W - 104, 76, 7, fill=1, stroke=0)
    c.setFillColor(HexColor('#00bcd4'))
    c.setFont("Helvetica-Bold", 10)
    c.drawString(60, y - 22, "Interpretation")
    c.setFillColor(HexColor('#d6e5ef'))
    c.setFont("Helvetica", 8.5)
    c.drawString(60, y - 39, "The displayed final score is a unified screening signal and is not proof of authorship.")
    c.drawString(60, y - 54, "Highlighted passages combine internal signals with secure AI review flags and require human review.")
    c.drawString(60, y - 69, "The internal analytical method is intentionally not disclosed in the user-facing report.")

    c.setFillColor(HexColor('#060d1a'))
    c.rect(0, 0, W, 26, fill=1, stroke=0)
    c.setFillColor(HexColor('#f4a300'))
    c.rect(0, 25, W, 1, fill=1, stroke=0)
    c.setFillColor(HexColor('#8aa0b8'))
    c.setFont("Helvetica", 7)
    c.drawCentredString(W/2, 8, "Document Review Report — Confidential")


def _draw_gemini_review_page(c, W, H, result, doc_name, doc_date):
    """English PDF page for the structured Gemini review."""
    review = result.get("gemini_review") or {}
    c.setFillColor(HexColor('#f8fbff'))
    c.rect(0, 0, W, H, fill=1, stroke=0)
    c.setFillColor(HexColor('#060d1a'))
    c.rect(0, H - 44, W, 44, fill=1, stroke=0)
    c.setFillColor(HexColor('#7c3aed'))
    c.rect(0, H - 45.5, W, 1.5, fill=1, stroke=0)
    c.setFillColor(HexColor('#ffffff'))
    c.setFont('Helvetica-Bold', 14)
    c.drawString(40, H - 28, 'AI-Assisted Academic Review')
    c.setFillColor(HexColor('#9db4c4'))
    c.setFont('Helvetica', 7.5)
    c.drawRightString(W - 40, H - 27, "Secure AI review")

    y = H - 78
    ai_score = review.get('ai_assistance_likelihood_score', 0)
    score = review.get('review_priority_score', 0)
    citation = review.get('citation_integrity_risk_score', 0)
    confidence = str(review.get('confidence', 'low')).title()
    box_w = (W - 116) / 4
    for i, (label, value, color) in enumerate([
        ('AI-assistance estimate', _format_percentage(ai_score), '#c0392b'),
        ('Review priority', f'{score}/100', '#7c3aed'),
        ('Citation-support risk', f'{citation}/100', '#f4a300'),
        ('Model confidence', confidence, '#00a8cc'),
    ]):
        x = 40 + i * (box_w + 12)
        c.setFillColor(HexColor('#ffffff'))
        c.roundRect(x, y - 55, box_w, 55, 6, fill=1, stroke=0)
        c.setStrokeColor(HexColor(color))
        c.roundRect(x, y - 55, box_w, 55, 6, fill=0, stroke=1)
        c.setFillColor(HexColor(color))
        c.setFont('Helvetica-Bold', 15)
        c.drawCentredString(x + box_w / 2, y - 24, value)
        c.setFillColor(HexColor('#455a64'))
        c.setFont('Helvetica', 7.5)
        c.drawCentredString(x + box_w / 2, y - 43, label)
    y -= 82

    def draw_wrapped(title, text, y_pos, bullet=False, max_lines=9):
        c.setFillColor(HexColor('#12263a'))
        c.setFont('Helvetica-Bold', 10)
        c.drawString(40, y_pos, title)
        y_pos -= 15
        c.setFillColor(HexColor('#334e68'))
        c.setFont('Helvetica', 8.2)
        lines = []
        if isinstance(text, list):
            for item in text:
                prefix = '• ' if bullet else ''
                lines.extend(textwrap.wrap(prefix + str(item), width=92) or [''])
        else:
            lines = textwrap.wrap(str(text or ''), width=96) or ['']
        for line in lines[:max_lines]:
            c.drawString(48, y_pos, line)
            y_pos -= 11
        return y_pos - 8

    y = draw_wrapped('Assessment summary', review.get('summary_en', ''), y, max_lines=9)
    y = draw_wrapped('Key reasons', review.get('key_reasons_en', []), y, bullet=True, max_lines=8)
    y = draw_wrapped('Recommended actions', review.get('recommended_actions_en', []), y, bullet=True, max_lines=6)

    flags = review.get('flagged_passages') or []
    if y > 120 and flags:
        c.setFillColor(HexColor('#12263a'))
        c.setFont('Helvetica-Bold', 10)
        c.drawString(40, y, 'Flagged passages')
        y -= 16
        c.setFont('Helvetica', 7.8)
        c.setFillColor(HexColor('#334e68'))
        for item in flags[:4]:
            line = f"{item.get('passage_id')} | severity {item.get('severity')}/5 | {item.get('reason_en', '')}"
            for wrapped in textwrap.wrap(line, width=96)[:2]:
                c.drawString(48, y, wrapped)
                y -= 10
            y -= 3
            if y < 70:
                break

    c.setFillColor(HexColor('#fff8e1'))
    c.roundRect(40, 42, W - 80, 34, 5, fill=1, stroke=0)
    c.setStrokeColor(HexColor('#f4a300'))
    c.roundRect(40, 42, W - 80, 34, 5, fill=0, stroke=1)
    c.setFillColor(HexColor('#6d4c00'))
    c.setFont('Helvetica', 7.5)
    c.drawString(48, 61, 'This is a triage review, not proof of AI authorship and not a plagiarism/similarity percentage.')
    c.drawString(48, 50, 'Only selected excerpts were used by the secure AI review; references were excluded before review.')




def generate_report_pdf_text_only(text: str, doc_name: str = "Document",
                                  analysis_result: dict | None = None) -> bytes:
    """
    نسخة احتياطية: عندما لا يكون هناك ملف PDF أصلي.
    تُنشئ تقريراً نصياً مع تظليل الفقرات المشبوهة.
    المراجع لا تُظلَّل.
    """
    if not RLAB_OK:
        return b""

    buf    = io.BytesIO()
    result = analysis_result or _ai_fingerprint_score(text)
    _combined_report = result.get('combined_analysis') or {}
    pct = float(_combined_report.get('score', result['percentage']))
    col_h = _combined_report.get('color', result['color'])
    allow_highlighting = _highlighting_allowed(pct)
    result['highlight_suppressed_low_score'] = not allow_highlighting
    result['highlight_target_words'] = 0 if not allow_highlighting else None
    result['highlight_covered_words'] = 0 if not allow_highlighting else None
    result['highlight_total_words'] = result.get('n_words', 0)
    result['highlight_achieved_pct'] = 0.0 if not allow_highlighting else None
    result['highlight_selected_blocks'] = 0 if not allow_highlighting else None
    today  = datetime.now().strftime("%Y-%m-%d  %H:%M")

    W, H   = A4
    ML, MR = 52, 52
    MT, MB = 45, 40
    TW     = W - ML - MR
    LH     = 13.5
    PS     = 9.0

    c = rl_canvas.Canvas(buf, pagesize=A4)

    # غلاف
    _draw_cover_rl(c, W, H, result, doc_name, today)
    c.showPage()

    # صفحات المحتوى
    page_num = 2
    y = H - MT - 36

    # رأس الصفحة
    _rl_page_header(c, W, H, page_num, doc_name, pct)

    # legend
    c.setFillColor(HexColor('#00bcd4'))
    c.setFont("Helvetica-Bold", 12)
    c.drawString(ML, y, "Document Content Analysis")
    y -= 7
    c.setStrokeColor(HexColor('#1e3a5f'))
    c.setLineWidth(0.7)
    c.line(ML, y, W - MR, y)
    y -= 18

    lx = ML
    for lcol, border, ltxt in [
        ('#fff9c4', '#f4a300', '   AI-Suspicious (orange highlight)'),
        ('#f8fafc', '#dce8f5', '   Normal / Human text'),
        ('#f3e5f5', '#c897d8', '   References — not analyzed'),
    ]:
        c.setFillColor(HexColor(lcol))
        c.roundRect(lx, y - 3, 14, 10, 2, fill=1, stroke=0)
        c.setStrokeColor(HexColor(border))
        c.setLineWidth(0.4)
        c.roundRect(lx, y - 3, 14, 10, 2, fill=0, stroke=1)
        c.setFillColor(HexColor('#333333'))
        c.setFont("Helvetica", 8)
        c.drawString(lx + 17, y + 3, ltxt)
        lx += 158
    y -= 22

    # فقرات النص
    paras_raw = [p.strip() for p in text.split('\n') if p.strip()]
    in_refs = False
    para_groups = []
    for para in paras_raw:
        if _is_reference_header(para):
            in_refs = True
        is_ref  = in_refs or _is_reference_line(para)
        is_susp = allow_highlighting and (not is_ref) and _para_is_suspicious(para, result)
        para_groups.append((para, is_susp, is_ref))

    _report_total_words = sum(_count_text_words(pt) for pt, _, is_ref in para_groups if not is_ref)
    _report_highlight_words = sum(_count_text_words(pt) for pt, is_susp, is_ref in para_groups if is_susp and not is_ref)
    result['highlight_total_words'] = _report_total_words
    result['highlight_covered_words'] = _report_highlight_words
    result['highlight_target_words'] = _report_highlight_words
    result['highlight_achieved_pct'] = round(
        100.0 * _report_highlight_words / max(_report_total_words, 1), 1
    )
    result['highlight_selected_blocks'] = sum(1 for _, is_susp, is_ref in para_groups if is_susp and not is_ref)

    def new_page_rl():
        nonlocal page_num, y
        _rl_page_footer(c, W, page_num)
        c.showPage()
        page_num += 1
        y = H - MT - 40
        _rl_page_header(c, W, H, page_num, doc_name, pct)

    def draw_para_rl(txt, is_susp, is_ref):
        nonlocal y

        if is_ref:
            bg, border, tc, fn, fs = '#f3e5f5', '#c897d8', '#6a1b9a', "Helvetica-Oblique", 9.0
        elif is_susp:
            bg, border, tc, fn, fs = '#fff9c4', '#f4a300', '#5d3a00', "Helvetica", 10.5
        else:
            bg, border, tc, fn, fs = '#f8fafc', '#dce8f5', '#1a1a2a', "Helvetica", 10.5

        avg_char_w = fs * 0.52
        max_c = max(10, int(TW / avg_char_w))
        words_ = txt.split()
        lines  = []
        cur    = ""
        for w in words_:
            test = (cur + " " + w).strip()
            if len(test) <= max_c:
                cur = test
            else:
                if cur: lines.append(cur)
                cur = w
        if cur: lines.append(cur)

        bh = len(lines) * LH + 8

        if y - bh < MB + 26:
            new_page_rl()

        # خلفية
        c.setFillColor(HexColor(bg))
        c.roundRect(ML - 5, y - bh + 4, TW + 10, bh + 2, 4, fill=1, stroke=0)
        if not is_ref:
            c.setStrokeColor(HexColor(border))
            c.setLineWidth(0.3)
            c.roundRect(ML - 5, y - bh + 4, TW + 10, bh + 2, 4, fill=0, stroke=1)

        # شريط جانبي للمشبوه — مثل تظليل Turnitin
        if is_susp:
            c.setFillColor(HexColor('#f4a300'))
            c.rect(ML - 5, y - bh + 4, 4, bh + 2, fill=1, stroke=0)

        # النص
        c.setFillColor(HexColor(tc))
        c.setFont(fn, fs)
        ty = y
        for line in lines:
            if ty < MB + 26:
                new_page_rl()
                ty = y
            c.drawString(ML + 6, ty, line)
            ty -= LH

        # علامة AI
        if is_susp:
            tag_y = y + 1
            c.setFillColor(HexColor('#f4a300'))
            c.roundRect(W - MR - 28, tag_y - 1, 28, 12, 3, fill=1, stroke=0)
            c.setFillColor(HexColor('#ffffff'))
            c.setFont("Helvetica-Bold", 6.5)
            c.drawCentredString(W - MR - 14, tag_y + 4, "AI ✓")

        y = ty - PS

    for pt, ps_, pr in para_groups:
        draw_para_rl(pt, ps_, pr)

    # صفحة ملخص
    _rl_page_footer(c, W, page_num)
    c.showPage()
    page_num += 1
    _draw_summary_page(c, W, H, result, doc_name, today)
    if (result.get("gemini_review") or {}).get("available"):
        c.showPage()
        _draw_gemini_review_page(c, W, H, result, doc_name, today)
    c.save()
    return buf.getvalue()


def _rl_page_header(c, W, H, page_num, doc_name, pct):
    c.setFillColor(HexColor('#060d1a'))
    c.rect(0, H - 30, W, 30, fill=1, stroke=0)
    c.setFillColor(HexColor('#00bcd4'))
    c.rect(0, H - 31, W, 1.5, fill=1, stroke=0)
    c.setFillColor(HexColor('#aaaaaa'))
    c.setFont("Helvetica", 7.5)
    c.drawString(18, H - 19, "AI Fingerprint Detector")
    c.drawCentredString(W/2, H - 19, doc_name[:42])
    c.drawRightString(W - 18, H - 19, f"AI: {_format_percentage(pct)}  |  p.{page_num}")


def _rl_page_footer(c, W, page_num):
    c.setFillColor(HexColor('#060d1a'))
    c.rect(0, 0, W, 24, fill=1, stroke=0)
    c.setFillColor(HexColor('#f4a300'))
    c.rect(0, 23, W, 1, fill=1, stroke=0)
    c.setFillColor(HexColor('#555'))
    c.setFont("Helvetica", 7)
    c.drawCentredString(W/2, 7, f"Page {page_num}  —  AI Fingerprint Detector  —  Confidential")


# ══════════════════════════════════════════════════════════════════════════════
# HELPERS: قراءة الملفات
# ══════════════════════════════════════════════════════════════════════════════

def extract_text_from_pdf(file_bytes: bytes) -> str:
    if not FITZ_OK: return ""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        return "\n".join(page.get_text() for page in doc)
    except Exception:
        return ""


def extract_text_from_docx(file_bytes: bytes) -> str:
    if not DOCX_OK:
        return ""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        return "\n".join(
            p.text for p in _iter_docx_paragraphs(doc) if (p.text or "").strip()
        )
    except Exception:
        return ""



def _iter_docx_paragraphs(container):
    """
    يمر على فقرات Word والجداول بترتيب ظهورها الحقيقي في المستند.
    هذا يمنع انتقال حالة "داخل المراجع" إلى أجزاء سابقة بسبب ترتيب غير صحيح.
    """
    try:
        if isinstance(container, _Document):
            parent_element = container.element.body
            parent = container
        elif isinstance(container, _Cell):
            parent_element = container._tc
            parent = container
        else:
            for p in getattr(container, "paragraphs", []):
                yield p
            return

        for child in parent_element.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                table = Table(child, parent)
                seen_cells = set()
                for row in table.rows:
                    for cell in row.cells:
                        # الخلايا المدمجة قد تظهر أكثر من مرة؛ لا نكرر نصها.
                        # احتفظ بعنصر XML نفسه، لا بـ id رقمي قد تعيد بايثون استخدامه
                        # أثناء إنشاء wrappers للخلايا؛ إعادة استخدام id كانت تُسقط خلايا صحيحة.
                        cell_key = cell._tc
                        if cell_key in seen_cells:
                            continue
                        seen_cells.add(cell_key)
                        yield from _iter_docx_paragraphs(cell)
    except Exception:
        return


def _set_run_docx_shading(run, fill: str = "9FE2FF"):
    """
    يضيف تظليل خلفية حقيقي على run داخل DOCX
    للحفاظ على النص الأصلي وتنسيقه بالكامل.
    """
    try:
        r_pr = run._r.get_or_add_rPr()
        shd = r_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            r_pr.append(shd)
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
    except Exception:
        pass


def generate_highlighted_docx_bytes(docx_bytes: bytes,
                                    text: str = "",
                                    doc_name: str = "Document",
                                    analysis_result: dict | None = None) -> bytes:
    """
    يطبّق نفس مبدأ PDF لكن مباشرة على ملف Word الأصلي:
    - لا يعيد كتابة التقرير كنص جديد
    - يحتفظ بالتنسيق والجداول والأشكال كما هي
    - يظلل الفقرات المشبوهة داخل DOCX نفسه
    - يستثني قسم المراجع من التظليل
    """
    if not DOCX_OK:
        return b""

    # احتياطياً: إذا كانت النتيجة *% نعيد ملف Word الأصلي بلا أي تعديل.
    if text:
        result = analysis_result or _ai_fingerprint_score(text)
        combined = result.get('combined_analysis') or {}
        if not _highlighting_allowed(combined.get('score', result.get('percentage', 0))):
            return docx_bytes
    else:
        result = analysis_result or {}

    try:
        doc = docx.Document(io.BytesIO(docx_bytes))
    except Exception:
        return b""

    in_refs = False
    highlighted = 0

    try:
        for para in _iter_docx_paragraphs(doc):
            para_text = (para.text or "").strip()
            if not para_text:
                continue

            if _is_reference_header(para_text):
                in_refs = True

            is_ref = in_refs or _is_reference_line(para_text)
            is_susp = (not is_ref) and _para_is_suspicious(para_text, result)

            if not is_susp:
                continue

            # نطبّق التظليل على كل run موجود للحفاظ على تنسيق كل جزء كما هو
            non_empty_run_found = False
            for run in para.runs:
                if (run.text or "").strip():
                    _set_run_docx_shading(run, fill="9FE2FF")
                    non_empty_run_found = True

            # إذا كانت الفقرة بلا runs واضحة، ننشئ run أخيراً بدون المساس بالمحتوى
            if not non_empty_run_found and para_text:
                run = para.add_run("")
                _set_run_docx_shading(run, fill="9FE2FF")

            highlighted += 1
    except Exception:
        return b""

    out = io.BytesIO()
    try:
        doc.save(out)
        return out.getvalue()
    except Exception:
        return b""



_DOCX_CONVERSION_DIAGNOSTIC = {
    "ok": False,
    "engine": "",
    "message": "لم تبدأ محاولة التحويل بعد.",
    "details": "",
}


def _set_docx_conversion_diagnostic(ok: bool, engine: str, message: str, details: str = "") -> None:
    """يحفظ آخر حالة تحويل Word حتى تعرض الواجهة سبب النجاح أو الفشل الحقيقي."""
    global _DOCX_CONVERSION_DIAGNOSTIC
    _DOCX_CONVERSION_DIAGNOSTIC = {
        "ok": bool(ok),
        "engine": str(engine or ""),
        "message": str(message or ""),
        "details": str(details or "")[:5000],
    }


def get_last_docx_conversion_diagnostic() -> dict:
    """يعيد نسخة آمنة من آخر تشخيص لتحويل Word إلى PDF."""
    return dict(_DOCX_CONVERSION_DIAGNOSTIC)


def convert_docx_to_pdf_bytes(docx_bytes: bytes) -> bytes:
    """
    يحوّل DOCX إلى PDF مع الحفاظ على تخطيط المستند قدر الإمكان.

    ترتيب المحركات:
    1) Microsoft Word عبر docx2pdf على Windows/macOS.
    2) Microsoft Word COM على Windows.
    3) LibreOffice Writer على Linux وStreamlit Community Cloud.

    ملاحظة نشر مهمة:
    - requirements.txt يثبت مكتبات Python فقط.
    - على Streamlit Community Cloud يجب تثبيت LibreOffice كاعتماد نظام
      من خلال packages.txt في جذر المستودع.
    """
    import os
    import platform
    import shutil
    import socket
    import subprocess
    import tempfile
    from pathlib import Path

    _set_docx_conversion_diagnostic(False, "", "بدأت محاولة تحويل Word إلى PDF.")

    if not docx_bytes or len(docx_bytes) < 100:
        _set_docx_conversion_diagnostic(
            False,
            "input",
            "ملف Word فارغ أو غير صالح.",
        )
        return b""

    def _read_valid_pdf(pdf_path: str) -> bytes:
        try:
            if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 100:
                with open(pdf_path, "rb") as f:
                    data = f.read()
                if data.startswith(b"%PDF"):
                    return data
        except Exception:
            pass
        return b""

    with tempfile.TemporaryDirectory(prefix="ai_detector_docx_") as tmpdir:
        docx_path = os.path.join(tmpdir, "input.docx")
        expected_pdf_path = os.path.join(tmpdir, "input.pdf")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)

        system_name = platform.system()
        errors = []

        # المسار الأدق على الأجهزة التي تحتوي Microsoft Word.
        if system_name in {"Windows", "Darwin"}:
            try:
                from docx2pdf import convert as docx2pdf_convert

                docx2pdf_convert(docx_path, expected_pdf_path)
                data = _read_valid_pdf(expected_pdf_path)
                if data:
                    _set_docx_conversion_diagnostic(
                        True,
                        "Microsoft Word / docx2pdf",
                        "تم تصدير Word إلى PDF بواسطة Microsoft Word.",
                    )
                    return data
                errors.append("docx2pdf انتهى دون إنشاء PDF صالح.")
            except Exception as exc:
                errors.append(f"docx2pdf: {type(exc).__name__}: {exc}")

        # Word COM على Windows.
        if os.name == "nt":
            word_app = None
            word_doc = None
            pythoncom_initialized = False
            try:
                import pythoncom
                import win32com.client

                pythoncom.CoInitialize()
                pythoncom_initialized = True
                word_app = win32com.client.DispatchEx("Word.Application")
                word_app.Visible = False
                word_app.DisplayAlerts = 0
                word_doc = word_app.Documents.Open(
                    os.path.abspath(docx_path),
                    ConfirmConversions=False,
                    ReadOnly=True,
                    AddToRecentFiles=False,
                )
                word_doc.ExportAsFixedFormat(
                    OutputFileName=os.path.abspath(expected_pdf_path),
                    ExportFormat=17,
                    OpenAfterExport=False,
                    OptimizeFor=0,
                    Range=0,
                    Item=0,
                    IncludeDocProps=True,
                    KeepIRM=True,
                    CreateBookmarks=1,
                    DocStructureTags=True,
                    BitmapMissingFonts=True,
                    UseISO19005_1=False,
                )
                data = _read_valid_pdf(expected_pdf_path)
                if data:
                    _set_docx_conversion_diagnostic(
                        True,
                        "Microsoft Word COM",
                        "تم تصدير Word إلى PDF بواسطة Microsoft Word.",
                    )
                    return data
                errors.append("Word COM انتهى دون إنشاء PDF صالح.")
            except Exception as exc:
                errors.append(f"Word COM: {type(exc).__name__}: {exc}")
            finally:
                try:
                    if word_doc is not None:
                        word_doc.Close(False)
                except Exception:
                    pass
                try:
                    if word_app is not None:
                        word_app.Quit()
                except Exception:
                    pass
                if pythoncom_initialized:
                    try:
                        pythoncom.CoUninitialize()
                    except Exception:
                        pass

        # LibreOffice: المحرك المطلوب على Streamlit Community Cloud/Linux.
        env = os.environ.copy()
        env["TMPDIR"] = tmpdir
        env["HOME"] = tmpdir
        env["XDG_RUNTIME_DIR"] = tmpdir
        env["SAL_USE_VCLPLUGIN"] = "svp"
        env.setdefault("LANG", "C.UTF-8")
        env.setdefault("LC_ALL", "C.UTF-8")

        # حل احتياطي لبيئات Linux المقيدة التي تمنع AF_UNIX.
        af_unix = getattr(socket, "AF_UNIX", None)
        if af_unix is not None and os.name != "nt":
            try:
                sock = socket.socket(af_unix, socket.SOCK_STREAM)
                sock.close()
            except OSError:
                shim_path = "/tmp/lo_socket_shim.so"
                if os.path.exists(shim_path):
                    env["LD_PRELOAD"] = shim_path

        candidates = [shutil.which("soffice"), shutil.which("libreoffice")]
        if os.name == "nt":
            candidates.extend([
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            ])
        elif system_name == "Darwin":
            candidates.append("/Applications/LibreOffice.app/Contents/MacOS/soffice")
        else:
            candidates.extend([
                "/usr/bin/soffice",
                "/usr/bin/libreoffice",
                "/usr/local/bin/soffice",
                "/usr/local/bin/libreoffice",
            ])

        soffice_bin = next((c for c in candidates if c and os.path.exists(c)), None)
        if not soffice_bin:
            details = "\n".join(errors)
            _set_docx_conversion_diagnostic(
                False,
                "LibreOffice",
                "LibreOffice غير مثبت في بيئة الاستضافة. أضف ملف packages.txt إلى جذر مستودع GitHub ثم أعد تشغيل التطبيق.",
                details,
            )
            return b""

        # تأكد أن الملف التنفيذي يعمل، وهذا يعطي رسالة أوضح من فشل صامت.
        try:
            version_proc = subprocess.run(
                [soffice_bin, "--version"],
                env=env,
                capture_output=True,
                timeout=30,
                text=True,
            )
            version_text = (version_proc.stdout or version_proc.stderr or "").strip()
            if version_proc.returncode != 0:
                errors.append(
                    f"LibreOffice --version returncode={version_proc.returncode}: {version_text}"
                )
        except Exception as exc:
            version_text = ""
            errors.append(f"LibreOffice version check: {type(exc).__name__}: {exc}")

        profile_dir = Path(tmpdir, "lo_profile")
        profile_dir.mkdir(parents=True, exist_ok=True)
        cmd = [
            soffice_bin,
            f"-env:UserInstallation={profile_dir.as_uri()}",
            "--headless",
            "--invisible",
            "--nologo",
            "--nodefault",
            "--nolockcheck",
            "--norestore",
            "--nofirststartwizard",
            "--convert-to", "pdf:writer_pdf_Export",
            "--outdir", tmpdir,
            docx_path,
        ]

        try:
            proc = subprocess.run(
                cmd,
                env=env,
                capture_output=True,
                timeout=240,
                text=True,
            )
        except subprocess.TimeoutExpired as exc:
            _set_docx_conversion_diagnostic(
                False,
                "LibreOffice",
                "انتهت مهلة تحويل Word إلى PDF داخل بيئة الاستضافة.",
                f"timeout={exc.timeout}; errors={' | '.join(errors)}",
            )
            return b""
        except Exception as exc:
            _set_docx_conversion_diagnostic(
                False,
                "LibreOffice",
                "تعذر تشغيل LibreOffice داخل بيئة الاستضافة.",
                f"{type(exc).__name__}: {exc}\n" + "\n".join(errors),
            )
            return b""

        # LibreOffice قد يعيد تسمية الملف؛ لذلك نبحث عن أي PDF ناتج.
        pdf_candidates = [Path(expected_pdf_path)] + sorted(Path(tmpdir).glob("*.pdf"))
        for candidate in pdf_candidates:
            data = _read_valid_pdf(str(candidate))
            if data:
                details = "\n".join(
                    part for part in [
                        version_text,
                        (proc.stdout or "").strip(),
                        (proc.stderr or "").strip(),
                    ] if part
                )
                _set_docx_conversion_diagnostic(
                    True,
                    "LibreOffice Writer",
                    "تم تصدير Word إلى PDF بواسطة LibreOffice Writer.",
                    details,
                )
                return data

        details = "\n".join(
            part for part in [
                f"returncode={proc.returncode}",
                (proc.stdout or "").strip(),
                (proc.stderr or "").strip(),
                "\n".join(errors),
            ] if part
        )
        _set_docx_conversion_diagnostic(
            False,
            "LibreOffice",
            "عمل LibreOffice لكنه لم يُنشئ ملف PDF صالحًا.",
            details,
        )
        return b""


def _build_logo_data_uri() -> str:
    svg = """<svg xmlns='http://www.w3.org/2000/svg' width='220' height='220' viewBox='0 0 220 220'>
      <defs>
        <linearGradient id='lg1' x1='0%' y1='0%' x2='100%' y2='100%'>
          <stop offset='0%' stop-color='#08150f'/>
          <stop offset='100%' stop-color='#0d3324'/>
        </linearGradient>
        <linearGradient id='lg2' x1='0%' y1='0%' x2='100%' y2='100%'>
          <stop offset='0%' stop-color='#22e5a8'/>
          <stop offset='50%' stop-color='#ffb703'/>
          <stop offset='100%' stop-color='#ff6b4a'/>
        </linearGradient>
        <linearGradient id='lg3' x1='0%' y1='0%' x2='100%' y2='100%'>
          <stop offset='0%' stop-color='#22e5a8' stop-opacity='0.3'/>
          <stop offset='100%' stop-color='#ffb703' stop-opacity='0.1'/>
        </linearGradient>
        <filter id='glow'>
          <feGaussianBlur stdDeviation='3' result='coloredBlur'/>
          <feMerge><feMergeNode in='coloredBlur'/><feMergeNode in='SourceGraphic'/></feMerge>
        </filter>
      </defs>
      <!-- خلفية -->
      <rect x='0' y='0' width='220' height='220' rx='44' fill='url(#lg1)'/>
      <!-- حلقات توهج خلفية -->
      <circle cx='110' cy='110' r='90' fill='none' stroke='url(#lg2)' stroke-width='0.5' opacity='0.2'/>
      <circle cx='110' cy='110' r='72' fill='none' stroke='url(#lg2)' stroke-width='0.5' opacity='0.3'/>
      <!-- شبكة نقاط خفية -->
      <g opacity='0.12' fill='#22e5a8'>
        <circle cx='40' cy='40' r='1.5'/><circle cx='70' cy='40' r='1.5'/><circle cx='100' cy='40' r='1.5'/>
        <circle cx='130' cy='40' r='1.5'/><circle cx='160' cy='40' r='1.5'/><circle cx='180' cy='40' r='1.5'/>
        <circle cx='40' cy='70' r='1.5'/><circle cx='180' cy='70' r='1.5'/>
        <circle cx='40' cy='180' r='1.5'/><circle cx='70' cy='180' r='1.5'/><circle cx='130' cy='180' r='1.5'/>
        <circle cx='160' cy='180' r='1.5'/><circle cx='180' cy='180' r='1.5'/>
      </g>
      <!-- دائرة رئيسية متوهجة -->
      <circle cx='110' cy='105' r='52' fill='url(#lg3)' filter='url(#glow)'/>
      <circle cx='110' cy='105' r='44' fill='none' stroke='url(#lg2)' stroke-width='2.5' opacity='0.9'/>
      <!-- خطوط بصمة إصبع -->
      <g fill='none' stroke='url(#lg2)' stroke-linecap='round' opacity='0.7'>
        <path d='M90 105 Q110 88 130 105 Q110 122 90 105' stroke-width='2'/>
        <path d='M83 105 Q110 80 137 105 Q110 130 83 105' stroke-width='1.5'/>
        <path d='M76 105 Q110 72 144 105 Q110 138 76 105' stroke-width='1.2'/>
        <path d='M70 105 Q110 65 150 105 Q110 145 70 105' stroke-width='1'/>
      </g>
      <!-- نقطة مركزية -->
      <circle cx='110' cy='105' r='6' fill='#22e5a8' filter='url(#glow)'/>
      <circle cx='110' cy='105' r='3' fill='#ffffff'/>
      <!-- خطوط ماسحة رادار -->
      <line x1='110' y1='105' x2='110' y2='65' stroke='#22e5a8' stroke-width='1.5' opacity='0.6'/>
      <line x1='110' y1='105' x2='140' y2='80' stroke='#ffb703' stroke-width='1' opacity='0.5'/>
      <!-- نص AI -->
      <text x='110' y='168' text-anchor='middle' font-size='13' font-family='Arial Black,Arial' font-weight='900'
            letter-spacing='4' fill='url(#lg2)' opacity='0.9'>AI·FP</text>
      <!-- إطار خارجي فاخر -->
      <rect x='3' y='3' width='214' height='214' rx='42' fill='none' stroke='url(#lg2)' stroke-width='1.5' opacity='0.4'/>
    </svg>"""
    return "data:image/svg+xml;base64," + base64.b64encode(svg.encode("utf-8")).decode("ascii")


logo_data_uri = _build_logo_data_uri()

st.markdown(f"""
<style>
@import url('https://fonts.googleapis.com/css2?family=Cairo:wght@400;600;700;900&family=Space+Grotesk:wght@400;600;700&display=swap');

* {{ box-sizing: border-box; }}

body,
[data-testid="stAppViewContainer"],
[data-testid="stVerticalBlock"],
.main {{
    direction: rtl;
    font-family: 'Cairo', 'Segoe UI', Tahoma, sans-serif;
}}

pre, code, .stat-num, .layer-val, .ltr {{
    direction: ltr !important;
    text-align: left !important;
    font-family: 'Space Grotesk', 'Courier New', monospace;
}}

[data-testid="stAppViewContainer"] {{
    background:
        radial-gradient(ellipse 80% 60% at 10% 0%, rgba(34,229,168,0.07) 0%, transparent 60%),
        radial-gradient(ellipse 60% 70% at 90% 100%, rgba(255,183,3,0.09) 0%, transparent 60%),
        radial-gradient(ellipse 50% 50% at 50% 50%, rgba(255,107,74,0.04) 0%, transparent 70%),
        linear-gradient(160deg, #071410 0%, #0a1f18 40%, #10190f 100%);
    min-height: 100vh;
}}

[data-testid="stHeader"] {{
    background: rgba(7,9,26,0.85) !important;
    backdrop-filter: blur(20px);
    border-bottom: 1px solid rgba(34,229,168,0.12);
}}

.main .block-container {{
    max-width: 980px;
    padding-top: 2rem;
    padding-bottom: 4rem;
    padding-right: 1.4rem;
    padding-left: 1.4rem;
}}

.hero-card {{
    background: linear-gradient(145deg, rgba(13,16,42,0.95), rgba(10,14,35,0.98));
    border: 1px solid rgba(34,229,168,0.18);
    border-radius: 28px;
    padding: 1.8rem 2rem 1.5rem;
    box-shadow: 0 0 0 1px rgba(255,183,3,0.08) inset, 0 30px 80px rgba(0,0,0,0.5);
    margin-bottom: 1.4rem;
    position: relative;
    overflow: hidden;
}}

.hero-card::before {{
    content: '';
    position: absolute;
    top: -80px; right: -80px;
    width: 260px; height: 260px;
    background: radial-gradient(circle, rgba(34,229,168,0.07), transparent 70%);
    pointer-events: none;
}}

.hero-card::after {{
    content: '';
    position: absolute;
    bottom: -60px; left: -60px;
    width: 200px; height: 200px;
    background: radial-gradient(circle, rgba(255,183,3,0.06), transparent 70%);
    pointer-events: none;
}}

.hero-card-topbar {{
    position: absolute;
    top: 0; left: 0; right: 0; height: 3px;
    background: linear-gradient(90deg, #22e5a8, #ffb703, #ff6b4a, #22e5a8);
    background-size: 200% 100%;
    animation: shimmer 4s linear infinite;
    border-radius: 28px 28px 0 0;
}}

@keyframes shimmer {{
    0%   {{ background-position: 0% 0%; }}
    100% {{ background-position: 200% 0%; }}
}}

.hero-grid {{
    display: grid;
    grid-template-columns: 110px 1fr;
    gap: 1.4rem;
    align-items: center;
    direction: rtl;
}}

.hero-logo {{
    width: 100px; height: 100px;
    border-radius: 24px;
    display: flex;
    align-items: center;
    justify-content: center;
    background: linear-gradient(145deg, #0b1f16, #08150f);
    border: 1px solid rgba(34,229,168,0.25);
    box-shadow: 0 0 30px rgba(34,229,168,0.15), 0 10px 30px rgba(0,0,0,0.4);
    overflow: hidden;
    flex-shrink: 0;
}}

.hero-logo img {{
    width: 92px; height: 92px;
    object-fit: contain;
}}

.hero-text {{ text-align: right; }}

.hero-title {{
    font-size: 2.1rem;
    font-weight: 900;
    background: linear-gradient(135deg, #22e5a8 0%, #ffb703 50%, #ff6b4a 100%);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
    line-height: 1.1;
    margin-bottom: 0.45rem;
    letter-spacing: -0.5px;
    direction: ltr;
    text-align: left;
}}

.hero-subtitle {{
    color: rgba(180,200,240,0.75);
    font-size: 0.93rem;
    line-height: 1.9;
    direction: rtl;
    text-align: right;
}}

.hero-badges {{
    display: flex;
    flex-wrap: wrap;
    gap: 0.5rem;
    margin-top: 1rem;
    justify-content: flex-end;
    direction: rtl;
}}

.hero-badge {{
    background: rgba(34,229,168,0.07);
    border: 1px solid rgba(34,229,168,0.18);
    color: #7be8c4;
    border-radius: 999px;
    padding: 0.3rem 0.85rem;
    font-size: 0.78rem;
    font-weight: 700;
    white-space: nowrap;
    transition: all 0.2s;
}}

.features-wrap {{
    background: rgba(10,14,35,0.85);
    border: 1px solid rgba(255,183,3,0.15);
    border-radius: 24px;
    padding: 1.4rem;
    box-shadow: 0 20px 50px rgba(0,0,0,0.3);
    margin-bottom: 0.6rem;
    position: relative;
    overflow: hidden;
}}

.features-wrap::before {{
    content: '';
    position: absolute;
    top: 0; left: 0; right: 0; height: 2px;
    background: linear-gradient(90deg, transparent, rgba(255,183,3,0.6), transparent);
}}

.features-title {{
    font-weight: 900;
    color: #cdf0e0;
    margin-bottom: 1rem;
    font-size: 1rem;
    text-align: right;
}}

.features-grid {{
    display: grid;
    grid-template-columns: repeat(2, 1fr);
    gap: 0.85rem;
}}

.feature-card {{
    background: rgba(15,20,50,0.7);
    border: 1px solid rgba(34,229,168,0.10);
    border-radius: 18px;
    padding: 1.1rem;
    transition: transform 0.25s, box-shadow 0.25s, border-color 0.25s;
    text-align: right;
    position: relative;
    overflow: hidden;
}}

.feature-card:hover {{
    transform: translateY(-3px);
    box-shadow: 0 16px 40px rgba(0,0,0,0.4), 0 0 20px rgba(34,229,168,0.07);
    border-color: rgba(34,229,168,0.22);
}}

.feature-icon {{
    width: 40px; height: 40px;
    border-radius: 12px;
    display: inline-flex;
    align-items: center;
    justify-content: center;
    background: linear-gradient(135deg, rgba(34,229,168,0.15), rgba(255,183,3,0.15));
    border: 1px solid rgba(34,229,168,0.15);
    font-size: 1.15rem;
    margin-bottom: 0.55rem;
}}

.feature-card strong {{
    display: block;
    color: #d9f5e6;
    margin-bottom: 0.3rem;
    font-size: 0.92rem;
    font-weight: 800;
}}

.feature-card span {{
    color: rgba(140,165,210,0.7);
    font-size: 0.82rem;
    line-height: 1.75;
}}

.section-card {{
    background: rgba(10,14,35,0.88);
    border: 1px solid rgba(34,229,168,0.14);
    border-radius: 24px;
    padding: 1.4rem;
    box-shadow: 0 20px 50px rgba(0,0,0,0.3);
    position: relative;
}}

.section-card::before {{
    content: '';
    position: absolute;
    top: 0; left: 0; right: 0; height: 2px;
    background: linear-gradient(90deg, transparent, rgba(34,229,168,0.5), transparent);
    border-radius: 24px 24px 0 0;
}}

.meter-wrap {{
    background: rgba(10,14,35,0.9);
    border: 1px solid rgba(34,229,168,0.15);
    border-radius: 24px;
    padding: 2.4rem 2rem 2rem;
    margin: 1.5rem 0;
    text-align: center;
    box-shadow: 0 30px 70px rgba(0,0,0,0.4);
    position: relative;
    overflow: hidden;
}}

.meter-wrap::before {{
    content: '';
    position: absolute;
    top: 0; left: 0; right: 0; height: 2px;
    background: linear-gradient(90deg, transparent, rgba(34,229,168,0.6), rgba(255,183,3,0.6), transparent);
}}

.meter-pct {{
    font-size: 5.5rem;
    font-weight: 900;
    line-height: 1;
    letter-spacing: -4px;
    font-family: 'Space Grotesk', sans-serif;
    direction: ltr;
    display: block;
}}

.meter-label {{
    font-size: 1.2rem;
    margin-top: 0.6rem;
    font-weight: 800;
    direction: rtl;
}}

.meter-human {{
    font-size: 0.9rem;
    color: rgba(140,165,210,0.65);
    margin-top: 0.5rem;
    direction: rtl;
    font-family: 'Space Grotesk', sans-serif;
}}

.bar-track {{
    background: rgba(255,255,255,0.06);
    border-radius: 999px;
    height: 14px;
    margin: 1.5rem 0;
    overflow: hidden;
    border: 1px solid rgba(255,255,255,0.04);
}}

.bar-fill {{
    height: 100%;
    border-radius: 999px;
    transition: width 0.8s cubic-bezier(.4, 0, .2, 1);
}}

.layer-card {{
    background: rgba(12,17,45,0.85);
    border: 1px solid rgba(34,229,168,0.10);
    border-radius: 16px;
    padding: 1rem 1.2rem;
    margin-bottom: 0.75rem;
    direction: rtl;
    text-align: right;
    transition: border-color 0.2s;
    overflow: hidden;
}}

.layer-card:hover {{
    border-color: rgba(34,229,168,0.22);
}}

.layer-title {{
    font-weight: 800;
    font-size: 0.9rem;
    color: #bfe8d4;
    direction: rtl;
    display: inline-block;
}}

.layer-val {{
    float: left;
    font-weight: 900;
    font-size: 0.92rem;
    direction: ltr !important;
    text-align: left !important;
    font-family: 'Space Grotesk', monospace;
}}

.layer-bar {{
    clear: both;
    height: 8px;
    background: rgba(255,255,255,0.06);
    border-radius: 999px;
    margin-top: 0.75rem;
    overflow: hidden;
}}

.layer-bar-fill {{
    height: 100%;
    border-radius: 999px;
}}

.hit-tag {{
    display: inline-block;
    background: rgba(255,107,107,0.08);
    border: 1px solid rgba(255,107,107,0.22);
    border-radius: 8px;
    padding: 4px 12px;
    margin: 3px;
    font-size: 0.77rem;
    font-family: 'Space Grotesk', 'Courier New', monospace;
    color: #ff9999;
    direction: ltr !important;
    text-align: left !important;
    transition: background 0.2s;
}}

.stat-row {{
    display: flex;
    gap: 0.75rem;
    flex-wrap: wrap;
    margin: 1rem 0;
    direction: rtl;
}}

.stat-box {{
    flex: 1;
    min-width: 95px;
    background: rgba(12,17,45,0.85);
    border: 1px solid rgba(34,229,168,0.12);
    border-radius: 16px;
    padding: 0.9rem 0.8rem;
    text-align: center;
    transition: transform 0.2s, border-color 0.2s;
}}

.stat-box:hover {{
    transform: translateY(-2px);
    border-color: rgba(34,229,168,0.28);
}}

.stat-num {{
    font-size: 1.5rem;
    font-weight: 900;
    color: #22e5a8;
    direction: ltr !important;
    text-align: center !important;
    font-family: 'Space Grotesk', sans-serif;
    display: block;
}}

.stat-lbl {{
    font-size: 0.7rem;
    color: rgba(140,165,210,0.6);
    margin-top: 0.25rem;
    direction: rtl;
}}

div[data-testid="stTabs"] {{
    direction: rtl;
}}

div[data-testid="stTabs"] button[role="tab"] {{
    border-radius: 12px !important;
    background: rgba(12,17,45,0.8) !important;
    border: 1px solid rgba(34,229,168,0.12) !important;
    color: rgba(140,180,220,0.8) !important;
    font-weight: 700 !important;
    font-size: 0.9rem !important;
    font-family: 'Cairo', sans-serif !important;
}}

div[data-testid="stTabs"] button[aria-selected="true"] {{
    background: linear-gradient(135deg, rgba(34,229,168,0.15), rgba(255,183,3,0.15)) !important;
    color: #22e5a8 !important;
    border-color: rgba(34,229,168,0.3) !important;
}}

div[data-testid="stTextArea"] textarea {{
    background: rgba(7,9,26,0.8) !important;
    border-radius: 16px !important;
    border: 1.5px solid rgba(34,229,168,0.15) !important;
    color: #d9f5e6 !important;
    direction: rtl;
    text-align: right;
    font-size: 0.95rem !important;
    font-family: 'Cairo', sans-serif !important;
}}

div[data-testid="stTextArea"] textarea::placeholder {{
    text-align: right;
    color: rgba(100,130,180,0.45) !important;
}}

div[data-testid="stFileUploader"] section {{
    background: rgba(7,9,26,0.6) !important;
    border-radius: 16px !important;
    border: 1.5px dashed rgba(34,229,168,0.2) !important;
    color: rgba(140,180,220,0.7) !important;
}}

.stButton > button,
.stDownloadButton > button {{
    border: none !important;
    border-radius: 14px !important;
    font-size: 0.97rem !important;
    font-weight: 800 !important;
    padding: 0.78rem 1.5rem !important;
    width: 100% !important;
    transition: all 0.25s cubic-bezier(.4, 0, .2, 1);
    font-family: 'Cairo', sans-serif !important;
}}

.stButton > button[kind="primary"],
.stDownloadButton > button {{
    background: linear-gradient(135deg, #22e5a8, #ffb703) !important;
    color: white !important;
    box-shadow: 0 8px 25px rgba(34,229,168,0.25), 0 4px 10px rgba(0,0,0,0.3) !important;
}}

.stButton > button[kind="primary"]:hover,
.stDownloadButton > button:hover {{
    transform: translateY(-2px);
    box-shadow: 0 14px 35px rgba(34,229,168,0.35), 0 6px 15px rgba(0,0,0,0.3) !important;
}}

.stButton > button:not([kind="primary"]) {{
    background: rgba(12,17,45,0.85) !important;
    color: rgba(140,180,220,0.85) !important;
    border: 1.5px solid rgba(34,229,168,0.18) !important;
}}

.stButton > button:not([kind="primary"]):hover {{
    background: rgba(34,229,168,0.08) !important;
    border-color: rgba(34,229,168,0.35) !important;
    color: #22e5a8 !important;
    transform: translateY(-1px);
}}

.stAlert,
div[data-testid="stAlert"] {{
    border-radius: 14px !important;
    direction: rtl !important;
    text-align: right !important;
}}

div[data-testid="stAlert"] {{
    background: rgba(34,229,168,0.06) !important;
    border: 1px solid rgba(34,229,168,0.18) !important;
}}

div[data-testid="stAlert"] p {{
    direction: rtl !important;
    text-align: right !important;
    color: rgba(140,200,240,0.85) !important;
}}

h3 {{
    text-align: right !important;
    direction: rtl !important;
    color: #7fd9ae !important;
    font-weight: 800 !important;
    margin: 1.2rem 0 0.6rem !important;
    font-family: 'Cairo', sans-serif !important;
}}

hr {{
    border: none !important;
    border-top: 1px solid rgba(34,229,168,0.12) !important;
    margin: 1.5rem 0 !important;
}}

.info-note {{
    text-align: center;
    font-size: 0.8rem;
    color: rgba(140,165,210,0.5);
    margin-top: 0.5rem;
    direction: rtl;
}}

code {{
    background: rgba(34,229,168,0.08) !important;
    border: 1px solid rgba(34,229,168,0.15) !important;
    color: #22e5a8 !important;
    border-radius: 6px !important;
    padding: 2px 6px !important;
    direction: ltr !important;
    text-align: left !important;
    font-family: 'Space Grotesk', 'Courier New', monospace !important;
}}

[data-testid="stMarkdownContainer"] p,
[data-testid="stMarkdownContainer"] li {{
    color: rgba(160,185,225,0.85);
    direction: rtl;
    text-align: right;
    line-height: 1.85;
}}

[data-testid="stMarkdownContainer"] strong {{
    color: #d9f5e6;
    font-weight: 800;
}}

[data-testid="stSuccess"] {{
    background: rgba(0,200,100,0.08) !important;
    border-color: rgba(0,200,100,0.2) !important;
    border-radius: 14px !important;
    direction: rtl !important;
    text-align: right !important;
}}

[data-testid="stError"] {{
    background: rgba(255,80,80,0.08) !important;
    border-color: rgba(255,80,80,0.2) !important;
    border-radius: 14px !important;
    direction: rtl !important;
    text-align: right !important;
}}

[data-testid="stWarning"] {{
    background: rgba(255,170,0,0.08) !important;
    border-color: rgba(255,170,0,0.2) !important;
    border-radius: 14px !important;
    direction: rtl !important;
    text-align: right !important;
}}

/* ── لافتة التحديثات المتحركة ─────────────────────────────────────────── */
.update-banner {{
    position: relative;
    overflow: hidden;
    border-radius: 16px;
    background: linear-gradient(90deg, rgba(34,229,168,0.16), rgba(255,183,3,0.16), rgba(255,107,74,0.16));
    border: 1px solid rgba(34,229,168,0.3);
    box-shadow: 0 10px 30px rgba(0,0,0,0.35);
    padding: 0.65rem 0;
    margin-bottom: 1.4rem;
}}

.update-banner-track {{
    display: flex;
    width: max-content;
    white-space: nowrap;
    animation: marquee-scroll 22s linear infinite;
}}

.update-banner:hover .update-banner-track {{
    animation-play-state: paused;
}}

.update-banner-track span {{
    display: inline-block;
    padding: 0 2.2rem;
    font-weight: 800;
    font-size: 0.95rem;
    font-family: 'Cairo', sans-serif;
    color: #eafff4;
    direction: rtl;
}}

.update-banner-track span b {{
    background: linear-gradient(135deg, #22e5a8, #ffb703);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
}}

@keyframes marquee-scroll {{
    0%   {{ transform: translateX(0%); }}
    100% {{ transform: translateX(-50%); }}
}}

/* ── تقييمات العملاء ───────────────────────────────────────────────────── */
.ratings-summary {{
    text-align: center;
    padding: 0.4rem 0 1rem;
}}

.ratings-avg {{
    font-size: 3rem;
    font-weight: 900;
    line-height: 1;
    font-family: 'Space Grotesk', sans-serif;
    background: linear-gradient(135deg, #22e5a8, #ffb703);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
    direction: ltr;
}}

.ratings-stars-big {{
    font-size: 1.6rem;
    letter-spacing: 4px;
    color: #ffb703;
    margin-top: 0.3rem;
    direction: ltr;
}}

.ratings-count {{
    font-size: 0.82rem;
    color: rgba(160,220,195,0.65);
    margin-top: 0.3rem;
    direction: rtl;
}}

.review-item {{
    background: rgba(12,45,32,0.55);
    border: 1px solid rgba(34,229,168,0.14);
    border-radius: 14px;
    padding: 0.8rem 1rem;
    margin-bottom: 0.6rem;
    direction: rtl;
    text-align: right;
}}

.review-head {{
    display: flex;
    justify-content: space-between;
    align-items: center;
    margin-bottom: 0.25rem;
}}

.review-name {{
    font-weight: 800;
    color: #d9f5e6;
    font-size: 0.88rem;
}}

.review-stars {{
    color: #ffb703;
    font-size: 0.95rem;
    letter-spacing: 2px;
    direction: ltr;
}}

.review-comment {{
    color: rgba(190,225,210,0.8);
    font-size: 0.85rem;
    line-height: 1.7;
}}

.review-date {{
    color: rgba(150,190,175,0.45);
    font-size: 0.72rem;
    margin-top: 0.3rem;
    direction: ltr;
    text-align: left;
}}

.rate-prompt {{
    font-weight: 800;
    color: #cdf0e0;
    margin: 0.6rem 0 0.4rem;
    text-align: right;
    direction: rtl;
}}

@media (max-width: 720px) {{
    .hero-grid, .features-grid {{
        grid-template-columns: 1fr;
    }}
    .hero-logo {{
        margin: 0 auto;
    }}
    .hero-title {{
        font-size: 1.6rem;
        text-align: right;
    }}
}}
</style>
""", unsafe_allow_html=True)


# ── Gemini service configuration (owner-managed, hidden from users) ───────────
_GEMINI_SERVICE_ENABLED = _gemini_service_enabled()
_GEMINI_MODEL = _configured_gemini_model()
_GEMINI_KEY_AVAILABLE = bool(_get_gemini_api_key())


# ── Hero Section ──────────────────────────────────────────────────────────────
st.markdown(f"""
<div class="hero-card">
    <div class="hero-card-topbar"></div>
    <div class="hero-grid">
        <div class="hero-logo">
            <img src="{logo_data_uri}" alt="AI Fingerprint Detector">
        </div>
        <div class="hero-text">
            <div class="hero-title">AI Fingerprint Detector</div>
            <div class="hero-subtitle">
                فحص أكاديمي مدعوم من الذكاء الاصطناعي مع محرك داخلي متعدد المؤشرات للنصوص الإنجليزية،
                وتقرير مُظلَّل لملفات PDF وWord مع الحفاظ على التنسيق والجداول والأشكال.
            </div>
            <div class="hero-badges">
                <span class="hero-badge">🎯 تظليل دقيق على الملف الأصلي</span>
                <span class="hero-badge">📄 PDF و Word (تنسيق كامل)</span>
                <span class="hero-badge">🚫 المراجع خارج التظليل</span>
                <span class="hero-badge">🤖 مدعوم من الذكاء الاصطناعي</span>
            </div>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)


# ── لافتة تحديثات متحركة ─────────────────────────────────────────────────────
_banner_msg = (
    "✨ <b>الإصدار الاحترافي الجديد:</b> نسبة موحدة للنص المحتمل صياغته بالذكاء الاصطناعي "
    "&nbsp;•&nbsp; 🎯 لا تظليل من 0–20% ويبدأ التظليل من 21% "
    "&nbsp;•&nbsp; ✅ خطة مراجعة ذكية قابلة للتنفيذ في نهاية التقرير "
    "&nbsp;•&nbsp; 📄 تقرير احترافي يحافظ على الملف الأصلي والمراجع"
)
st.markdown(f"""
<div class="update-banner">
    <div class="update-banner-track">
        <span>{_banner_msg}</span>
        <span>{_banner_msg}</span>
    </div>
</div>
""", unsafe_allow_html=True)


# ── تقييمات العملاء ───────────────────────────────────────────────────────────
st.markdown('<div class="section-card">', unsafe_allow_html=True)
st.markdown("### ⭐ تقييمات العملاء")

_ratings_data = _load_ratings()
_avg_rating   = sum(r["stars"] for r in _ratings_data) / len(_ratings_data) if _ratings_data else 0.0
_full_stars   = int(round(_avg_rating))
_stars_disp   = "★" * _full_stars + "☆" * (5 - _full_stars)

st.markdown(f"""
<div class="ratings-summary">
    <span class="ratings-avg">{_avg_rating:.1f}</span>
    <div class="ratings-stars-big">{_stars_disp}</div>
    <div class="ratings-count">بناءً على {len(_ratings_data)} تقييم من عملائنا</div>
</div>
""", unsafe_allow_html=True)

with st.expander("💬 اطّلع على آراء العملاء"):
    for _r in reversed(_ratings_data[-10:]):
        _stxt = "★" * _r["stars"] + "☆" * (5 - _r["stars"])
        _comment = _r.get("comment") or ""
        st.markdown(f"""
        <div class="review-item">
            <div class="review-head">
                <span class="review-name">{_r.get('name', 'عميل')}</span>
                <span class="review-stars">{_stxt}</span>
            </div>
            {f'<div class="review-comment">{_comment}</div>' if _comment else ''}
            <div class="review-date">{_r.get('date', '')}</div>
        </div>
        """, unsafe_allow_html=True)

st.markdown('<div class="rate-prompt">قيّم تجربتك معنا:</div>', unsafe_allow_html=True)
_new_stars = st.feedback("stars", key="rating_feedback")

_rc1, _rc2 = st.columns([2, 2])
with _rc1:
    _reviewer_name = st.text_input(
        "اسمك (اختياري)", key="reviewer_name",
        label_visibility="collapsed", placeholder="اسمك (اختياري)",
    )
with _rc2:
    _reviewer_comment = st.text_input(
        "تعليقك (اختياري)", key="reviewer_comment",
        label_visibility="collapsed", placeholder="تعليقك (اختياري)",
    )

if st.button("إرسال التقييم", use_container_width=True):
    if _new_stars is None:
        st.warning("الرجاء اختيار عدد النجوم أولاً.")
    else:
        _save_rating({
            "name": (_reviewer_name or "").strip() or "عميل",
            "stars": int(_new_stars) + 1,
            "comment": (_reviewer_comment or "").strip(),
            "date": datetime.now().strftime("%Y-%m-%d"),
        })
        st.success("✅ شكراً على تقييمك! تم إضافته إلى آراء العملاء.")
        st.rerun()

st.markdown('</div>', unsafe_allow_html=True)


# ── Features ──────────────────────────────────────────────────────────────────
st.markdown("""
<div class="features-wrap">
    <div class="features-title">ماذا يُقدّم البرنامج؟</div>
    <div class="features-grid">
        <div class="feature-card">
            <div class="feature-icon">🔍</div>
            <strong>مراجعة ذكية متكاملة</strong>
            <span>يحلل المحرك الداخلي النص أولًا، ثم تستخدم المراجعة المدعومة من الذكاء الاصطناعي مقتطفات مختارة بلا مراجع لإصدار تفسير واضح وتحديد المقاطع التي تحتاج تحققًا بشريًا.</span>
        </div>
        <div class="feature-card">
            <div class="feature-icon">📑</div>
            <strong>التقرير المظلل: PDF و Word (تنسيق كامل)</strong>
            <span>عند رفع PDF يتم التظليل فوق الصفحات الأصلية مباشرة. وعند رفع DOCX يتم تصدير الملف الأصلي أولاً إلى PDF بواسطة Microsoft Word عند توفره، ثم يُنشأ عليه التقرير نفسه تماماً مع الحفاظ على الجداول والصور والأشكال والتنسيق.</span>
        </div>
        <div class="feature-card">
            <div class="feature-icon">🚫</div>
            <strong>المراجع محمية من التظليل</strong>
            <span>يستبعد قسم المراجع والاستشهادات الرقمية والمؤلف–السنة من الحساب، ولا يطبّق على قسم المراجع أي تظليل.</span>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)


# ── Input Section ─────────────────────────────────────────────────────────────
st.markdown('<div class="section-card">', unsafe_allow_html=True)

st.markdown("### رفع الملف للفحص")
st.info("📌 يفحص البرنامج النصوص الإنجليزية بتحليل مشترك إلزامي بين المحرك الداخلي ونموذج ذكي ثابت. لا تُعتمد النتيجة ولا يُصدر التقرير ما لم تنجح المرحلتان معًا. تُخفى النسب من 0 إلى 20 بصيغة *% دون تظليل، ويبدأ التظليل وإظهار الرقم من 21% فأعلى.")

st.session_state["smart_review_consent"] = True
_gemini_consent = st.checkbox(
    "المراجعة المدعومة من الذكاء الاصطناعي مفعلة تلقائيًا للمقتطفات المختارة",
    key="smart_review_consent",
    disabled=True,
    help="مفعلة تلقائيًا دون تدخل المستخدم. لا يُرسل الملف كاملًا ولا قسم المراجع.",
)

input_text         = ""
uploaded_name      = "Document"
original_pdf_bytes = None
original_docx_bytes = None

uploaded = st.file_uploader(
    "ارفع ملف PDF أو Word:",
    type=["pdf", "docx"],
    key="file_upload",
)

if uploaded:
    uploaded_name = uploaded.name
    raw = uploaded.read()
    ext = uploaded.name.lower().rsplit(".", 1)[-1]

    if ext == "pdf":
        original_pdf_bytes = raw
        input_text = extract_text_from_pdf(raw)
        if input_text:
            st.success(f"✅ تم استخراج النص من ملف PDF ({len(input_text.split()):,} كلمة)")
        else:
            st.error("تعذّر استخراج النص من ملف PDF. تأكد من أن الملف قابل للقراءة وأن PyMuPDF مثبتة.")
    elif ext == "docx":
        original_docx_bytes = raw
        input_text = extract_text_from_docx(raw)
        if input_text:
            st.success(f"✅ تم استخراج النص من ملف Word ({len(input_text.split()):,} كلمة)")
            with st.spinner("جارٍ تصدير ملف Word الأصلي إلى PDF مع الحفاظ على التنسيق والجداول والصور والأشكال..."):
                converted_pdf = convert_docx_to_pdf_bytes(raw)
                conversion_diag = get_last_docx_conversion_diagnostic()
            if converted_pdf:
                original_pdf_bytes = converted_pdf
                engine_name = conversion_diag.get("engine") or "محرك التحويل"
                st.info(
                    f"✅ تم تصدير Word الأصلي إلى PDF بواسطة {engine_name} دون أي تعديل مسبق. "
                    "سيُنشأ التقرير عليه بنفس غلاف وتظليل وملخص تقرير PDF تماماً."
                )
            else:
                st.error(
                    "تعذّر تحويل Word إلى PDF داخل بيئة الاستضافة، ولذلك لا يمكن إنشاء تقرير "
                    "مطابق لتقرير PDF مع الحفاظ على الجداول والصور والتنسيق."
                )
                st.warning(
                    conversion_diag.get("message")
                    or "على Streamlit Community Cloud أضف LibreOffice إلى packages.txt، وليس requirements.txt فقط."
                )
                details = (conversion_diag.get("details") or "").strip()
                if details:
                    with st.expander("التفاصيل الفنية للتحويل"):
                        st.code(details)
        else:
            st.error("تعذّر استخراج النص من ملف Word. تأكد من تثبيت python-docx.")


# ── Buttons ───────────────────────────────────────────────────────────────────
col_btn, col_clear = st.columns([3, 1])
with col_btn:
    analyze_btn = st.button("🔍 فحص الآن", use_container_width=True, type="primary")
with col_clear:
    if st.button("🗑️ مسح", use_container_width=True):
        st.session_state.pop("last_result", None)
        st.session_state.pop("last_text", None)
        st.session_state.pop("last_name", None)
        st.session_state.pop("last_pdf_bytes", None)
        st.session_state.pop("last_docx_bytes", None)
        st.rerun()

if analyze_btn and input_text and input_text.strip():
    _active_gemini_key = _get_gemini_api_key()
    if not _GEMINI_SERVICE_ENABLED:
        st.session_state.pop("last_result", None)
        st.error("خدمة المراجعة الذكية غير مفعلة حاليًا من إدارة التطبيق.")
    elif not _active_gemini_key:
        st.session_state.pop("last_result", None)
        st.error("خدمة المراجعة الذكية غير متاحة مؤقتًا. يرجى التواصل مع إدارة التطبيق.")
    elif not st.session_state.get("smart_review_consent", True):
        st.session_state.pop("last_result", None)
        st.error("يجب الموافقة على إرسال مقتطفات مختارة قبل بدء الفحص.")
    else:
        try:
            with st.spinner("جارٍ تنفيذ التحليل المشترك الإلزامي بالمحرك الداخلي والنموذج الذكي الثابت..."):
                result = _ai_fingerprint_score(input_text.strip())
                _gemini_source_text = _prepare_analysis_text(input_text.strip()).get(
                    "main_text_raw", result.get("analysis_text", input_text.strip())
                )
                result["gemini_review"] = _gemini_assisted_analysis(
                    _gemini_source_text,
                    result,
                    _active_gemini_key,
                    _GEMINI_MODEL,
                )
                result["combined_analysis"] = _combined_probability(result, result["gemini_review"])
                if not result["gemini_review"].get("available") or not result.get("combined_analysis"):
                    raise GeminiAPIError("لم يكتمل الدمج بين المحرك الداخلي والمراجعة الذكية.")
                result["joint_analysis_complete"] = True
        except InsufficientEnglishTextError as exc:
            st.session_state.pop("last_result", None)
            st.error(f"تعذّر الفحص: {exc}")
        except EnglishOnlyError:
            st.session_state.pop("last_result", None)
            st.error("هذا الإصدار مخصص للإنجليزية فقط. تم رفض الملف لأنه عربي أو مختلط أو غير إنجليزي، ولم تُصدر أي نسبة.")
        except GeminiAPIError as exc:
            st.session_state.pop("last_result", None)
            st.session_state.pop("last_text", None)
            st.session_state.pop("last_name", None)
            st.session_state.pop("last_pdf_bytes", None)
            st.session_state.pop("last_docx_bytes", None)
            st.error(f"فشل التحليل الذكي: {exc}")
            st.warning("لم يتم اعتماد النتيجة الداخلية منفردة، ولم يُنشأ التقرير لأن التحليل المشترك لم يكتمل.")
        except Exception as exc:
            st.session_state.pop("last_result", None)
            st.error(f"حدث خطأ أثناء تحليل البصمة: {exc}")
        else:
            st.session_state["last_result"]     = result
            st.session_state["last_text"]       = input_text.strip()
            st.session_state["last_name"]       = uploaded_name
            st.session_state["last_pdf_bytes"]  = original_pdf_bytes
            st.session_state["last_docx_bytes"] = original_docx_bytes

st.markdown('</div>', unsafe_allow_html=True)


# ── Results ───────────────────────────────────────────────────────────────────
# أي نتيجة من إصدار أقدم لا تحمل ختم اكتمال التحليل المشترك لا يجوز عرضها أو تنزيل تقريرها.
if "last_result" in st.session_state and not st.session_state["last_result"].get("joint_analysis_complete"):
    st.session_state.pop("last_result", None)
    st.session_state.pop("last_text", None)
    st.session_state.pop("last_name", None)
    st.session_state.pop("last_pdf_bytes", None)
    st.session_state.pop("last_docx_bytes", None)
    st.warning("تم إلغاء نتيجة غير مكتملة من جلسة سابقة. أعد الفحص لإتمام التحليل المشترك.")

if "last_result" in st.session_state:
    st.markdown("### النتيجة النهائية للفحص")
    r = st.session_state["last_result"]
    local_pct = float(r["percentage"])
    gemini_review = r.get("gemini_review") or {}
    combined = r.get("combined_analysis") or _combined_probability(r, gemini_review)
    pct = float(combined.get("score", local_pct))
    clr = combined.get("color") or r["color"]
    pct_label = _format_percentage(pct)
    meter_width = 0.0 if _low_score_is_masked(pct) else max(0.0, min(100.0, pct))
    main_verdict = combined.get("verdict_ar") or r["verdict"]
    confidence_label = combined.get("confidence_label", "غير محددة")

    st.markdown(f"""
    <div class="meter-wrap">
        <div style="font-size:.9rem;font-weight:800;color:#d9f5e6;margin-bottom:.35rem">النتيجة النهائية المعتمدة</div>
        <div class="meter-pct" style="color:{clr}">{pct_label}</div>
        <div class="meter-label" style="color:{clr}">{main_verdict}</div>
        <div class="meter-human">ثقة النتيجة المركبة: {confidence_label}</div>
        <div class="bar-track">
            <div class="bar-fill" style="width:{meter_width}%;background:{clr}"></div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    st.warning(
        "النسبة الظاهرة هي تقدير النص المحتمل صياغته أو تطويره بمساعدة الذكاء الاصطناعي بعد دمج نتائج الفحص. "
        "تُعرض *% ولا يعمل التظليل عندما تكون النتيجة من 0 إلى 20، وتظهر النسبة الرقمية ويبدأ التظليل من 21% فأعلى. "
        "النتيجة مؤشر فحص وليست إثباتًا قاطعًا لهوية الكاتب أو نسبة اقتباس علمي."
    )

    _pc1, _pc2, _pc3 = st.columns(3)
    _pc1.metric("المؤشر الداخلي", _format_percentage(combined.get("local_score", local_pct)))
    _pc2.metric("تقدير الذكاء الاصطناعي", _format_percentage(combined.get("gemini_score", 0)))
    _pc3.metric("مقدار الاختلاف", f"{combined.get('disagreement', 0):.0f} نقطة")
    if float(combined.get("disagreement", 0)) >= 30:
        st.info("يوجد اختلاف واضح بين المحركين؛ لذلك خُفضت ثقة النتيجة تلقائيًا ويُنصح بالمراجعة البشرية.")

    if gemini_review.get("available"):
        st.markdown("#### تفسير المراجعة الذكية")
        st.write(gemini_review.get("summary_ar", ""))
        _gc1, _gc2, _gc3, _gc4 = st.columns(4)
        _gc1.metric("نسبة المساعدة بالذكاء الاصطناعي", _format_percentage(gemini_review.get("ai_assistance_likelihood_score", 0)))
        _gc2.metric("مخاطر دعم الاستشهادات", f"{gemini_review.get('citation_integrity_risk_score', 0)}/100")
        _gc3.metric("المقاطع التي تمت مراجعتها", gemini_review.get("passages_reviewed", 0))
        _gc4.metric("المقاطع المحددة", len(gemini_review.get("flagged_passages") or []))

        if gemini_review.get("key_reasons_ar"):
            st.markdown("**أسباب رفع أولوية المراجعة:**")
            for _reason in gemini_review["key_reasons_ar"]:
                st.markdown(f"- {_reason}")
        if gemini_review.get("counter_evidence_ar"):
            st.markdown("**أدلة معاكسة تقلل الاشتباه:**")
            for _reason in gemini_review["counter_evidence_ar"]:
                st.markdown(f"- {_reason}")

        _flags = gemini_review.get("flagged_passages") or []
        if _flags:
            with st.expander("المقاطع التي حددتها المراجعة الذكية", expanded=True):
                for _flag in _flags:
                    st.markdown(
                        f"**{_flag.get('passage_id')} — شدة {_flag.get('severity')}/5**  \n"
                        f"{_flag.get('reason_ar')}"
                    )
                    st.code(_flag.get("quote", ""), language=None)

        if gemini_review.get("recommended_actions_ar"):
            with st.expander("إجراءات المراجعة المقترحة"):
                for _action in gemini_review["recommended_actions_ar"]:
                    st.markdown(f"- {_action}")

    with st.expander("المؤشر المحلي المساند — ليس النتيجة الرئيسية"):
        st.write(f"المؤشر الداخلي: {_format_percentage(local_pct)}")
        st.write(f"كفاية الأدلة المحلية: {r.get('confidence_percentage', 0):.0f}%")
        st.caption(f"وزن المؤشر الداخلي في النتيجة الحالية: {combined.get('local_weight', 0)*100:.0f}%، ووزن المراجع الذكي: {combined.get('gemini_weight', 0)*100:.0f}%؛ تتغير الأوزان حسب كفاية الأدلة والثقة.")
    if r.get("reference_words_excluded", 0) > 0:
        st.success(
            f"✅ تم استبعاد {r['reference_words_excluded']:,} كلمة من قسم المراجع "
            f"و{r.get('inline_citation_words_removed', 0):,} كلمة من علامات الاستشهاد قبل الحساب."
        )
    elif not r.get("reference_section_found", False):
        st.info("ℹ️ لم يُكتشف قسم مراجع مستقل، مع حذف الاستشهادات المضمنة التي أمكن التعرف عليها.")

    st.markdown(f"""
    <div class="stat-row">
        <div class="stat-box"><div class="stat-num">{r['n_words']:,}</div><div class="stat-lbl">كلمة إنجليزية</div></div>
        <div class="stat-box"><div class="stat-num">{r['n_sents']}</div><div class="stat-lbl">جملة</div></div>
        <div class="stat-box"><div class="stat-num">{r.get('chunk_count', 1)}</div><div class="stat-lbl">أقسام تحليلية</div></div>
        <div class="stat-box"><div class="stat-num">{r.get('evidence_families', 0)}/10</div><div class="stat-lbl">مجموعات أدلة</div></div>
        <div class="stat-box"><div class="stat-num">{r['reference_words_excluded']:,}</div><div class="stat-lbl">كلمة مراجع مستبعدة</div></div>
        <div class="stat-box"><div class="stat-num">{r['confidence_percentage']:.0f}%</div><div class="stat-lbl">كفاية الأدلة</div></div>
    </div>
    """, unsafe_allow_html=True)

    if (gemini_review.get("flagged_passages") or []):
        st.info(
            "أضافت المراجعة الذكية قائمة مستقلة بالمقاطع التي تحتاج مراجعة، وتم دمجها مباشرة مع تظليل المحرك الداخلي. "
            "ويجب قراءة التظليل مع الأسباب والسياق الأكاديمي الكامل، ولا يُعامل كإثبات اقتباس."
        )
    else:
        st.success("لم تحدد المراجعة الذكية مقاطع إضافية بعينها في العينة المرسلة.")

    st.divider()


    # ── PDF / DOCX Report Download ─────────────────────────────────────────────
    if RLAB_OK:
        st.markdown("### 📥 تحميل التقرير الكامل")

        saved_text                 = st.session_state.get("last_text", "")
        saved_name                 = st.session_state.get("last_name", "Document")
        saved_pdf_bytes            = st.session_state.get("last_pdf_bytes", None)
        saved_docx_bytes           = st.session_state.get("last_docx_bytes", None)

        if saved_text and saved_pdf_bytes:
            # ── تقرير واحد موحد: PDF أصلي أو PDF مُصدَّر من Word الأصلي ───
            is_converted_docx = saved_docx_bytes is not None
            label_spinner = (
                "جارٍ توليد تقرير PDF الموحد من ملف Word المُصدَّر مع الحفاظ على تخطيطه..."
                if is_converted_docx
                else "جارٍ توليد التقرير من ملف PDF الأصلي..."
            )
            with st.spinner(label_spinner):
                pdf_bytes = generate_report_pdf_from_original(
                    saved_pdf_bytes, saved_text, saved_name, analysis_result=r
                )
            if pdf_bytes:
                report_filename = f"AI_Report_{saved_name.replace('.', '_')}.pdf"
                btn_label = (
                    "⬇️  تحميل تقرير Word بصيغة PDF (مطابق لمسار تقرير PDF)"
                    if is_converted_docx
                    else "⬇️  تحميل التقرير المظلل من PDF الأصلي"
                )
                st.download_button(
                    label=btn_label,
                    data=pdf_bytes,
                    file_name=report_filename,
                    mime="application/pdf",
                    use_container_width=True,
                )

                no_highlight_note = (
                    " · النتيجة من 0 إلى 20 وتظهر *% دون تظليل"
                    if not _highlighting_allowed((r.get('combined_analysis') or {}).get('score', r.get('percentage', 0)))
                    else " · تظليل أزرق سماوي فوق جميع المقاطع المؤهلة للمراجعة"
                )
                info_detail = (
                    "✅ تقرير Word خرج عبر نفس مسار تقرير PDF: غلاف احترافي · صفحات Word بعد تصديرها كما هي "
                    "بجداولها وصورها وأشكالها ورؤوسها وتذييلاتها · المراجع مستبعدة · صفحة ملخص · خطة مراجعة ذكية في النهاية"
                    + no_highlight_note
                    if is_converted_docx
                    else
                    "✅ التقرير يحتوي على: صفحة غلاف · الصفحات الأصلية كما هي · المراجع بدون تظليل · ملخص الإشارات · خطة مراجعة ذكية في النهاية"
                    + no_highlight_note
                )
                st.markdown(
                    f"<div class='info-note'>{info_detail}</div>",
                    unsafe_allow_html=True,
                )
            else:
                st.warning("تعذّر توليد التقرير. تأكد من تثبيت PyMuPDF و reportlab.")

        elif saved_text and saved_docx_bytes:
            # ── حل أخير فقط إذا تعذرت كل مسارات الحفاظ على الأصل ────────────
            with st.spinner("جارٍ توليد التقرير النصي الاحتياطي..."):
                pdf_bytes = generate_report_pdf_text_only(saved_text, saved_name, analysis_result=r)
            if pdf_bytes:
                report_filename = f"AI_Report_{saved_name.replace('.', '_')}.pdf"
                st.download_button(
                    label="⬇️  تحميل التقرير النصي الاحتياطي",
                    data=pdf_bytes,
                    file_name=report_filename,
                    mime="application/pdf",
                    use_container_width=True,
                )
                st.markdown(
                    "<div class='info-note'>"
                    "⚠️ هذا المسار احتياطي فقط عند تعذّر الحفاظ على الأصل."
                    "</div>",
                    unsafe_allow_html=True,
                )
            else:
                st.warning("تعذّر توليد التقرير. تأكد من تثبيت reportlab.")
    else:
        st.info("💡 لتفعيل تصدير PDF: `pip install reportlab`")

    st.markdown(
        f"<div style='text-align:center;font-size:.82rem;opacity:.5;direction:ltr'>"
        f"FINAL SCORE: {pct_label} | Local evidence sufficiency: {r['confidence_percentage']:.0f}%"
        f"</div>",
        unsafe_allow_html=True,
    )

elif analyze_btn and not (input_text and input_text.strip()):
    st.warning("الرجاء رفع ملف PDF أو Word أولاً.")
